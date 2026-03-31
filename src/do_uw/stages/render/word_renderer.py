"""Word document assembly orchestrator.

Creates a complete D&O underwriting worksheet as a Word document,
delegating to section renderers for each major section. Handles
document setup: custom styles, margins, header/footer, TOC.

Phase 35-06: Density-aware rendering with pre-computed narratives.
Phase 60-01: Shared context layer — sections receive context dict
from build_template_context() instead of raw AnalysisState.
"""

from __future__ import annotations

import logging
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document  # type: ignore[import-untyped]
from docx.enum.section import WD_ORIENT  # type: ignore[import-untyped]
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
from docx.shared import Inches, Pt, RGBColor  # type: ignore[import-untyped]

from do_uw.models.density import DensityLevel
from do_uw.models.state import AnalysisState
from do_uw.stages.render.design_system import DesignSystem, setup_styles
from do_uw.stages.render.docx_helpers import (
    add_page_number,
    add_section_divider,
    add_toc_field,
)
from do_uw.brain.manifest_schema import load_manifest
from do_uw.stages.render.md_renderer import build_template_context

logger = logging.getLogger(__name__)

# Map manifest section IDs to (module_path, function_name) for Word rendering.
# Sections with None are not rendered as standalone Word sections
# (identity is the cover page, red_flags is embedded in executive summary,
# appendix sections like sources/qa_audit/coverage are HTML-only).
_SECTION_RENDERER_MAP: dict[str, tuple[str, str] | None] = {
    "identity": None,  # Cover page handled separately by _add_title_page
    "executive_summary": ("do_uw.stages.render.sections.sect1_executive", "render_section_1"),
    "red_flags": None,  # Embedded in executive summary for Word
    "company_operations": ("do_uw.stages.render.sections.sect2_company", "render_section_2"),
    "business_profile": ("do_uw.stages.render.sections.sect2_company", "render_section_2"),  # Legacy alias
    "financial_health": ("do_uw.stages.render.sections.sect3_financial", "render_section_3"),
    "market_activity": ("do_uw.stages.render.sections.sect4_market", "render_section_4"),
    "governance": ("do_uw.stages.render.sections.sect5_governance", "render_section_5"),
    "litigation": ("do_uw.stages.render.sections.sect6_litigation", "render_section_6"),
    "sector_industry": None,  # Placeholder — HTML-only for now (Phase 133)
    "scoring": ("do_uw.stages.render.sections.sect7_scoring", "render_section_7"),
    "sources": None,  # HTML-only appendix
    "qa_audit": None,  # HTML-only appendix
    "market_overflow": None,  # HTML-only appendix
    "coverage": None,  # HTML-only appendix
}

# Map manifest section IDs to Word display names for density/narrative lookup
_SECTION_DISPLAY_NAMES: dict[str, str] = {
    "executive_summary": "Section 1: Executive Summary",
    "company_operations": "Section 2: Company & Operations",
    "business_profile": "Section 2: Company Profile",
    "financial_health": "Section 3: Financial Health",
    "market_activity": "Section 4: Market & Trading",
    "governance": "Section 5: Governance & Leadership",
    "litigation": "Section 6: Litigation & Regulatory",
    "sector_industry": "Section 7: Sector & Industry Analysis",
    "scoring": "Section 8: Scoring & Risk Assessment",
}

# Map section display names to density section IDs and narrative field names
_SECTION_DENSITY_MAP: dict[str, str] = {
    "Section 1: Executive Summary": "executive_summary",
    "Section 2: Company & Operations": "company",
    "Section 2: Company Profile": "company",
    "Section 3: Financial Health": "financial",
    "Section 4: Market & Trading": "market",
    "Section 5: Governance & Leadership": "governance",
    "Section 6: Litigation & Regulatory": "litigation",
    "Section 7: Sector & Industry Analysis": "sector_industry",
    "Section 8: Scoring & Risk Assessment": "scoring",
}

# Colors for density indicators (amber for ELEVATED, red for CRITICAL)
_DENSITY_COLORS: dict[str, tuple[int, int, int]] = {
    DensityLevel.ELEVATED: (0xFF, 0xB8, 0x00),  # Amber
    DensityLevel.CRITICAL: (0xCC, 0x00, 0x00),  # Red
}

_DENSITY_LABELS: dict[str, str] = {
    DensityLevel.ELEVATED: "ELEVATED CONCERN",
    DensityLevel.CRITICAL: "CRITICAL RISK",
}


# Section renderer type: callable(doc, context, ds, **kwargs) -> None
# Sections receive context dict from build_template_context().
# Section 4 (Market) also accepts chart_dir=Path.
SectionRenderer = Any


def _get_section_renderers() -> list[tuple[str, SectionRenderer | None]]:
    """Get available section renderers in manifest-declared order.

    Phase 76-02: Uses output manifest for section ordering instead of
    hardcoded list. Sections not in _SECTION_RENDERER_MAP are skipped
    (HTML-only appendices). Calibration Notes appended as system section.

    Validates that every manifest section ID is present in the renderer
    map to catch drift between the manifest and Word renderer.

    Returns list of (section_display_name, renderer_fn) tuples.
    Renderer is None if the section module is not yet implemented.
    """
    manifest = load_manifest()
    renderers: list[tuple[str, SectionRenderer | None]] = []

    # Validate: every manifest section must be accounted for in renderer map
    for section in manifest.sections:
        if section.id not in _SECTION_RENDERER_MAP:
            logger.warning(
                "Manifest section '%s' (%s) has no entry in _SECTION_RENDERER_MAP — "
                "add it as None (skip) or provide a renderer to keep Word/HTML in sync",
                section.id,
                section.name,
            )

    for section in manifest.sections:
        # Look up renderer for this manifest section
        if section.id not in _SECTION_RENDERER_MAP:
            continue

        renderer_info = _SECTION_RENDERER_MAP[section.id]
        if renderer_info is None:
            # Explicitly skipped section (identity cover, red_flags embedded, HTML-only)
            continue

        module_path, func_name = renderer_info
        display_name = _SECTION_DISPLAY_NAMES.get(section.id, section.name)
        renderer: SectionRenderer | None = _try_import_renderer(module_path, func_name)
        renderers.append((display_name, renderer))

    # Calibration Notes is a system section, not in the manifest
    cal_renderer: SectionRenderer | None = _try_import_renderer(
        "do_uw.stages.render.sections.sect_calibration", "render_calibration_section"
    )
    renderers.append(("Calibration Notes", cal_renderer))

    return renderers


def _try_import_renderer(
    module_path: str, func_name: str
) -> SectionRenderer | None:
    """Try to import a section renderer function.

    Returns None if the module or function doesn't exist yet.
    """
    try:
        import importlib

        module = importlib.import_module(module_path)
        return getattr(module, func_name, None)
    except (ImportError, ModuleNotFoundError):
        return None


def _add_header(doc: Any, ticker: str, company_name: str) -> None:
    """Add document header with company name and generation date."""
    section: Any = doc.sections[0]
    header: Any = section.header
    header.is_linked_to_previous = False

    paragraph: Any = header.paragraphs[0]
    paragraph.clear()
    run: Any = paragraph.add_run(f"{company_name} ({ticker})")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x14, 0x46)

    date_run: Any = paragraph.add_run(
        f"  |  Generated: {datetime.now(tz=UTC).strftime('%Y-%m-%d')}"
    )
    date_run.font.size = Pt(8)
    date_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _add_footer(doc: Any) -> None:
    """Add document footer with confidentiality notice and page number."""
    section: Any = doc.sections[0]
    footer: Any = section.footer
    footer.is_linked_to_previous = False

    paragraph: Any = footer.paragraphs[0]
    paragraph.clear()
    run: Any = paragraph.add_run("Angry Dolphin Underwriting | Confidential")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Add separator and page number
    sep_run: Any = paragraph.add_run("  |  Page ")
    sep_run.font.size = Pt(7)
    sep_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    add_page_number(paragraph)


def _set_page_margins(doc: Any, ds: DesignSystem) -> None:
    """Set page margins and orientation on the document."""
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = ds.page_margin
        section.bottom_margin = ds.page_margin
        section.left_margin = ds.page_margin
        section.right_margin = ds.page_margin


def _add_title_page(doc: Any, state: AnalysisState) -> None:
    """Add title / executive summary header to the document."""
    company_name = "Unknown Company"
    if state.company and state.company.identity:
        sv_name = state.company.identity.legal_name
        if sv_name is not None:
            company_name = str(sv_name.value)

    # Logo at top, centered
    logo_path = Path(__file__).parent.parent.parent / "assets" / "logo.png"
    if logo_path.exists():
        logo_para: Any = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.paragraph_format.space_after = Pt(12)
        logo_run: Any = logo_para.add_run()
        logo_run.add_picture(str(logo_path), width=Inches(1.75))

    # Title
    title_para: Any = doc.add_paragraph(style="DOHeading1")
    title_run: Any = title_para.add_run(
        "Angry Dolphin -- D&O Underwriting Worksheet"
    )
    _ = title_run  # style applied by paragraph style

    # Subtitle with company info
    subtitle: Any = doc.add_paragraph(style="DOHeading2")
    sub_run: Any = subtitle.add_run(f"{company_name} ({state.ticker})")
    _ = sub_run

    # Generation date
    date_para: Any = doc.add_paragraph(style="DOBody")
    date_run: Any = date_para.add_run(
        f"Generated: {datetime.now(tz=UTC).strftime('%B %d, %Y')}"
    )
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _add_density_indicator(
    doc: Any, section_name: str, state: AnalysisState
) -> None:
    """Add a colored density-tier indicator paragraph for a section.

    CLEAN: no indicator (clean is normal, no noise).
    ELEVATED: amber-colored "ELEVATED CONCERN" text.
    CRITICAL: red-colored "CRITICAL RISK" text.

    Args:
        doc: The python-docx Document.
        section_name: Display name like "Section 3: Financial Health".
        state: AnalysisState with section_densities populated.
    """
    section_id = _SECTION_DENSITY_MAP.get(section_name)
    if section_id is None:
        return  # Non-analytical sections (Calibration, Meeting Prep) skip

    if state.analysis is None:
        return

    density = state.analysis.section_densities.get(section_id)
    if density is None:
        return

    level = density.level
    if level == DensityLevel.CLEAN:
        return  # No indicator for clean sections

    label = _DENSITY_LABELS.get(level)
    color = _DENSITY_COLORS.get(level)
    if label is None or color is None:
        return

    para: Any = doc.add_paragraph(style="DOBody")
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(4)
    run: Any = para.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*color)


def _add_section_narrative(
    doc: Any, section_name: str, state: AnalysisState
) -> None:
    """Add a pre-computed narrative paragraph for a section (OUT-03).

    If pre_computed_narratives has a narrative for this section, render it
    as the section's opening paragraph in formal research report voice.

    If no pre-computed narrative is available, does nothing (section
    renderers still generate their own content).

    Args:
        doc: The python-docx Document.
        section_name: Display name like "Section 3: Financial Health".
        state: AnalysisState with pre_computed_narratives populated.
    """
    section_id = _SECTION_DENSITY_MAP.get(section_name)
    if section_id is None:
        return

    if state.analysis is None or state.analysis.pre_computed_narratives is None:
        return

    narratives = state.analysis.pre_computed_narratives
    narrative_text = getattr(narratives, section_id, None)
    if not narrative_text:
        return

    para: Any = doc.add_paragraph(style="DOBody")
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(6)

    # Narrative body
    body_run: Any = para.add_run(narrative_text)
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_placeholder_section(
    doc: Any, section_name: str, ds: DesignSystem
) -> None:
    """Add a placeholder paragraph for an unimplemented section."""
    para: Any = doc.add_paragraph(style="DOHeading2")
    para.add_run(section_name)

    body: Any = doc.add_paragraph(style="DOBody")
    run: Any = body.add_run(f"{section_name} -- To be implemented")
    run.font.color.rgb = ds.color_text_light
    run.italic = True


def render_word_document(
    state: AnalysisState, output_path: Path, ds: DesignSystem,
    chart_dir: Path | None = None,
) -> Path:
    """Render a complete D&O underwriting worksheet as a Word document.

    Creates the document with custom styles, margins, header/footer,
    TOC, and delegates to section renderers. Sections not yet
    implemented get placeholder paragraphs.

    Args:
        state: The complete AnalysisState with all pipeline data.
        output_path: Where to save the .docx file.
        ds: Design system with visual constants.

    Returns:
        The output_path where the document was saved.
    """
    doc: Any = Document()
    setup_styles(doc)
    _set_page_margins(doc, ds)

    # Resolve company name for header
    company_name = "Unknown Company"
    if state.company and state.company.identity:
        sv_name = state.company.identity.legal_name
        if sv_name is not None:
            company_name = str(sv_name.value)

    # Header and footer
    _add_header(doc, state.ticker, company_name)
    _add_footer(doc)

    # Title page
    _add_title_page(doc, state)

    # Table of Contents
    add_toc_field(doc)

    # Phase 60-01: Build shared context once, pass to all sections.
    # The context dict contains all extracted/computed data that sections need.
    # _state key provides backward-compat escape hatch for data not yet in
    # context_builders — goal is zero context["_state"] accesses by end of Phase 60.
    context = build_template_context(state, chart_dir=chart_dir)
    context["_state"] = state

    # Render each section with gold dividers between major sections
    renderers = _get_section_renderers()
    for section_name, renderer_fn in renderers:
        add_section_divider(doc)

        # Phase 35-06: Add density indicator and pre-computed narrative
        # before the section's data tables/content (OUT-03 compliance)
        _add_density_indicator(doc, section_name, state)
        _add_section_narrative(doc, section_name, state)

        if renderer_fn is not None:
            try:
                # Section 4 (Market) uses chart_dir for pre-generated PNG embedding
                if section_name == "Section 4: Market & Trading" and chart_dir is not None:
                    renderer_fn(doc, context, ds, chart_dir=chart_dir)
                else:
                    renderer_fn(doc, context, ds)
                logger.info("Rendered: %s", section_name)
            except Exception:
                logger.exception("Failed to render: %s", section_name)
                _add_placeholder_section(doc, section_name, ds)
        else:
            _add_placeholder_section(doc, section_name, ds)
            logger.info("Placeholder: %s (not yet implemented)", section_name)

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Saved Word document: %s", output_path)

    return output_path


__all__ = [
    "_add_density_indicator",
    "_add_section_narrative",
    "render_word_document",
]
