"""Helper utilities for python-docx document manipulation.

Provides table creation with styled headers, cell shading, page numbers,
table of contents, border control, and sourced paragraphs. All helpers
operate on untyped python-docx objects via Any annotations.
"""

from __future__ import annotations

from typing import Any, cast

from docx.oxml import OxmlElement  # type: ignore[import-untyped]
from docx.oxml.ns import qn  # type: ignore[import-untyped]
from docx.shared import Inches, Pt, RGBColor  # type: ignore[import-untyped]

from do_uw.stages.render.design_system import DesignSystem


def _oxml(tag: str) -> Any:
    """Create an OxmlElement with Any type for pyright strict compliance."""
    return cast(Any, OxmlElement(tag))


def set_cell_shading(cell: Any, hex_color: str) -> None:
    """Apply background shading to a table cell.

    Creates a NEW OxmlElement each time (do NOT reuse across cells).

    Args:
        cell: A python-docx table cell.
        hex_color: Hex color string WITHOUT '#' prefix (e.g., "1A1446").
    """
    shading_elm: Any = _oxml("w:shd")
    shading_elm.set(qn("w:fill"), hex_color)
    shading_elm.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell: Any, **kwargs: dict[str, str]) -> None:
    """Set individual cell borders.

    Args:
        cell: A python-docx table cell.
        **kwargs: Border specs keyed by position (top, bottom, left, right).
            Each value is a dict with keys: sz (size), val (style), color.
            Example: top={"sz": "6", "val": "single", "color": "1A1446"}
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders: Any = _oxml("w:tcBorders")

    for edge, attrs in kwargs.items():
        element: Any = _oxml(f"w:{edge}")
        for attr_key, attr_val in attrs.items():
            element.set(qn(f"w:{attr_key}"), str(attr_val))
        tc_borders.append(element)

    tc_pr.append(tc_borders)


def set_table_column_widths(table: Any, widths: list[Inches]) -> None:
    """Set column widths on a table.

    Args:
        table: A python-docx Table.
        widths: List of Inches values, one per column.
    """
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = width


def add_styled_table(
    doc: Any,
    headers: list[str],
    rows: list[list[str]],
    ds: DesignSystem,
) -> Any:
    """Create a table with navy header row and alternating shading.

    Args:
        doc: The python-docx Document.
        headers: Column header strings.
        rows: List of row data (each row is a list of cell strings).
        ds: Design system for colors and fonts.

    Returns:
        The created Table object.
    """
    table: Any = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row: navy background, white bold text
    header_row: Any = table.rows[0]
    for idx, header_text in enumerate(headers):
        cell: Any = header_row.cells[idx]
        set_cell_shading(cell, ds.header_bg)
        paragraph: Any = cell.paragraphs[0]
        paragraph.clear()
        run: Any = paragraph.add_run(header_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = ds.font_body
        run.font.size = Pt(9)

    # Data rows: alternating shading
    for row_idx, row_data in enumerate(rows):
        row_obj: Any = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row_obj.cells[col_idx]
            if row_idx % 2 == 1:
                set_cell_shading(cell, ds.row_alt)
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            run = paragraph.add_run(cell_text)
            run.font.name = ds.font_body
            run.font.size = Pt(9)
            run.font.color.rgb = ds.color_text

    return table


def add_data_table(
    doc: Any,
    headers: list[str],
    rows: list[list[str]],
    ds: DesignSystem,
    col_widths: list[float] | None = None,
) -> Any:
    """Create a data table with optional column width control.

    Like add_styled_table but with numeric column right-alignment
    and optional column widths.

    Args:
        doc: The python-docx Document.
        headers: Column header strings.
        rows: List of row data.
        ds: Design system for colors and fonts.
        col_widths: Optional list of column widths in inches.

    Returns:
        The created Table object.
    """
    table: Any = add_styled_table(doc, headers, rows, ds)

    # Apply column widths if specified
    if col_widths is not None:
        widths = [Inches(w) for w in col_widths]
        set_table_column_widths(table, widths)

    # Right-align numeric columns (heuristic: header contains $, %, #, or "Score")
    numeric_cols: set[int] = set()
    for idx, hdr in enumerate(headers):
        if any(marker in hdr for marker in ("$", "%", "#", "Score", "Count", "Rate")):
            numeric_cols.add(idx)

    if numeric_cols:
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]

        for row_obj in table.rows:
            for col_idx in numeric_cols:
                if col_idx < len(row_obj.cells):
                    cell_para: Any = row_obj.cells[col_idx].paragraphs[0]
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return table


def add_risk_indicator(paragraph: Any, level: str, ds: DesignSystem) -> None:
    """Add a colored risk level tag to a paragraph.

    Adds text like "[HIGH]" in the risk heat spectrum color.

    Args:
        paragraph: A python-docx Paragraph.
        level: Risk level string (CRITICAL, HIGH, ELEVATED, etc.).
        ds: Design system for risk colors.
    """
    from do_uw.stages.render.design_system import get_risk_color

    color_hex = get_risk_color(level)
    run: Any = paragraph.add_run(f" [{level.upper()}]")
    run.bold = True
    run.font.size = Pt(9)
    # Convert hex string to RGB
    r = int(color_hex[1:3], 16)
    g = int(color_hex[3:5], 16)
    b = int(color_hex[5:7], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    _ = ds  # ds reserved for future use


def add_page_number(paragraph: Any) -> None:
    """Add an auto-updating page number field to a paragraph.

    Uses fldChar + instrText with PAGE instruction for Word to
    auto-update on open.

    Args:
        paragraph: A python-docx Paragraph (typically in footer).
    """
    run: Any = paragraph.add_run()

    # Begin field char
    fld_char_begin: Any = _oxml("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    # Field instruction
    instr_text: Any = _oxml("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    run._r.append(instr_text)

    # Separate field char
    fld_char_separate: Any = _oxml("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)

    # End field char
    fld_char_end: Any = _oxml("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def add_toc_field(doc: Any) -> None:
    """Insert a Table of Contents field code.

    Word will regenerate the TOC when the document is opened.
    Uses OxmlElement with TOC instruction.

    Args:
        doc: The python-docx Document.
    """
    paragraph: Any = doc.add_paragraph()
    run: Any = paragraph.add_run()

    # Begin complex field
    fld_char_begin: Any = _oxml("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    # TOC instruction
    instr_text: Any = _oxml("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r' TOC \o "1-3" \h \z \u '
    run._r.append(instr_text)

    # Separate
    fld_char_separate: Any = _oxml("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)

    # Placeholder text
    placeholder_run: Any = paragraph.add_run("[Table of Contents -- Update in Word]")
    placeholder_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    placeholder_run.font.size = Pt(9)

    # End complex field
    end_run: Any = paragraph.add_run()
    fld_char_end: Any = _oxml("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_char_end)


def add_sourced_paragraph(
    doc: Any, text: str, citation: str, ds: DesignSystem
) -> Any:
    """Add a paragraph with main text and a smaller citation suffix.

    Args:
        doc: The python-docx Document.
        text: Main paragraph text.
        citation: Citation string (e.g., "[SEC 10-K, 2024-12-31, HIGH]").
        ds: Design system for fonts and colors.

    Returns:
        The created Paragraph.
    """
    paragraph: Any = doc.add_paragraph(style="DOBody")
    run: Any = paragraph.add_run(text)
    run.font.name = ds.font_body
    run.font.size = ds.size_body
    run.font.color.rgb = ds.color_text

    if citation:
        cite_run: Any = paragraph.add_run(f"  {citation}")
        cite_run.font.name = ds.font_mono
        cite_run.font.size = ds.size_small
        cite_run.font.color.rgb = ds.color_text_light

    return paragraph


def add_section_divider(doc: Any) -> None:
    """Add a gold accent divider line before a major section.

    Creates a thin gold horizontal rule as a visual separator between
    major sections of the document.

    Args:
        doc: The python-docx Document.
    """
    paragraph: Any = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    # Add a bottom border to the paragraph using gold color
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders: Any = _oxml("w:pBdr")
    bottom: Any = _oxml("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")  # 0.75pt line
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "FFD000")  # AD Gold
    p_borders.append(bottom)
    p_pr.append(p_borders)


__all__ = [
    "add_data_table",
    "add_page_number",
    "add_risk_indicator",
    "add_section_divider",
    "add_sourced_paragraph",
    "add_styled_table",
    "add_toc_field",
    "set_cell_border",
    "set_cell_shading",
    "set_table_column_widths",
]
