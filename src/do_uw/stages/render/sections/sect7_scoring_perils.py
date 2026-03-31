"""Section 7 peril-organized scoring: summary table and deep dive sections.

Renders brain framework peril assessment as a structured scoring overview:
- Peril summary table: 8-row overview with risk level and active chain counts
- Per-peril deep dives: causal chain narratives with trigger/amplifier/mitigator
  format for each active peril

Integrates with scoring_peril_data.py which extracts and cross-references
brain perils/chains with pipeline signal_results.
"""

from __future__ import annotations

from typing import Any

from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

from do_uw.stages.render.design_system import DesignSystem
from do_uw.stages.render.docx_helpers import (
    add_risk_indicator,
    set_cell_shading,
)

# Risk level -> hex shading color (no '#' prefix, for cell shading)
_RISK_SHADING: dict[str, str] = {
    "HIGH": "FCE8E6",       # Light red
    "ELEVATED": "FFCCBC",   # Orange-medium
    "MODERATE": "FFF3CD",   # Light amber
    "LOW": "DCEEF8",        # Light blue
}

# Risk level -> text color for emphasis
_RISK_TEXT_COLORS: dict[str, RGBColor] = {
    "HIGH": RGBColor(0xCC, 0x00, 0x00),      # Dark red
    "ELEVATED": RGBColor(0xCC, 0x33, 0x00),   # Red-orange
    "MODERATE": RGBColor(0xE6, 0x73, 0x00),   # Orange
    "LOW": RGBColor(0x4A, 0x90, 0xD9),        # Blue
}


def _risk_shading_color(level: str) -> str:
    """Map risk level to hex shading color. Defaults to neutral gray."""
    return _RISK_SHADING.get(level.upper(), "F2F4F8")


def _risk_text_color(level: str) -> RGBColor:
    """Map risk level to text color. Defaults to standard body color."""
    return _RISK_TEXT_COLORS.get(level.upper(), RGBColor(0x33, 0x33, 0x33))


# ---------------------------------------------------------------------------
# Peril Summary Table
# ---------------------------------------------------------------------------


def render_peril_summary(
    doc: Any, peril_data: dict[str, Any], ds: DesignSystem
) -> None:
    """Render D&O Peril Summary table -- overview of all 8 perils.

    Shows a color-coded summary of each peril's risk level, active
    causal chain count, and key evidence snippets. Similar style to
    the plaintiff heat map in sect7_peril_map.py.

    Args:
        doc: python-docx Document to render into.
        peril_data: Dict from extract_peril_scoring() with all_perils,
            active_count, perils (active only), highest_peril.
        ds: Design system for visual constants.
    """
    heading: Any = doc.add_paragraph(style="DOHeading3")
    heading.add_run("D&O Claim Peril Assessment")

    all_perils = peril_data.get("all_perils", [])
    if not all_perils:
        para: Any = doc.add_paragraph(style="DOBody")
        para.add_run("Peril assessment not available.")
        return

    # Summary stats
    active_count = peril_data.get("active_count", 0)
    total_count = len(all_perils)
    intro: Any = doc.add_paragraph(style="DOBody")
    intro.add_run(
        f"{active_count} of {total_count} D&O claim perils show active risk signals. "
        f"Active perils have at least one causal chain with triggered checks."
    )

    # Build color-coded table with cell shading per risk level
    headers = ["Peril", "Risk Level", "Active Chains", "Key Evidence"]
    n_rows = len(all_perils)
    table: Any = doc.add_table(rows=1 + n_rows, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    header_row: Any = table.rows[0]
    for idx, hdr_text in enumerate(headers):
        cell: Any = header_row.cells[idx]
        set_cell_shading(cell, ds.header_bg)
        paragraph: Any = cell.paragraphs[0]
        paragraph.clear()
        run: Any = paragraph.add_run(hdr_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = ds.font_body
        run.font.size = Pt(9)

    # Data rows with risk-level color coding
    for row_idx, peril in enumerate(all_perils):
        row_obj: Any = table.rows[row_idx + 1]
        risk = peril.get("risk_level", "LOW")
        chains_text = (
            f"{peril.get('active_chain_count', 0)}"
            f"/{peril.get('total_chain_count', 0)}"
        )
        evidence_items = peril.get("key_evidence", [])
        evidence = "; ".join(evidence_items[:2]) or "\u2014"
        # Col 0: Peril name
        cell_name: Any = row_obj.cells[0]
        para_name: Any = cell_name.paragraphs[0]
        para_name.clear()
        run_name: Any = para_name.add_run(peril.get("name", "Unknown"))
        run_name.font.name = ds.font_body
        run_name.font.size = Pt(9)
        run_name.bold = True
        run_name.font.color.rgb = ds.color_text

        # Col 1: Risk level (color-coded)
        cell_risk: Any = row_obj.cells[1]
        set_cell_shading(cell_risk, _risk_shading_color(risk))
        para_risk: Any = cell_risk.paragraphs[0]
        para_risk.clear()
        run_risk: Any = para_risk.add_run(risk)
        run_risk.font.name = ds.font_body
        run_risk.font.size = Pt(9)
        run_risk.bold = True
        run_risk.font.color.rgb = _risk_text_color(risk)

        # Col 2: Active chains
        cell_chains: Any = row_obj.cells[2]
        para_chains: Any = cell_chains.paragraphs[0]
        para_chains.clear()
        run_chains: Any = para_chains.add_run(chains_text)
        run_chains.font.name = ds.font_body
        run_chains.font.size = Pt(9)

        # Col 3: Key evidence
        cell_ev: Any = row_obj.cells[3]
        para_ev: Any = cell_ev.paragraphs[0]
        para_ev.clear()
        run_ev: Any = para_ev.add_run(evidence)
        run_ev.font.name = ds.font_body
        run_ev.font.size = Pt(9)
        run_ev.font.color.rgb = ds.color_text_light


# ---------------------------------------------------------------------------
# Per-Peril Deep Dives
# ---------------------------------------------------------------------------


def render_peril_deep_dives(
    doc: Any, peril_data: dict[str, Any], ds: DesignSystem
) -> None:
    """Render deep dive for each active peril with causal chain narratives.

    Only active perils (those with at least one triggered chain) get
    deep dive sections. Inactive perils are shown only in the summary
    table above.

    Args:
        doc: python-docx Document to render into.
        peril_data: Dict from extract_peril_scoring().
        ds: Design system for visual constants.
    """
    active_perils = peril_data.get("perils", [])
    if not active_perils:
        return

    heading: Any = doc.add_paragraph(style="DOHeading3")
    heading.add_run("Active Peril Analysis")

    for peril in active_perils:
        _render_single_peril(doc, peril, ds)


def _render_single_peril(
    doc: Any, peril: dict[str, Any], ds: DesignSystem
) -> None:
    """Render a single peril with its active causal chains.

    Includes peril header with risk indicator, frequency/severity context,
    and each active chain rendered as trigger/amplifier/mitigator narrative.
    """
    # Peril header with risk level indicator
    risk_level = peril.get("risk_level", "LOW")
    sub_heading: Any = doc.add_paragraph(style="DOHeading3")
    name_text = peril.get("name", "Unknown Peril")
    sub_heading.add_run(f"{name_text}: {risk_level}")
    add_risk_indicator(sub_heading, risk_level, ds)

    # Context line: frequency | severity | typical settlement range
    freq = peril.get("frequency", "unknown")
    sev = peril.get("severity", "unknown")
    settlement = peril.get("typical_settlement_range", "")
    ctx: Any = doc.add_paragraph(style="DOBody")
    ctx_text = f"Frequency: {freq} | Severity: {sev}"
    if settlement:
        ctx_text += f" | Typical range: {settlement}"
    ctx_run: Any = ctx.add_run(ctx_text)
    ctx_run.font.size = ds.size_small
    ctx_run.font.color.rgb = ds.color_text_light

    # Active chain count summary
    chains = peril.get("chains", [])
    active_chains = [c for c in chains if c.get("active")]
    total_chains = len(chains)
    summary: Any = doc.add_paragraph(style="DOBody")
    summary.add_run(
        f"{len(active_chains)} of {total_chains} causal chains active."
    )

    # Each active chain
    for chain in chains:
        if not chain.get("active"):
            continue
        _render_chain_narrative(doc, chain, ds)


def _render_chain_narrative(
    doc: Any, chain: dict[str, Any], ds: DesignSystem
) -> None:
    """Render a causal chain as trigger -> amplifier -> mitigator narrative.

    Shows chain name, description, then structured lists of triggered
    checks, active amplifiers, and active mitigators with evidence.
    """
    # Chain name (bold) + risk level
    chain_para: Any = doc.add_paragraph(style="DOBody")
    name_run: Any = chain_para.add_run(
        f"Chain: {chain.get('name', 'Unknown')}"
    )
    name_run.bold = True
    chain_risk = chain.get("risk_level", "")
    if chain_risk:
        chain_para.add_run(f" [{chain_risk}]")

    # Description (italic, smaller)
    description = chain.get("description", "")
    if description:
        desc_para: Any = doc.add_paragraph(style="DOBody")
        desc_run: Any = desc_para.add_run(description)
        desc_run.italic = True
        desc_run.font.size = ds.size_small

    # Build evidence map for trigger annotations
    evidence_map = chain.get("evidence_map", {})

    # Triggers (what started it)
    triggers = chain.get("triggered_triggers", [])
    if triggers:
        t_para: Any = doc.add_paragraph(style="DOBody")
        t_label: Any = t_para.add_run("  Triggers: ")
        t_label.bold = True
        t_label.font.size = ds.size_body
        for t in triggers:
            evidence = evidence_map.get(t, "")
            line = f"{t}"
            if evidence:
                line += f" \u2014 {evidence}"
            t_para.add_run(f"\n    \u2022 {line}")

    # Amplifiers (what makes it worse)
    amplifiers = chain.get("active_amplifiers", [])
    if amplifiers:
        a_para: Any = doc.add_paragraph(style="DOBody")
        a_label: Any = a_para.add_run("  Amplifiers: ")
        a_label.bold = True
        a_label.font.size = ds.size_body
        for a in amplifiers:
            a_para.add_run(f"\n    \u2022 {a}")

    # Mitigators (what helps)
    mitigators = chain.get("active_mitigators", [])
    if mitigators:
        m_para: Any = doc.add_paragraph(style="DOBody")
        m_label: Any = m_para.add_run("  Mitigators: ")
        m_label.bold = True
        m_label.font.size = ds.size_body
        for m in mitigators:
            m_para.add_run(f"\n    \u2022 {m}")

    # Historical context (filing rate, median severity)
    filing_rate = chain.get("historical_filing_rate")
    median_sev = chain.get("median_severity_usd")
    if filing_rate is not None or median_sev is not None:
        hist: Any = doc.add_paragraph(style="DOBody")
        parts: list[str] = []
        if filing_rate is not None:
            parts.append(f"Historical filing rate: {filing_rate * 100:.0f}%")
        if median_sev is not None:
            parts.append(f"Median severity: ${median_sev / 1e6:.0f}M")
        hist_run: Any = hist.add_run("  " + " | ".join(parts))
        hist_run.font.size = ds.size_small
        hist_run.font.color.rgb = ds.color_text_light


__all__ = ["render_peril_deep_dives", "render_peril_summary"]
