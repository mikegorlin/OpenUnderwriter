"""Section 1-08: Market Context renderer.

Displays market pricing intelligence from the BENCHMARK stage.
Shows peer segment pricing, rate-on-line ranges, trend direction,
and mispricing alerts when data is available. Gracefully handles
missing pricing data with a brief notice.

Data path: state.executive_summary.deal_context.market_intelligence

Phase 60-01: Migrated from state access to shared context dict.
"""

from __future__ import annotations

from typing import Any

from do_uw.models.executive_summary import MarketIntelligence
from do_uw.stages.render.design_system import DesignSystem
from do_uw.stages.render.docx_helpers import add_styled_table


def render_market_context(
    doc: Any, context: dict[str, Any], ds: DesignSystem
) -> None:
    """Render Market Context section within the Executive Summary.

    Reads MarketIntelligence from deal_context. When no pricing data
    exists, renders a brief notice and returns. When data is available,
    renders a summary paragraph, a metrics table, and any mispricing
    alerts.

    Args:
        doc: The python-docx Document.
        context: Shared context dict with _state for backward compat.
        ds: Design system for colors and fonts.
    """
    mi = _get_market_intelligence(context)

    if mi is None or not mi.has_data:
        _render_no_data(doc, mi)
        return

    _render_heading(doc)
    _render_summary(doc, mi, ds)
    _render_metrics_table(doc, mi, ds)
    _render_alerts(doc, mi, ds)


def _get_market_intelligence(
    context: dict[str, Any],
) -> MarketIntelligence | None:
    """Extract MarketIntelligence from context, returning None if absent."""
    state = context["_state"]
    if state.executive_summary is None:
        return None
    deal = state.executive_summary.deal_context
    return deal.market_intelligence


def _render_no_data(doc: Any, mi: MarketIntelligence | None) -> None:
    """Render brief notice when market pricing data is unavailable."""
    para: Any = doc.add_paragraph(style="DOHeading3")
    para.add_run("Market Context")
    body: Any = doc.add_paragraph(style="DOBody")
    if mi is not None and mi.peer_count > 0:
        body.add_run(
            f"Market pricing data insufficient "
            f"({mi.peer_count} peers found, confidence: {mi.confidence_level})."
        )
    else:
        body.add_run("No market pricing data available.")


def _render_heading(doc: Any) -> None:
    """Add Market Context heading."""
    para: Any = doc.add_paragraph(style="DOHeading3")
    para.add_run("Market Context")


def _render_summary(doc: Any, mi: MarketIntelligence, ds: DesignSystem) -> None:
    """Render summary paragraph describing market position."""
    _ = ds
    parts: list[str] = []

    if mi.segment_label:
        parts.append(f"Segment: {mi.segment_label}.")

    parts.append(
        f"Based on {mi.peer_count} comparable quotes "
        f"(confidence: {mi.confidence_level})."
    )

    if mi.trend_direction and mi.trend_direction != "INSUFFICIENT_DATA":
        trend_desc = mi.trend_direction.lower().replace("_", " ")
        mag = (
            f" ({mi.trend_magnitude_pct:+.1f}%)"
            if mi.trend_magnitude_pct is not None
            else ""
        )
        parts.append(f"Market trend: {trend_desc}{mag}.")

    if mi.data_window:
        parts.append(f"Data window: {mi.data_window}.")

    body: Any = doc.add_paragraph(style="DOBody")
    body.add_run(" ".join(parts))


def _render_metrics_table(
    doc: Any, mi: MarketIntelligence, ds: DesignSystem
) -> None:
    """Render pricing metrics table."""
    rows: list[list[str]] = []

    if mi.median_rate_on_line is not None:
        rows.append([
            "Median Rate-on-Line",
            f"{mi.median_rate_on_line:.4f}",
        ])

    if mi.ci_low is not None and mi.ci_high is not None:
        rows.append([
            "95% Confidence Interval",
            f"{mi.ci_low:.4f} - {mi.ci_high:.4f}",
        ])

    if mi.trend_direction and mi.trend_direction != "INSUFFICIENT_DATA":
        trend_val = mi.trend_direction
        if mi.trend_magnitude_pct is not None:
            trend_val = f"{mi.trend_direction} ({mi.trend_magnitude_pct:+.1f}%)"
        rows.append(["Market Trend", trend_val])

    rows.append(["Peer Count", str(mi.peer_count)])
    rows.append(["Confidence", mi.confidence_level])

    if mi.segment_label:
        rows.append(["Segment", mi.segment_label])

    if rows:
        add_styled_table(doc, ["Metric", "Value"], rows, ds)


def _render_alerts(doc: Any, mi: MarketIntelligence, ds: DesignSystem) -> None:
    """Render mispricing alerts if any are present."""
    alerts: list[str] = []
    if mi.mispricing_alert:
        alerts.append(mi.mispricing_alert)
    if mi.model_vs_market_alert:
        alerts.append(mi.model_vs_market_alert)

    if not alerts:
        return

    from docx.shared import RGBColor  # type: ignore[import-untyped]

    heading: Any = doc.add_paragraph(style="DOBody")
    run: Any = heading.add_run("Mispricing Signals")
    run.bold = True
    run.font.color.rgb = ds.color_primary

    for alert_text in alerts:
        para: Any = doc.add_paragraph(style="DOBody")
        alert_run: Any = para.add_run(alert_text)
        alert_run.font.color.rgb = RGBColor(0xE6, 0x73, 0x00)  # Orange warning


__all__ = ["render_market_context"]
