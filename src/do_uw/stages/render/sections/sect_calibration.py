"""Calibration Notes section for Word document renderer.

Renders system intelligence status, recent calibration changes,
discovery findings, and pending feedback. Only renders when
calibration data is available.

Phase 60-03: Migrated from AnalysisState to shared context dict.
"""

from __future__ import annotations

import logging
from typing import Any

from docx.shared import Pt  # type: ignore[import-untyped]

from do_uw.stages.render.design_system import DesignSystem

logger = logging.getLogger(__name__)


def render_calibration_section(
    doc: Any, context: dict[str, Any], ds: DesignSystem
) -> None:
    """Render calibration notes section in Word document.

    Uses the markdown calibration notes renderer for content,
    then formats as Word paragraphs and tables.

    Args:
        doc: python-docx Document instance.
        context: Shared context dict from build_template_context().
        ds: DesignSystem for visual styling.
    """
    try:
        from do_uw.stages.render.context_builders import (
            render_calibration_notes,
        )

        # TODO(phase-60): render_calibration_notes takes AnalysisState directly
        state = context["_state"]
        notes_md = render_calibration_notes(state)
    except Exception:
        notes_md = ""

    if not notes_md:
        return

    # Parse the markdown and render as Word content
    _render_markdown_as_word(doc, notes_md, ds)


def _render_markdown_as_word(
    doc: Any, markdown: str, ds: DesignSystem
) -> None:
    """Convert simple markdown to Word paragraphs.

    Handles: ## headings, ### headings, - bullets, | tables,
    and plain text paragraphs.
    """
    lines = markdown.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        if line.startswith("## "):
            para: Any = doc.add_paragraph(style="DOHeading2")
            para.add_run(line[3:])
        elif line.startswith("### "):
            para = doc.add_paragraph(style="DOHeading3")
            para.add_run(line[4:])
        elif line.startswith("| ") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|--"):
            # Table: collect header + separator + rows
            headers = [h.strip() for h in line.split("|")[1:-1]]
            i += 1  # skip separator line
            i += 1
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().split("|")[1:-1]]
                rows.append(cells)
                i += 1
            _add_simple_table(doc, headers, rows, ds)
            continue  # already advanced i
        elif line.startswith("- "):
            para = doc.add_paragraph(style="DOBody")
            # Parse bold markers **text**
            _add_formatted_run(para, line[2:])
        else:
            para = doc.add_paragraph(style="DOBody")
            _add_formatted_run(para, line)

        i += 1


def _add_formatted_run(para: Any, text: str) -> None:
    """Add a run with simple **bold** formatting support."""
    parts = text.split("**")
    for idx, part in enumerate(parts):
        if not part:
            continue
        run: Any = para.add_run(part)
        if idx % 2 == 1:  # odd parts are bold
            run.bold = True


def _add_simple_table(
    doc: Any,
    headers: list[str],
    rows: list[list[str]],
    ds: DesignSystem,
) -> None:
    """Add a simple table to the Word document."""
    col_count = len(headers)
    table: Any = doc.add_table(
        rows=1 + len(rows), cols=col_count
    )
    table.style = "Table Grid"

    # Header row
    for j, header in enumerate(headers):
        cell: Any = table.rows[0].cells[j]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < col_count:
                cell = table.rows[i + 1].cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)


__all__ = ["render_calibration_section"]
