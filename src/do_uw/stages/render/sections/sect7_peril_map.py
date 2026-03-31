"""Section 7 peril map: heat map grid, bear cases, settlement, tower risk.

Renders Phase 27 peril assessment: 7x2 plaintiff heat map, bear case
narratives, DDL settlement prediction, and tower risk characterization.
"""

from __future__ import annotations

from typing import Any

from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

from do_uw.models.peril import BearCase, PerilMap, PlaintiffAssessment
from do_uw.stages.render.design_system import DesignSystem
from do_uw.stages.render.docx_helpers import (
    add_styled_table,
    set_cell_shading,
)
from do_uw.stages.render.formatters import format_currency

# Probability bands -> hex (no '#' prefix, for cell shading)
_PROB_COLORS: dict[str, str] = {
    "VERY_LOW": "DCEEF8",  # Light blue (matches design_system highlight_good)
    "LOW": "FFF3CD",  # Light amber
    "MODERATE": "FFE0B2",  # Orange-light
    "ELEVATED": "FFCCBC",  # Orange-medium
    "HIGH": "FCE8E6",  # Light red (matches design_system highlight_bad)
}

# Severity bands -> hex (no '#' prefix)
_SEV_COLORS: dict[str, str] = {
    "NUISANCE": "DCEEF8",  # Light blue
    "MINOR": "FFF3CD",  # Light amber
    "MODERATE": "FFE0B2",  # Orange-light
    "SIGNIFICANT": "FFCCBC",  # Orange-medium
    "SEVERE": "FCE8E6",  # Light red
}

# Probability band -> RGBColor for text emphasis
_PROB_TEXT_COLORS: dict[str, RGBColor] = {
    "VERY_LOW": RGBColor(0x4A, 0x90, 0xD9),  # Blue
    "LOW": RGBColor(0x99, 0x88, 0x00),  # Dark amber
    "MODERATE": RGBColor(0xE6, 0x73, 0x00),  # Orange
    "ELEVATED": RGBColor(0xCC, 0x33, 0x00),  # Red-orange
    "HIGH": RGBColor(0xCC, 0x00, 0x00),  # Red
}


def _band_color(band_value: str, color_map: dict[str, str]) -> str:
    """Map band string to hex shading color. Defaults to neutral gray."""
    return color_map.get(band_value.upper(), "F2F4F8")


def _format_plaintiff_type(plaintiff_type: str) -> str:
    """Format plaintiff type for display (e.g., SHAREHOLDERS -> Shareholders)."""
    return plaintiff_type.replace("_", " ").title()


def _render_plaintiff_heat_map(
    doc: Any,
    assessments: list[PlaintiffAssessment],
    ds: DesignSystem,
) -> None:
    """Render 7x2 heat map grid with color-coded probability/severity cells."""
    heading: Any = doc.add_paragraph(style="DOHeading3")
    heading.add_run("Plaintiff Peril Heat Map")

    if not assessments:
        para: Any = doc.add_paragraph(style="DOBody")
        para.add_run("No plaintiff assessments available.")
        return

    # Build table: header + N data rows
    headers = ["Plaintiff Type", "Probability", "Severity"]
    n_rows = len(assessments)
    table: Any = doc.add_table(rows=1 + n_rows, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    header_row: Any = table.rows[0]
    for idx, hdr_text in enumerate(headers):
        cell: Any = header_row.cells[idx]
        set_cell_shading(cell, ds.header_bg)
        paragraph: Any = cell.paragraphs[0]
        paragraph.clear()
        run: Any = paragraph.add_run(hdr_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = ds.font_body
        run.font.size = Pt(9)

    # Data rows
    for row_idx, assessment in enumerate(assessments):
        row_obj: Any = table.rows[row_idx + 1]

        # Col 0: Plaintiff type
        cell_type: Any = row_obj.cells[0]
        para_type: Any = cell_type.paragraphs[0]
        para_type.clear()
        run_type: Any = para_type.add_run(
            _format_plaintiff_type(assessment.plaintiff_type)
        )
        run_type.font.name = ds.font_body
        run_type.font.size = Pt(9)
        run_type.bold = True
        run_type.font.color.rgb = ds.color_text

        # Col 1: Probability (color-coded)
        cell_prob: Any = row_obj.cells[1]
        prob_color = _band_color(assessment.probability_band, _PROB_COLORS)
        set_cell_shading(cell_prob, prob_color)
        para_prob: Any = cell_prob.paragraphs[0]
        para_prob.clear()
        run_prob: Any = para_prob.add_run(
            assessment.probability_band.replace("_", " ")
        )
        run_prob.font.name = ds.font_body
        run_prob.font.size = Pt(9)
        run_prob.bold = True
        text_color = _PROB_TEXT_COLORS.get(
            assessment.probability_band.upper(),
            ds.color_text,
        )
        run_prob.font.color.rgb = text_color

        # Col 2: Severity (color-coded)
        cell_sev: Any = row_obj.cells[2]
        sev_color = _band_color(assessment.severity_band, _SEV_COLORS)
        set_cell_shading(cell_sev, sev_color)
        para_sev: Any = cell_sev.paragraphs[0]
        para_sev.clear()
        run_sev: Any = para_sev.add_run(
            assessment.severity_band.replace("_", " ")
        )
        run_sev.font.name = ds.font_body
        run_sev.font.size = Pt(9)
        run_sev.bold = True
        run_sev.font.color.rgb = ds.color_text

    # Notes for lenses with probability >= MODERATE
    elevated_bands = {"MODERATE", "ELEVATED", "HIGH"}
    for assessment in assessments:
        if assessment.probability_band.upper() in elevated_bands:
            top_finding = (
                assessment.key_findings[0]
                if assessment.key_findings
                else "Multiple contributing factors"
            )
            note_para: Any = doc.add_paragraph(style="DOBody")
            run_note: Any = note_para.add_run(
                f"{_format_plaintiff_type(assessment.plaintiff_type)}: "
                f"{assessment.triggered_signal_count} triggered checks. "
                f"Key: {top_finding}"
            )
            run_note.font.size = ds.size_small
            run_note.font.color.rgb = ds.color_text_light


def _render_bear_cases(
    doc: Any,
    bear_cases: list[BearCase],
    ds: DesignSystem,
) -> None:
    """Render bear case narratives: committee summary + evidence chain."""
    heading: Any = doc.add_paragraph(style="DOHeading3")
    heading.add_run("Bear Cases")

    if not bear_cases:
        para: Any = doc.add_paragraph(style="DOBody")
        para.add_run(
            "No bear cases identified -- analysis did not find sufficient "
            "evidence to construct litigation scenarios."
        )
        return

    for bear_case in bear_cases:
        # Sub-heading: theory label
        theory_label = bear_case.theory.replace("_", " ").title()
        sub_heading: Any = doc.add_paragraph(style="DOHeading3")
        sub_heading.add_run(f"Bear Case: {theory_label}")

        # Committee summary (bold, prominent)
        summary_para: Any = doc.add_paragraph(style="DOBody")
        summary_run: Any = summary_para.add_run(bear_case.committee_summary)
        summary_run.bold = True

        # Defense assessment (if available)
        if bear_case.defense_assessment:
            defense_para: Any = doc.add_paragraph(style="DOBody")
            defense_run: Any = defense_para.add_run(
                f"Defense: {bear_case.defense_assessment}"
            )
            defense_run.italic = True

        # Evidence chain
        if bear_case.evidence_chain:
            evidence_heading: Any = doc.add_paragraph(style="DOBody")
            ev_run: Any = evidence_heading.add_run("Evidence Chain")
            ev_run.bold = True
            ev_run.font.size = ds.size_body

            for idx, evidence in enumerate(bear_case.evidence_chain, 1):
                ev_para: Any = doc.add_paragraph(style="DOBody")
                ev_text = (
                    f"{idx}. [{evidence.signal_id}] "
                    f"{evidence.description}"
                )
                if evidence.source:
                    ev_text += f" (Source: {evidence.source})"
                ev_content_run: Any = ev_para.add_run(ev_text)
                ev_content_run.font.size = ds.size_small

        # Footer: probability and severity
        footer_para: Any = doc.add_paragraph(style="DOBody")
        footer_run: Any = footer_para.add_run(
            f"Probability: {bear_case.probability_band.replace('_', ' ')} | "
            f"Severity: {bear_case.severity_estimate.replace('_', ' ')}"
        )
        footer_run.font.size = ds.size_small
        footer_run.font.color.rgb = ds.color_text_light


def _render_settlement_summary(
    doc: Any,
    context: dict[str, Any],
    ds: DesignSystem,
) -> None:
    """Render settlement prediction (Phase 27 DDL or legacy severity scenarios)."""
    heading: Any = doc.add_paragraph(style="DOHeading3")
    heading.add_run("Settlement Prediction")

    # Try Phase 27 settlement prediction first
    settlement_pred = _get_settlement_prediction(context)
    if settlement_pred is not None:
        _render_ddl_settlement_table(doc, settlement_pred, ds)
        return

    # Fallback to legacy severity scenarios
    # TODO(phase-60): use context["scoring"] when available
    state = context.get("_state")
    if state is not None and state.scoring and state.scoring.severity_scenarios:
        ss = state.scoring.severity_scenarios
        if ss.scenarios:
            headers = [
                "Scenario", "DDL", "Settlement",
                "Defense Costs", "Total Exposure",
            ]
            rows: list[list[str]] = []
            for scenario in ss.scenarios:
                rows.append([
                    f"{scenario.label} ({scenario.percentile}th)",
                    format_currency(scenario.ddl_amount, compact=True),
                    format_currency(
                        scenario.settlement_estimate, compact=True
                    ),
                    format_currency(
                        scenario.defense_cost_estimate, compact=True
                    ),
                    format_currency(
                        scenario.total_exposure, compact=True
                    ),
                ])
            add_styled_table(doc, headers, rows, ds)

            # Methodology note
            note_para: Any = doc.add_paragraph(style="DOBody")
            note_run: Any = note_para.add_run(
                "Note: Settlement estimates based on market cap decline "
                "scenarios and historical settlement-to-DDL ratios. "
                "Percentile bands reflect industry empirical data."
            )
            note_run.font.size = ds.size_small
            note_run.font.color.rgb = ds.color_text_light
            note_run.italic = True
            return

    para: Any = doc.add_paragraph(style="DOBody")
    para.add_run("Settlement prediction not available.")


def _get_settlement_prediction(
    context: dict[str, Any],
) -> dict[str, Any] | None:
    """Extract settlement prediction dict from context."""
    # TODO(phase-60): use context["settlement_prediction"] when available
    state = context.get("_state")
    if state is not None and state.analysis and state.analysis.settlement_prediction:
        return state.analysis.settlement_prediction
    return None


def _render_ddl_settlement_table(
    doc: Any, pred: dict[str, Any], ds: DesignSystem
) -> None:
    """Render DDL-based settlement prediction table (Phase 27 format)."""
    scenarios = pred.get("scenarios", [])
    if not scenarios:
        para: Any = doc.add_paragraph(style="DOBody")
        para.add_run("Settlement prediction computed but no scenarios available.")
        return

    headers = [
        "Scenario", "DDL", "Settlement",
        "Defense Costs", "Total Exposure",
    ]
    rows: list[list[str]] = []
    for scenario in scenarios:
        label = scenario.get("label", "Unknown")
        percentile = scenario.get("percentile", "")
        display_name = f"{label} ({percentile}th)" if percentile else label
        rows.append([
            display_name,
            format_currency(scenario.get("ddl_amount", 0), compact=True),
            format_currency(
                scenario.get("settlement_estimate", 0), compact=True
            ),
            format_currency(
                scenario.get("defense_cost_estimate", 0), compact=True
            ),
            format_currency(
                scenario.get("total_exposure", 0), compact=True
            ),
        ])

    add_styled_table(doc, headers, rows, ds)

    # Market cap context
    market_cap = pred.get("market_cap", 0)
    if market_cap and market_cap > 0:
        ctx_para: Any = doc.add_paragraph(style="DOBody")
        ctx_run: Any = ctx_para.add_run(
            f"Based on market cap: {format_currency(market_cap, compact=True)}"
        )
        ctx_run.font.size = ds.size_small
        ctx_run.font.color.rgb = ds.color_text_light

    # Methodology note
    method = pred.get("methodology", "")
    note_para: Any = doc.add_paragraph(style="DOBody")
    note_text = (
        "Note: DDL-based settlement prediction using empirical "
        "settlement-to-DDL ratios calibrated from securities class action "
        "historical data. Defense costs estimated from case characteristics."
    )
    if method:
        note_text = f"Methodology: {method}"
    note_run: Any = note_para.add_run(note_text)
    note_run.font.size = ds.size_small
    note_run.font.color.rgb = ds.color_text_light
    note_run.italic = True


def _render_tower_characterization(
    doc: Any,
    context: dict[str, Any],
    ds: DesignSystem,
) -> None:
    """Render per-layer expected loss share (analytical, not prescriptive)."""
    heading: Any = doc.add_paragraph(style="DOHeading3")
    heading.add_run("Tower Risk Characterization")

    pred = _get_settlement_prediction(context)
    if pred is None:
        para: Any = doc.add_paragraph(style="DOBody")
        para.add_run("Tower characterization not available.")
        return

    tower_risk = pred.get("tower_risk_characterization", {})
    if not tower_risk:
        para = doc.add_paragraph(style="DOBody")
        para.add_run("Tower characterization not available.")
        return

    layers = tower_risk.get("layers", [])
    if not layers:
        para = doc.add_paragraph(style="DOBody")
        para.add_run("Tower characterization not available.")
        return

    headers = ["Layer", "Expected Loss Share", "Risk Characterization"]
    rows: list[list[str]] = []
    for layer in layers:
        layer_name = layer.get("layer_type", "Unknown")
        share_pct = layer.get("expected_loss_share_pct", 0)
        description = layer.get("description", "")
        rows.append([
            layer_name.replace("_", " ").title(),
            f"{share_pct:.0f}%",
            description,
        ])

    add_styled_table(doc, headers, rows, ds)

    # Summary note
    if layers:
        primary = layers[0]
        primary_share = primary.get("expected_loss_share_pct", 0)
        if primary_share > 0:
            note_para: Any = doc.add_paragraph(style="DOBody")
            note_run: Any = note_para.add_run(
                f"Primary layer carries {primary_share:.0f}% of expected "
                f"loss exposure."
            )
            note_run.font.size = ds.size_small
            note_run.font.color.rgb = ds.color_text_light
            note_run.italic = True


def render_peril_map(
    doc: Any, context: dict[str, Any], ds: DesignSystem
) -> None:
    """Render complete peril map: heat map, bear cases, settlement, tower.

    Phase 60-02: Receives context dict from build_template_context().
    Uses context["_state"] escape hatch for analysis and executive_summary.
    """
    heading: Any = doc.add_paragraph(style="DOHeading2")
    heading.add_run("Peril Assessment")

    # Deserialize peril map from context
    peril_map = _deserialize_peril_map(context)
    if peril_map is None:
        para: Any = doc.add_paragraph(style="DOBody")
        para.add_run("Peril assessment not available.")
        return

    _render_plaintiff_heat_map(doc, peril_map.assessments, ds)
    _render_bear_cases(doc, peril_map.bear_cases, ds)
    _render_settlement_summary(doc, context, ds)
    _render_tower_characterization(doc, context, ds)

    # Mispricing detection alerts
    # TODO(phase-60): move executive_summary to context_builders
    state = context.get("_state")
    if (
        state is not None
        and state.executive_summary is not None
        and state.executive_summary.deal_context.market_intelligence
        is not None
    ):
        mi = state.executive_summary.deal_context.market_intelligence
        alerts: list[str] = []
        if mi.mispricing_alert is not None:
            alerts.append(mi.mispricing_alert)
        if mi.model_vs_market_alert is not None:
            alerts.append(mi.model_vs_market_alert)
        if alerts:
            alert_heading: Any = doc.add_paragraph(style="DOHeading3")
            alert_heading.add_run("Pricing Divergence Alert")
            for alert_text in alerts:
                p: Any = doc.add_paragraph(style="DOBody")
                run: Any = p.add_run(alert_text)
                run.bold = True
                run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)


def _deserialize_peril_map(context: dict[str, Any]) -> PerilMap | None:
    """Deserialize PerilMap from context's analysis.peril_map dict."""
    # TODO(phase-60): use context["peril_map"] when available
    state = context.get("_state")
    if state is None or state.analysis is None or state.analysis.peril_map is None:
        return None
    try:
        return PerilMap.model_validate(state.analysis.peril_map)
    except Exception:
        return None


__all__ = ["render_peril_map"]
