"""Audit risk, debt analysis, and critical accounting estimates.

Split from sect3_financial.py for 500-line compliance.
Renders detailed audit risk assessment, debt structure with maturity
schedule, covenant status, and critical accounting estimates from
LLM extraction of 10-K Item 9A and MD&A.

Phase 60-01: Migrated from state access to shared context dict.
"""

from __future__ import annotations

from typing import Any, cast

from do_uw.models.financials import AuditProfile, ExtractedFinancials
from do_uw.stages.render.design_system import DesignSystem
from do_uw.stages.render.docx_helpers import (
    add_risk_indicator,
    add_styled_table,
    set_cell_shading,
)
from do_uw.stages.render.context_builders._signal_fallback import (
    safe_get_result,
)
from do_uw.stages.render.formatters import (
    format_currency,
    format_number,
    na_if_none,
    safe_float,
)


def _get_signal_results(context: dict[str, Any]) -> dict[str, Any] | None:
    """Extract signal_results dict from context."""
    state = context.get("_state")
    if state is None or state.analysis is None:
        return None
    return state.analysis.signal_results


def render_audit_risk(
    doc: Any, context: dict[str, Any], ds: DesignSystem
) -> None:
    """Render audit risk assessment, debt analysis, and accounting estimates.

    If audit profile is clean (no material weaknesses, no going concern,
    no restatements), renders a concise one-liner instead of full detail.
    Debt structure and critical estimates always render if data available.

    Layout order:
    1. Audit Risk Assessment (full or concise)
    2. D&O context for audit red flags (only if not clean)
    3. Debt Structure table
    4. Debt Maturity Schedule
    5. Critical Accounting Estimates
    """
    # TODO(phase-60): move to context_builders
    state = context["_state"]
    financials = state.extracted.financials if state.extracted else None
    signal_results = _get_signal_results(context)

    # Issue-driven density: concise audit for clean companies
    if _is_audit_clean(financials):
        _render_audit_concise(doc, financials, ds)
    else:
        _render_audit_assessment(doc, financials, signal_results, ds)

    _render_debt_structure(doc, financials, ds)
    _render_debt_maturity(doc, financials, ds)
    _render_critical_estimates(doc, financials, ds)


def _is_audit_clean(financials: ExtractedFinancials | None) -> bool:
    """Check if audit profile has no red flags."""
    if financials is None:
        return False
    audit = financials.audit
    if audit is None:
        return False
    if audit.material_weaknesses:
        return False
    if audit.going_concern is not None and audit.going_concern.value is True:
        return False
    if audit.restatements:
        return False
    return True


def _render_audit_concise(
    doc: Any,
    financials: ExtractedFinancials | None,
    ds: DesignSystem,
) -> None:
    """Render concise audit summary for clean companies."""
    para: Any = doc.add_paragraph(style="DOHeading2")
    para.add_run("Audit Risk Assessment")

    audit = financials.audit if financials else None
    if audit is None:
        body: Any = doc.add_paragraph(style="DOBody")
        body.add_run("Audit profile data not available.")
        return

    auditor_name = _sv_str(audit.auditor_name)
    if auditor_name == "N/A":
        auditor_name = "Not identified"

    tenure_str = ""
    if audit.tenure_years is not None:
        tenure_str = f", {audit.tenure_years.value} year tenure"

    opinion_str = "unqualified opinion"
    if audit.opinion_type is not None:
        opinion_str = f"{audit.opinion_type.value.lower()} opinion"

    body = doc.add_paragraph(style="DOBody")
    body.add_run(
        f"Audit Risk: No concerns. {auditor_name}{tenure_str}, "
        f"{opinion_str}."
    )


def _render_audit_assessment(
    doc: Any,
    financials: ExtractedFinancials | None,
    signal_results: dict[str, Any] | None,
    ds: DesignSystem,
) -> None:
    """Render audit risk assessment table with D&O context from brain YAML."""
    para: Any = doc.add_paragraph(style="DOHeading2")
    para.add_run("Audit Risk Assessment")

    audit = financials.audit if financials else None
    if audit is None:
        body: Any = doc.add_paragraph(style="DOBody")
        body.add_run("Audit profile data not available.")
        return

    # Compute is_big4 from auditor name if not set by extraction
    auditor_name_str = _sv_str(audit.auditor_name)
    is_big4_display = _sv_bool(audit.is_big4)
    if is_big4_display == "N/A" and auditor_name_str != "N/A":
        is_big4_display = "Yes" if _infer_big4(auditor_name_str) else "No"

    # Compute opinion type fallback: if None, infer from other signals
    opinion_display = _sv_str(audit.opinion_type)
    if opinion_display == "N/A":
        # If no going concern + no material weaknesses = likely unqualified
        gc = audit.going_concern
        has_gc = gc is not None and gc.value is True
        has_mw = len(audit.material_weaknesses) > 0
        if not has_gc and not has_mw:
            opinion_display = "Unqualified (inferred)"
        elif has_gc:
            opinion_display = "Going Concern (inferred)"

    # Auditor display: use helpful fallback when name unavailable
    auditor_display = auditor_name_str
    if auditor_display == "N/A":
        auditor_display = "Not identified (review 10-K Item 9A)"

    # Main audit attributes table
    rows: list[list[str]] = []
    rows.append(["Auditor", auditor_display])
    rows.append(["Big 4 Firm", is_big4_display])
    rows.append(["Tenure (Years)", _sv_int(audit.tenure_years)])
    rows.append(["Opinion Type", opinion_display])
    rows.append(["Going Concern", _sv_bool(audit.going_concern)])
    rows.append(["Material Weaknesses", str(len(audit.material_weaknesses))])
    rows.append(["Significant Deficiencies", str(len(audit.significant_deficiencies))])
    rows.append(["Restatements", str(len(audit.restatements))])
    rows.append(["Critical Audit Matters (CAMs)", str(len(audit.critical_audit_matters))])

    if audit.remediation_status is not None:
        rows.append(["Remediation Status", _sv_str(audit.remediation_status)])

    table: Any = add_styled_table(doc, ["Attribute", "Value"], rows, ds)

    # Color code high-risk rows
    _color_audit_risks(table, audit)

    # D&O context from brain signal do_context (Phase 116-02 migration)
    _render_audit_signal_do_context(doc, audit, signal_results, ds)


def _color_audit_risks(table: Any, audit: AuditProfile) -> None:
    """Apply red shading to high-risk audit attributes."""
    # Going concern row (index 4 in rows, table row 5)
    gc = audit.going_concern
    if gc is not None and gc.value is True:
        set_cell_shading(table.rows[5].cells[1], "FCE8E6")

    # Material weaknesses row (index 5, table row 6)
    if audit.material_weaknesses:
        set_cell_shading(table.rows[6].cells[1], "FCE8E6")

    # Restatements row (index 7, table row 8)
    if audit.restatements:
        set_cell_shading(table.rows[8].cells[1], "FCE8E6")


def _render_audit_signal_do_context(
    doc: Any,
    audit: AuditProfile,
    signal_results: dict[str, Any] | None,
    ds: DesignSystem,
) -> None:
    """Render D&O context for audit red flags from brain signal do_context.

    Replaces the deleted _add_audit_do_context() function. Each audit finding
    now pulls its D&O commentary from the corresponding brain signal's
    do_context field instead of hardcoded Python strings.
    """
    if audit.material_weaknesses:
        sig = safe_get_result(signal_results, "FIN.ACCT.material_weakness")
        do_text = sig.do_context if sig and sig.do_context else ""
        if do_text:
            fp: Any = doc.add_paragraph(style="DOBody")
            fp.add_run(f"Material Weaknesses ({len(audit.material_weaknesses)}): {do_text}")
            add_risk_indicator(fp, "HIGH", ds)
        for mw in audit.material_weaknesses[:5]:
            detail: Any = doc.add_paragraph(style="DOBody")
            detail.add_run(f"  - {mw.value}").font.size = _pt(9)

    if audit.restatements:
        sig = safe_get_result(signal_results, "FIN.ACCT.restatement")
        do_text = sig.do_context if sig and sig.do_context else ""
        if do_text:
            fp: Any = doc.add_paragraph(style="DOBody")
            fp.add_run(f"Restatements ({len(audit.restatements)}): {do_text}")
            add_risk_indicator(fp, "CRITICAL", ds)
        for rst in audit.restatements[:5]:
            detail: Any = doc.add_paragraph(style="DOBody")
            rst_data: dict[str, str] = rst.value
            detail.add_run(f"  - {rst_data.get('date', 'N/A')}").font.size = _pt(9)

    gc = audit.going_concern
    if gc is not None and gc.value is True:
        sig = safe_get_result(signal_results, "FIN.ACCT.quality_indicators")
        do_text = sig.do_context if sig and sig.do_context else ""
        if do_text:
            fp = doc.add_paragraph(style="DOBody")
            fp.add_run(f"Going Concern: {do_text}")
            add_risk_indicator(fp, "CRITICAL", ds)

    if audit.critical_audit_matters:
        sig = safe_get_result(signal_results, "FIN.ACCT.internal_controls")
        do_text = sig.do_context if sig and sig.do_context else ""
        if do_text:
            fp = doc.add_paragraph(style="DOBody")
            fp.add_run(
                f"Critical Audit Matters ({len(audit.critical_audit_matters)}): {do_text}"
            )
            add_risk_indicator(fp, "ELEVATED", ds)
        for cam in audit.critical_audit_matters[:5]:
            detail = doc.add_paragraph(style="DOBody")
            detail.add_run(f"  - {cam.value}").font.size = _pt(9)

    if audit.tenure_years is not None:
        tenure: int = audit.tenure_years.value
        if tenure > 20:
            sig = safe_get_result(signal_results, "FIN.ACCT.auditor")
            do_text = sig.do_context if sig and sig.do_context else ""
            if do_text:
                fp = doc.add_paragraph(style="DOBody")
                fp.add_run(f"Extended Auditor Tenure ({tenure} years): {do_text}")
                add_risk_indicator(fp, "ELEVATED", ds)


def _render_debt_structure(
    doc: Any,
    financials: ExtractedFinancials | None,
    ds: DesignSystem,
) -> None:
    """Render debt structure table from LLM extraction."""
    para: Any = doc.add_paragraph(style="DOHeading2")
    para.add_run("Debt Analysis")

    if financials is None:
        body: Any = doc.add_paragraph(style="DOBody")
        body.add_run("Debt analysis data not available.")
        return

    rows: list[list[str]] = []

    # Leverage ratios
    if financials.leverage:
        lev = financials.leverage.value
        _add_ratio_row(rows, "Debt-to-Equity", lev.get("debt_to_equity"))
        _add_ratio_row(rows, "Debt-to-EBITDA", lev.get("debt_to_ebitda"))
        _add_ratio_row(rows, "Interest Coverage", lev.get("interest_coverage"))

    # Debt structure
    if financials.debt_structure:
        ds_val: dict[str, Any] = financials.debt_structure.value
        total = ds_val.get("total_debt")
        if total is not None:
            rows.append([
                "Total Debt",
                format_currency(safe_float(total), compact=True),
            ])

        # Instrument details from LLM extraction
        instruments: Any = ds_val.get("instruments")
        if isinstance(instruments, list):
            for raw_inst in cast(list[Any], instruments)[:10]:
                inst: dict[str, Any] = cast(dict[str, Any], raw_inst)
                name = str(inst.get("name", "Unknown"))
                itype = str(inst.get("type", ""))
                amount: Any = inst.get("outstanding")
                rate = str(inst.get("rate", ""))
                maturity = str(inst.get("maturity", ""))
                amt_str = (
                    format_currency(safe_float(amount), compact=True)
                    if amount is not None
                    else "N/A"
                )
                rows.append([
                    f"  {name} ({itype})" if itype else f"  {name}",
                    f"{amt_str}, {rate}, due {maturity}"
                    if maturity
                    else amt_str,
                ])

        # Covenants
        covenants: Any = ds_val.get("covenants")
        if isinstance(covenants, dict):
            cov_dict: dict[str, Any] = cast(dict[str, Any], covenants)
            status = str(cov_dict.get("status", "Unknown"))
            rows.append(["Covenant Status", status])

    # Refinancing risk
    if financials.refinancing_risk:
        rr: dict[str, Any] = financials.refinancing_risk.value
        risk_level = str(rr.get("risk_level", "N/A"))
        rows.append(["Refinancing Risk", risk_level])

    if rows:
        table: Any = add_styled_table(doc, ["Metric", "Value"], rows, ds)

        # Color code high leverage
        _color_debt_risks(table, financials)
    else:
        body = doc.add_paragraph(style="DOBody")
        body.add_run("No debt analysis data available.")


def _add_ratio_row(
    rows: list[list[str]], label: str, val: float | None,
) -> None:
    """Add a ratio row with formatting."""
    if val is None:
        rows.append([label, "N/A"])
    else:
        rows.append([label, f"{val:.2f}"])


def _color_debt_risks(
    table: Any, financials: ExtractedFinancials,
) -> None:
    """Apply conditional coloring to debt risk indicators."""
    if financials.leverage is None:
        return
    lev = financials.leverage.value
    dte = lev.get("debt_to_equity")
    if dte is not None and dte > 2.0:
        # D/E row is first data row (index 1)
        set_cell_shading(table.rows[1].cells[1], "FCE8E6")

    ic = lev.get("interest_coverage")
    if ic is not None and ic < 3.0:
        # Interest coverage is row 3
        if len(table.rows) > 3:
            set_cell_shading(table.rows[3].cells[1], "FFF3CD")


def _render_debt_maturity(
    doc: Any,
    financials: ExtractedFinancials | None,
    ds: DesignSystem,
) -> None:
    """Render debt maturity schedule from LLM extraction."""
    if financials is None or financials.debt_structure is None:
        return

    ds_val: dict[str, Any] = financials.debt_structure.value
    maturities: Any = ds_val.get("maturity_schedule")
    if not isinstance(maturities, list) or not maturities:
        return

    para: Any = doc.add_paragraph(style="DOHeading3")
    para.add_run("Debt Maturity Schedule")

    rows: list[list[str]] = []
    for raw_entry in cast(list[Any], maturities)[:10]:
        entry: dict[str, Any] = cast(dict[str, Any], raw_entry)
        year = str(entry.get("year", "N/A"))
        amount: Any = entry.get("amount")
        amt_str = (
            format_currency(safe_float(amount), compact=True)
            if amount is not None
            else "N/A"
        )
        rows.append([year, amt_str])

    if rows:
        add_styled_table(doc, ["Year", "Maturing Amount"], rows, ds)

        # Near-term maturity highlight
        near_term: Any = ds_val.get("near_term_maturities")
        if near_term is not None and safe_float(near_term) > 0:
            note: Any = doc.add_paragraph(style="DOBody")
            run: Any = note.add_run(
                f"Near-term maturities: {format_currency(safe_float(near_term), compact=True)}. "
                "Concentration of near-term maturities increases refinancing "
                "risk and may limit financial flexibility for litigation defense."
            )
def _render_critical_estimates(
    doc: Any,
    financials: ExtractedFinancials | None,
    ds: DesignSystem,
) -> None:
    """Render critical accounting estimates from LLM MD&A extraction."""
    if financials is None or financials.tax_indicators is None:
        return

    tax: dict[str, Any] = financials.tax_indicators.value
    etr: Any = tax.get("effective_tax_rate")
    deferred: Any = tax.get("deferred_tax_asset")
    utp: Any = tax.get("uncertain_tax_positions")

    if etr is None and deferred is None and utp is None:
        return

    para: Any = doc.add_paragraph(style="DOHeading3")
    para.add_run("Critical Accounting Estimates")

    rows: list[list[str]] = []

    if etr is not None:
        etr_f = safe_float(etr)
        etr_pct = etr_f * 100 if etr_f <= 1.0 else etr_f
        assessment = ""
        if etr_pct < 15:
            assessment = "Below statutory rate -- review required"
        elif etr_pct > 30:
            assessment = "Above statutory rate"
        else:
            assessment = "Normal range"
        rows.append(["Effective Tax Rate", f"{etr_pct:.1f}%", assessment])

    if deferred is not None:
        rows.append([
            "Deferred Tax Assets",
            format_currency(safe_float(deferred), compact=True),
            "",
        ])

    if utp is not None:
        rows.append([
            "Uncertain Tax Positions",
            format_currency(safe_float(utp), compact=True),
            "Potential SEC/IRS exposure",
        ])

    if rows:
        add_styled_table(
            doc, ["Estimate", "Value", "Assessment"], rows, ds,
        )


_BIG4_NAMES: set[str] = {
    "pricewaterhousecoopers", "pwc",
    "deloitte", "deloitte & touche",
    "ernst & young", "ernst&young", "ey",
    "kpmg",
}


def _infer_big4(auditor_name: str) -> bool:
    """Infer Big 4 status from auditor name string."""
    name_lower = auditor_name.lower()
    return any(b4 in name_lower for b4 in _BIG4_NAMES)


def _sv_str(sv: Any) -> str:
    """Extract string from SourcedValue or return N/A."""
    if sv is None:
        return "N/A"
    return na_if_none(sv.value)


def _sv_bool(sv: Any) -> str:
    """Extract bool from SourcedValue as Yes/No or N/A."""
    if sv is None:
        return "N/A"
    return "Yes" if sv.value else "No"


def _sv_int(sv: Any) -> str:
    """Extract int from SourcedValue or return N/A."""
    if sv is None:
        return "N/A"
    return format_number(sv.value)


def _pt(size: int) -> Any:
    """Create a Pt size value."""
    from docx.shared import Pt  # type: ignore[import-untyped]

    return Pt(size)


__all__ = ["render_audit_risk"]
