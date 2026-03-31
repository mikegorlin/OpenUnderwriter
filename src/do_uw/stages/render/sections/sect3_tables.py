"""Balance sheet, cash flow, and key ratio tables with conditional formatting.

Renders balance sheet and cash flow tables with multi-period data,
YoY change column, and conditional cell shading. NO GREEN ANYWHERE.
Also provides key ratios comparison (company vs peer group average).

Phase 60-01: Migrated from state access to shared context dict.
"""

from __future__ import annotations

from typing import Any

from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

from do_uw.models.financials import (
    FinancialLineItem,
    FinancialStatement,
)
from do_uw.stages.render.design_system import DesignSystem
from do_uw.stages.render.docx_helpers import set_cell_shading
from do_uw.stages.render.formatters import (
    format_compact,
    format_currency,
    format_percentage,
)

# Metric direction: True = higher is better, False = higher is worse
_METRIC_DIRECTION: dict[str, bool] = {
    # Income statement
    "Total Revenue": True,
    "Revenue": True,
    "Net Revenue": True,
    "Cost of Revenue": False,
    "Cost of Goods Sold": False,
    "Gross Profit": True,
    "Operating Income": True,
    "Operating Expenses": False,
    "EBITDA": True,
    "Net Income": True,
    "Earnings Per Share": True,
    "EPS (Basic)": True,
    "EPS (Diluted)": True,
    "R&D Expenses": False,
    "SG&A Expenses": False,
    "Interest Expense": False,
    "Tax Provision": False,
    "Gross Margin": True,
    "Operating Margin": True,
    "Net Margin": True,
    # Balance sheet
    "Total Assets": True,
    "Total Liabilities": False,
    "Total Debt": False,
    "Long-Term Debt": False,
    "Short-Term Debt": False,
    "Current Liabilities": False,
    "Total Equity": True,
    "Stockholders' Equity": True,
    "Cash and Equivalents": True,
    "Cash and Cash Equivalents": True,
    "Current Assets": True,
    "Goodwill": True,
    "Intangible Assets": True,
    "Accounts Receivable": True,
    "Inventory": False,
    "Accounts Payable": False,
    "Working Capital": True,
    # Cash flow
    "Operating Cash Flow": True,
    "Capital Expenditures": False,
    "Free Cash Flow": True,
    "Investing Cash Flow": True,
    "Financing Cash Flow": True,
    "Depreciation & Amortization": False,
    "Stock-Based Compensation": False,
    "Dividends Paid": False,
    "Share Repurchases": False,
}


def render_financial_tables(
    doc: Any, context: dict[str, Any], ds: DesignSystem
) -> None:
    """Render all financial statement tables with conditional formatting.

    Handles income statement, balance sheet, cash flow, and key ratios.
    Each table has multi-period columns and YoY change with conditional
    cell shading.
    """
    # TODO(phase-60): move to context_builders
    state = context["_state"]
    financials = (
        state.extracted.financials if state.extracted else None
    )
    stmts = financials.statements if financials else None
    if stmts is None:
        body: Any = doc.add_paragraph(style="DOBody")
        body.add_run("Financial statement data not available.")
        return

    if stmts.income_statement:
        _render_income_statement(doc, stmts.income_statement, ds)
    if stmts.balance_sheet:
        _render_balance_sheet(doc, stmts.balance_sheet, ds)
    if stmts.cash_flow:
        _render_cash_flow(doc, stmts.cash_flow, ds)

    # Key ratios comparison (company vs peers)
    _render_key_ratios(doc, context, ds)


def _render_income_statement(
    doc: Any,
    stmt: FinancialStatement,
    ds: DesignSystem,
) -> None:
    """Render income statement with margins and conditional formatting."""
    heading: Any = doc.add_paragraph(style="DOHeading2")
    heading.add_run("Income Statement")

    if not stmt.line_items:
        body: Any = doc.add_paragraph(style="DOBody")
        body.add_run("No income statement data available.")
        return

    _render_statement_table(doc, stmt, ds)


def _render_balance_sheet(
    doc: Any,
    stmt: FinancialStatement,
    ds: DesignSystem,
) -> None:
    """Render balance sheet with key ratios computed inline."""
    heading: Any = doc.add_paragraph(style="DOHeading2")
    heading.add_run("Balance Sheet")

    if not stmt.line_items:
        body: Any = doc.add_paragraph(style="DOBody")
        body.add_run("No balance sheet data available.")
        return

    _render_statement_table(doc, stmt, ds)

    # Compute and display inline ratios
    ratios = _compute_bs_ratios(stmt)
    if ratios:
        ratio_heading: Any = doc.add_paragraph(style="DOHeading3")
        ratio_heading.add_run("Key Balance Sheet Ratios")
        from do_uw.stages.render.docx_helpers import add_styled_table

        add_styled_table(
            doc,
            ["Ratio", "Value", "Assessment"],
            ratios,
            ds,
        )


def _compute_bs_ratios(stmt: FinancialStatement) -> list[list[str]]:
    """Compute key balance sheet ratios from line items."""
    rows: list[list[str]] = []

    ca = _latest_value(stmt, "Current Assets")
    cl = _latest_value(stmt, "Current Liabilities")
    ta = _latest_value(stmt, "Total Assets")
    tl = _latest_value(stmt, "Total Liabilities")
    td = _latest_value(stmt, "Total Debt")
    te = _latest_value(stmt, "Total Equity", "Stockholders' Equity")
    cash = _latest_value(stmt, "Cash and Cash Equivalents", "Cash and Equivalents")

    if ca is not None and cl is not None and cl != 0:
        ratio = ca / cl
        assessment = "Healthy" if ratio > 1.5 else ("Adequate" if ratio > 1.0 else "Concern")
        rows.append(["Current Ratio", f"{ratio:.2f}", assessment])

    if td is not None and te is not None and te != 0:
        ratio = td / te
        assessment = "Conservative" if ratio < 0.5 else ("Moderate" if ratio < 1.5 else "Elevated")
        rows.append(["Debt/Equity", f"{ratio:.2f}", assessment])

    if tl is not None and ta is not None and ta != 0:
        ratio = tl / ta
        assessment = "Low" if ratio < 0.4 else ("Moderate" if ratio < 0.6 else "High")
        rows.append(["Debt/Assets", f"{ratio:.2f}", assessment])

    if cash is not None and ta is not None and ta != 0:
        ratio = (cash / ta) * 100
        rows.append(["Cash/Assets", f"{ratio:.1f}%", ""])

    return rows


def _latest_value(
    stmt: FinancialStatement, *labels: str,
) -> float | None:
    """Get the latest period value for a line item by label."""
    for item in stmt.line_items:
        for label in labels:
            if label.lower() in item.label.lower():
                vals = list(item.values.values())
                if vals and vals[0] is not None:
                    return float(vals[0].value)
    return None


def _render_cash_flow(
    doc: Any,
    stmt: FinancialStatement,
    ds: DesignSystem,
) -> None:
    """Render cash flow statement with free cash flow emphasis."""
    heading: Any = doc.add_paragraph(style="DOHeading2")
    heading.add_run("Cash Flow Statement")

    if not stmt.line_items:
        body: Any = doc.add_paragraph(style="DOBody")
        body.add_run("No cash flow data available.")
        return

    _render_statement_table(doc, stmt, ds)

    # Free cash flow callout
    ocf = _latest_value(stmt, "Operating Cash Flow")
    capex = _latest_value(stmt, "Capital Expenditures")
    if ocf is not None and capex is not None:
        fcf = ocf - abs(capex)
        fcf_str = format_currency(fcf, compact=True)
        assessment = "positive" if fcf > 0 else "negative"
        note: Any = doc.add_paragraph(style="DOBody")
        run: Any = note.add_run(
            f"Free Cash Flow: {fcf_str} ({assessment}). "
        )
        if fcf < 0:
            run2: Any = note.add_run(
                "Negative FCF may constrain litigation defense capacity "
                "and signals potential funding risk."
            )
            _ = run2
        _ = run


def _render_key_ratios(
    doc: Any, context: dict[str, Any], ds: DesignSystem,
) -> None:
    """Render company key financial ratios.

    Shows peer comparison column only when peer data is available.
    """
    # TODO(phase-60): move to context_builders
    state = context["_state"]
    financials = state.extracted.financials if state.extracted else None
    if financials is None:
        return

    # Only render if we have company liquidity/leverage data
    if financials.liquidity is None and financials.leverage is None:
        return

    heading: Any = doc.add_paragraph(style="DOHeading3")
    heading.add_run("Key Ratios: Company vs. Peers")

    rows: list[list[str]] = []

    if financials.liquidity is not None:
        liq = financials.liquidity.value
        cr = liq.get("current_ratio")
        if cr is not None:
            rows.append(["Current Ratio", f"{cr:.2f}"])
        qr = liq.get("quick_ratio")
        if qr is not None:
            rows.append(["Quick Ratio", f"{qr:.2f}"])

    if financials.leverage is not None:
        lev = financials.leverage.value
        dte = lev.get("debt_to_equity")
        if dte is not None:
            rows.append(["Debt/Equity", f"{dte:.2f}"])
        ic = lev.get("interest_coverage")
        if ic is not None:
            rows.append(["Interest Coverage", f"{ic:.1f}x"])

    if rows:
        from do_uw.stages.render.docx_helpers import add_styled_table

        add_styled_table(
            doc,
            ["Ratio", "Value"],
            rows,
            ds,
        )


# ---------------------------------------------------------------------------
# Shared statement table renderer
# ---------------------------------------------------------------------------


def _render_statement_table(
    doc: Any,
    stmt: FinancialStatement,
    ds: DesignSystem,
) -> None:
    """Render a single financial statement as a table with conditional formatting.

    Suppresses:
    - Periods (columns) where ALL line items are N/A
    - Line items (rows) where ALL period values are N/A
    """
    raw_periods = stmt.periods if stmt.periods else []

    # First pass: count how many items have data for each period
    period_data_count: dict[str, int] = dict.fromkeys(raw_periods, 0)
    total_items = len(stmt.line_items)
    for item in stmt.line_items:
        for period in raw_periods:
            sv = item.values.get(period)
            if sv is not None:
                period_data_count[period] += 1

    # Filter: keep periods with >20% data coverage (avoids single-item artifacts)
    min_coverage = max(1, int(total_items * 0.2))
    periods = [
        p for p in raw_periods
        if period_data_count[p] >= min_coverage
    ]
    if not periods:
        body: Any = doc.add_paragraph(style="DOBody")
        body.add_run("No financial data available for this statement.")
        return

    headers = ["Metric", *periods]
    if len(periods) >= 2:
        headers.append("YoY Change")

    # Build rows, skipping items where ALL period values are N/A
    rows_data: list[tuple[FinancialLineItem, list[str]]] = []
    for item in stmt.line_items:
        row_cells: list[str] = [item.label]
        has_any_value = False
        for period in periods:
            sv = item.values.get(period)
            if sv is not None:
                row_cells.append(_format_value(sv.value, item.label))
                has_any_value = True
            else:
                row_cells.append("N/A")
        if not has_any_value:
            continue  # Skip entirely N/A rows
        if len(periods) >= 2:
            yoy_str = (
                format_percentage(item.yoy_change)
                if item.yoy_change is not None
                else "N/A"
            )
            row_cells.append(yoy_str)
        rows_data.append((item, row_cells))

    # Create table
    n_rows = len(rows_data) + 1
    n_cols = len(headers)
    table: Any = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    # Header row
    _style_header_row(table.rows[0], headers, ds)

    # Data rows with conditional formatting
    yoy_col_idx = len(headers) - 1 if len(periods) >= 2 else -1
    for row_idx, (item, cells) in enumerate(rows_data):
        row: Any = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(cells):
            cell: Any = row.cells[col_idx]
            if row_idx % 2 == 1:
                set_cell_shading(cell, ds.row_alt)

            para: Any = cell.paragraphs[0]
            para.clear()
            run: Any = para.add_run(cell_text)
            run.font.name = ds.font_body
            run.font.size = Pt(9)
            run.font.color.rgb = ds.color_text

            # Apply conditional formatting on YoY Change column
            if col_idx == yoy_col_idx and item.yoy_change is not None:
                shading = _get_conditional_shading(
                    item.yoy_change, item.label
                )
                if shading:
                    set_cell_shading(cell, shading)


def _style_header_row(
    header_row: Any, headers: list[str], ds: DesignSystem
) -> None:
    """Apply navy header styling to the first row."""
    for idx, header_text in enumerate(headers):
        cell: Any = header_row.cells[idx]
        set_cell_shading(cell, ds.header_bg)
        para: Any = cell.paragraphs[0]
        para.clear()
        run: Any = para.add_run(header_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = ds.font_body
        run.font.size = Pt(9)


def _get_conditional_shading(yoy_change: float, label: str) -> str | None:
    """Determine conditional shading color for a YoY change value.

    Returns hex color (no '#' prefix) for cell shading, or None.
    RED (#FCE8E6) for deteriorating (>10% = red, >5% = amber).
    BLUE (#DCEEF8) for improving. NOT green.
    AMBER (#FFF3CD) for slight/mixed changes.
    """
    higher_is_better = _METRIC_DIRECTION.get(label, True)
    abs_change = abs(yoy_change)

    # Trivial changes get no special formatting
    if abs_change < 1.0:
        return None

    # Determine if this change is positive or negative for the company
    is_improving = (
        (yoy_change > 0 and higher_is_better)
        or (yoy_change < 0 and not higher_is_better)
    )
    is_deteriorating = (
        (yoy_change < 0 and higher_is_better)
        or (yoy_change > 0 and not higher_is_better)
    )

    if is_deteriorating:
        if abs_change >= 10.0:
            return "FCE8E6"  # Red for significant deterioration
        return "FFF3CD"  # Amber for moderate deterioration
    if is_improving:
        if abs_change >= 10.0:
            return "DCEEF8"  # Blue for significant improvement (NOT green)
        return "FFF3CD"  # Amber for moderate improvement
    return None


def _format_value(value: float, label: str) -> str:
    """Format a financial value based on the metric label.

    Detects non-currency line items (shares, EPS, margins) and applies
    the appropriate formatter instead of always using format_currency.
    """
    label_lower = label.lower()
    if "margin" in label_lower or "%" in label_lower:
        return format_percentage(value)
    if "eps" in label_lower or "per share" in label_lower:
        return f"${value:.2f}"
    # Share counts and non-monetary counts -- no $ prefix
    if any(kw in label_lower for kw in (
        "shares", "share count", "weighted average",
        "diluted shares", "basic shares",
    )):
        return format_compact(value)
    return format_currency(value, compact=True)


__all__ = ["render_financial_tables"]
