"""Section 1: Executive Summary renderer (v2).

Narrative-first: thesis, sourced snapshot, tier, inherent risk with
decile and spectrum, key findings as narrative paragraphs, claim
probability with context, tower recommendation.

Delegates table/visual rendering to sect1_executive_tables.py.

Phase 60-01: Migrated from state access to shared context dict.
"""

from __future__ import annotations

from typing import Any

from do_uw.stages.render.design_system import DesignSystem
from do_uw.stages.render.docx_helpers import (
    add_risk_indicator,
    add_styled_table,
)
from do_uw.stages.render.sections.sect1_executive_tables import (
    add_tier_indicator,
    build_ceiling_line,
    build_factor_breakdown,
    render_claim_probability,
    render_inherent_risk,
    render_snapshot,
    render_tower_recommendation,
)
from do_uw.stages.render.sections.sect1_findings import (
    build_negative_narrative,
    build_positive_narrative,
)
from do_uw.stages.render.sections.sect1_helpers import (
    build_thesis_narrative,
)
from do_uw.stages.render.sections.sect1_market_context import (
    render_market_context,
)


def render_section_1(
    doc: Any, context: dict[str, Any], ds: DesignSystem
) -> None:
    """Render Executive Summary with rich narratives and visual indicators."""
    _render_heading(doc, ds)
    _render_data_quality_notice(doc, context, ds)
    _render_thesis(doc, context, ds)
    render_snapshot(doc, context, ds)
    _render_tier(doc, context, ds)
    render_inherent_risk(doc, context, ds)
    _render_key_negatives(doc, context, ds)
    _render_key_positives(doc, context, ds)
    render_claim_probability(doc, context, ds)
    render_tower_recommendation(doc, context, ds)
    render_market_context(doc, context, ds)


def _render_heading(doc: Any, ds: DesignSystem) -> None:
    """Add section heading."""
    para: Any = doc.add_paragraph(style="DOHeading1")
    run: Any = para.add_run("Executive Summary")
    _ = (run, ds)


def _render_data_quality_notice(
    doc: Any, context: dict[str, Any], ds: DesignSystem
) -> None:
    """Render data quality notice when blind spot detection was skipped.

    When no search API is configured, the worksheet must clearly warn
    the reader that web-based blind spot detection was NOT performed.
    This prevents silent false claims of "no issues found" when search
    was never executed.
    """
    _ = ds
    # TODO(phase-60): move acquired_data to context_builders
    state = context["_state"]
    if state.acquired_data is None:
        # No acquired data at all -- notice should render
        search_configured = False
    else:
        blind_spots = state.acquired_data.blind_spot_results
        search_configured = bool(
            blind_spots.get("search_configured", False)
        )

    if search_configured:
        return

    from docx.shared import RGBColor  # type: ignore[import-untyped]

    heading: Any = doc.add_paragraph(style="DOHeading2")
    heading.add_run("Data Quality Notice")

    para: Any = doc.add_paragraph(style="DOBody")
    run: Any = para.add_run(
        "IMPORTANT: Web-based blind spot detection was not performed "
        "for this analysis (no search API configured). This worksheet "
        "relies solely on SEC filing data and structured APIs. "
        "Publicly known events such as short seller reports, state AG "
        "actions, active securities class actions not disclosed in "
        "filings, news coverage, and social media sentiment may not "
        "be reflected. Manual verification of the litigation and "
        "regulatory landscape is strongly recommended."
    )
    run.bold = True
    run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)  # Dark red for warning


def _render_thesis(doc: Any, context: dict[str, Any], ds: DesignSystem) -> None:
    """Render the underwriting thesis as a rich narrative paragraph."""
    para: Any = doc.add_paragraph(style="DOHeading2")
    para.add_run("Underwriting Thesis")

    # TODO(phase-60): move benchmark to context_builders
    state = context["_state"]

    # Read pre-computed narrative, fallback to local computation
    narrative = (
        state.benchmark.thesis_narrative
        if state.benchmark and state.benchmark.thesis_narrative
        else build_thesis_narrative(state)
    )

    body: Any = doc.add_paragraph(style="DOBody")
    run: Any = body.add_run(narrative)
    run.font.color.rgb = ds.color_primary


def _render_tier(doc: Any, context: dict[str, Any], ds: DesignSystem) -> None:
    """Render tier classification with score breakdown."""
    para: Any = doc.add_paragraph(style="DOHeading3")
    para.add_run("Tier Classification")

    scoring = context.get("scoring")
    # TODO(phase-60): move scoring model access to context_builders
    state = context["_state"]
    tier_cls = state.scoring.tier if state.scoring else None
    if tier_cls is None:
        body: Any = doc.add_paragraph(style="DOBody")
        body.add_run("Tier classification not available.")
        return

    tier_name = tier_cls.tier.value
    composite = (
        f"{state.scoring.composite_score:.1f}/100" if state.scoring else "N/A"
    )
    quality = (
        f"{state.scoring.quality_score:.1f}/100" if state.scoring else "N/A"
    )

    factor_line = build_factor_breakdown(context)
    rows: list[list[str]] = [
        ["Tier", tier_name],
        ["Composite Score", composite],
        ["Quality Score", quality],
    ]
    if factor_line:
        rows.append(["Factor Breakdown", factor_line])

    ceiling_line = build_ceiling_line(context)
    if ceiling_line:
        rows.append(["Score Ceiling", ceiling_line])

    range_line = (
        f"{state.scoring.composite_score:.1f} falls in "
        f"{tier_name} range ({tier_cls.score_range_low}-{tier_cls.score_range_high})"
        if state.scoring
        else ""
    )
    if range_line:
        rows.append(["Position", range_line])

    rows.append(["Appetite", tier_cls.action or "N/A"])
    add_styled_table(doc, ["Attribute", "Value"], rows, ds)

    tier_para: Any = doc.add_paragraph(style="DOBody")
    tier_para.add_run(f"Classification: {tier_name}")
    add_tier_indicator(tier_para, tier_name, ds)


def _render_key_negatives(
    doc: Any, context: dict[str, Any], ds: DesignSystem
) -> None:
    """Render top negative findings as narrative paragraphs."""
    para: Any = doc.add_paragraph(style="DOHeading3")
    para.add_run("Underwriting Headwinds")

    state = context["_state"]
    findings = (
        state.executive_summary.key_findings
        if state.executive_summary
        else None
    )
    negatives = findings.negatives[:5] if findings else []

    if not negatives:
        body: Any = doc.add_paragraph(style="DOBody")
        body.add_run("No negative findings identified.")
        return

    for idx, finding in enumerate(negatives, 1):
        title, narrative = build_negative_narrative(finding, idx, context)

        # Title line with risk indicator
        tp: Any = doc.add_paragraph(style="DOBody")
        run: Any = tp.add_run(f"{idx}. {title}")
        run.bold = True
        risk_level = "CRITICAL" if finding.ranking_score >= 0.75 else (
            "HIGH" if finding.ranking_score >= 0.5 else "ELEVATED"
        )
        add_risk_indicator(tp, risk_level, ds)

        # Narrative paragraph
        np: Any = doc.add_paragraph(style="DOBody")
        np.add_run(narrative)


def _render_key_positives(
    doc: Any, context: dict[str, Any], ds: DesignSystem
) -> None:
    """Render top positive findings as narrative paragraphs."""
    para: Any = doc.add_paragraph(style="DOHeading3")
    para.add_run("Risk Mitigants")

    state = context["_state"]
    findings = (
        state.executive_summary.key_findings
        if state.executive_summary
        else None
    )
    positives = findings.positives[:5] if findings else []

    if not positives:
        body: Any = doc.add_paragraph(style="DOBody")
        body.add_run("No positive findings identified.")
        return

    for idx, finding in enumerate(positives, 1):
        title, narrative = build_positive_narrative(finding, idx, context)

        tp: Any = doc.add_paragraph(style="DOBody")
        run: Any = tp.add_run(f"{idx}. {title}")
        run.bold = True

        np: Any = doc.add_paragraph(style="DOBody")
        np.add_run(narrative)


__all__ = ["render_section_1"]
