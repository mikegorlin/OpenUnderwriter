"""Four-tier rendering helpers for visual distinction within sections.

Implements SC3 (four-tier display) and SC4 (underwriter education) with:
- Tier 1: Customary -- standard body text (render_customary_block)
- Tier 2: Objective -- highlighted callout for risk findings (render_objective_signal)
- Tier 3: Relative -- peer context with scenario education (render_scenario_context)
- Tier 4: Subjective -- question prompt cross-ref (add_meeting_prep_ref)

Each helper uses the DesignSystem for consistent styling. Signal colors
are risk-level-dependent. Scenario context uses caption style for
visual distinction from facts.
"""

from __future__ import annotations

from typing import Any

from do_uw.stages.render.design_system import DesignSystem, get_risk_color

# ---------------------------------------------------------------------------
# Risk level to shading color mapping (hex without '#' for XML)
# ---------------------------------------------------------------------------

_SIGNAL_SHADING: dict[str, str] = {
    "CRITICAL": "FCE8E6",   # Light red
    "HIGH": "FCE8E6",       # Light red
    "ELEVATED": "FFF3CD",   # Light amber
    "MODERATE": "DCEEF8",   # Light blue
    "LOW": "F2F4F8",        # Light gray
}


def render_objective_signal(
    doc: Any,
    ds: DesignSystem,
    signal_text: str,
    risk_level: str,
    evidence: str | None = None,
) -> None:
    """Render a visually highlighted callout for an objective risk finding.

    Uses risk-level-dependent background shading with bold signal text.
    Example: "[ELEVATED] Beneish M-Score = -1.42 (LIKELY MANIPULATOR)"

    Args:
        doc: The python-docx Document.
        ds: Design system with visual constants.
        signal_text: Primary signal description.
        risk_level: Risk level (CRITICAL, HIGH, ELEVATED, MODERATE, LOW).
        evidence: Optional supporting evidence text.
    """
    from docx.oxml.ns import qn  # type: ignore[import-untyped]
    from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

    para: Any = doc.add_paragraph(style="DOBody")

    # Add left border shading via XML for visual distinction
    shading_hex = _SIGNAL_SHADING.get(risk_level.upper(), "F2F4F8")
    pPr = para._element.get_or_add_pPr()
    shading = pPr.makeelement(qn("w:shd"), {
        qn("w:fill"): shading_hex,
        qn("w:val"): "clear",
    })
    pPr.append(shading)

    # Risk level tag
    tag_run: Any = para.add_run(f"[{risk_level.upper()}] ")
    tag_run.bold = True
    tag_run.font.size = Pt(9)
    color_hex = get_risk_color(risk_level)
    tag_run.font.color.rgb = RGBColor(
        int(color_hex[1:3], 16),
        int(color_hex[3:5], 16),
        int(color_hex[5:7], 16),
    )

    # Signal text
    sig_run: Any = para.add_run(signal_text)
    sig_run.bold = True
    sig_run.font.size = ds.size_body

    # Evidence (lighter weight)
    if evidence:
        ev_run: Any = para.add_run(f" -- {evidence}")
        ev_run.font.size = ds.size_small
        ev_run.font.color.rgb = ds.color_text_light


def render_scenario_context(
    doc: Any,
    ds: DesignSystem,
    scenario_text: str,
) -> None:
    """Render Level 2 'What COULD BE' educational context.

    Italic, indented caption-style text that provides industry claim
    rates, peer examples, or actuarial context. Visually distinct from
    factual data (uses DOCaption style).

    Example: "Companies with M-Score below -1.78 experience SCA filings
    at 2.3x the base rate."

    Args:
        doc: The python-docx Document.
        ds: Design system with visual constants.
        scenario_text: Educational scenario description.
    """
    from docx.shared import Pt  # type: ignore[import-untyped]

    para: Any = doc.add_paragraph(style="DOCaption")
    para.paragraph_format.left_indent = Pt(18)
    run: Any = para.add_run(f"Context: {scenario_text}")
    run.italic = True
    run.font.size = ds.size_caption
    run.font.color.rgb = ds.color_text_light


def add_meeting_prep_ref(
    doc: Any,
    ds: DesignSystem,
    question_ref: str,
) -> None:
    """Render Level 3 'What to ASK' cross-reference to meeting prep.

    Accent-colored italic text pointing the underwriter to the relevant
    meeting prep question. Uses descriptive text (not Q numbers) to
    avoid dependency on question numbering order.

    Example: "See Meeting Prep: Earnings Quality" in accent color.

    Args:
        doc: The python-docx Document.
        ds: Design system with visual constants.
        question_ref: Descriptive reference text (e.g., "Bear Case Scenarios").
    """
    from docx.shared import Pt  # type: ignore[import-untyped]

    para: Any = doc.add_paragraph(style="DOBody")
    ref_run: Any = para.add_run(f"See Meeting Prep: {question_ref}")
    ref_run.italic = True
    ref_run.font.size = Pt(9)
    ref_run.font.color.rgb = ds.color_accent


def render_customary_block(
    doc: Any,
    ds: DesignSystem,
    text: str,
) -> None:
    """Render standard body text for Customary tier data.

    Just adds a DOBody paragraph -- exists for consistency and readability
    in caller code that needs to explicitly mark tier distinctions.

    Args:
        doc: The python-docx Document.
        ds: Design system with visual constants.
        text: Body text content.
    """
    para: Any = doc.add_paragraph(style="DOBody")
    para.add_run(text)


__all__ = [
    "add_meeting_prep_ref",
    "render_customary_block",
    "render_objective_signal",
    "render_scenario_context",
]
