"""Shared helpers for ground truth validation tests.

Provides state loading, nested dict navigation, financial line item
lookup, accuracy tracking, and market cap classification utilities.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, cast

# ---------------------------------------------------------------------------
# Types and paths
# ---------------------------------------------------------------------------

JsonDict = dict[str, Any]
OUTPUT_DIR = Path(__file__).parent.parent.parent / "output"


# ---------------------------------------------------------------------------
# State loading and navigation
# ---------------------------------------------------------------------------


def load_state(ticker: str) -> JsonDict | None:
    """Load state.json for a ticker, returning None if not found."""
    state_path = OUTPUT_DIR / ticker / "state.json"
    if not state_path.exists():
        return None
    with state_path.open() as f:
        result: JsonDict = json.load(f)
        return result


def has_extraction(state: JsonDict) -> bool:
    """Check if extraction stage completed (extracted is not None)."""
    return state.get("extracted") is not None


def get_nested(data: JsonDict, *keys: str) -> Any:
    """Navigate nested dicts, returning None if any key is missing."""
    current: Any = data
    for key in keys:
        if not isinstance(current, dict):
            return None
        current = cast(Any, cast(JsonDict, current).get(key))
    return current


def sourced_value(data: JsonDict, *keys: str) -> Any:
    """Get the .value from a SourcedValue at a nested path."""
    obj = get_nested(data, *keys)
    if isinstance(obj, dict):
        return cast(Any, cast(JsonDict, obj).get("value"))
    return obj


def sourced_value_from_obj(obj: Any, field: str) -> Any:
    """Extract .value from a SourcedValue field on a dict object."""
    if not isinstance(obj, dict):
        return None
    val = cast(Any, cast(JsonDict, obj).get(field))
    if isinstance(val, dict):
        return cast(Any, cast(JsonDict, val).get("value"))
    return val


def find_line_item_latest(
    state: JsonDict,
    statement_type: str,
    xbrl_concept: str,
) -> float | None:
    """Find the latest value for an XBRL concept in a financial statement."""
    stmts = get_nested(
        state, "extracted", "financials", "statements", statement_type,
    )
    if not isinstance(stmts, dict):
        return None
    line_items: list[Any] = cast(
        list[Any], cast(JsonDict, stmts).get("line_items", []),
    )
    for raw_item in line_items:
        item = cast(JsonDict, raw_item)
        if cast(Any, item.get("xbrl_concept")) == xbrl_concept:
            values = cast(JsonDict, item.get("values", {}))
            if not values:
                continue
            first_sv = cast(Any, next(iter(values.values()), None))
            if isinstance(first_sv, dict):
                val = cast(Any, cast(JsonDict, first_sv).get("value"))
                if val is not None:
                    return float(val)
    return None


def find_line_item_by_label(
    state: JsonDict,
    statement_type: str,
    *label_substrings: str,
) -> float | None:
    """Find the latest value for a line item matching label substrings."""
    stmts = get_nested(
        state, "extracted", "financials", "statements", statement_type,
    )
    if not isinstance(stmts, dict):
        return None
    line_items: list[Any] = cast(
        list[Any], cast(JsonDict, stmts).get("line_items", []),
    )
    for raw_item in line_items:
        item = cast(JsonDict, raw_item)
        label = str(item.get("label", "")).lower()
        if any(sub in label for sub in label_substrings):
            values = cast(JsonDict, item.get("values", {}))
            if not values:
                continue
            first_sv = cast(Any, next(iter(values.values()), None))
            if isinstance(first_sv, dict):
                val = cast(Any, cast(JsonDict, first_sv).get("value"))
                if val is not None:
                    return float(val)
    return None


def market_cap_tier(market_cap: float | None) -> str | None:
    """Classify market cap into tier."""
    if market_cap is None:
        return None
    if market_cap >= 200_000_000_000:
        return "MEGA"
    if market_cap >= 10_000_000_000:
        return "LARGE"
    if market_cap >= 2_000_000_000:
        return "MID"
    if market_cap >= 300_000_000:
        return "SMALL"
    return "MICRO"


# ---------------------------------------------------------------------------
# Accuracy tracking
# ---------------------------------------------------------------------------

_accuracy_results: dict[str, dict[str, bool | None]] = {}


def record(ticker: str, field: str, *, passed: bool | None) -> None:
    """Record a field comparison result for accuracy reporting."""
    if ticker not in _accuracy_results:
        _accuracy_results[ticker] = {}
    _accuracy_results[ticker][field] = passed


def print_accuracy_report(ticker: str) -> None:
    """Print accuracy summary for a ticker."""
    results = _accuracy_results.get(ticker, {})
    if not results:
        return
    total = len(results)
    n_passed = sum(1 for v in results.values() if v is True)
    n_failed = sum(1 for v in results.values() if v is False)
    n_skipped = sum(1 for v in results.values() if v is None)
    denominator = total - n_skipped
    pct = (n_passed / denominator * 100) if denominator > 0 else 0.0
    print(f"\n{'=' * 60}")
    print(
        f"  {ticker} Ground Truth Accuracy: {n_passed}/{denominator} "
        f"fields match ({pct:.0f}%)",
    )
    print(f"  Passed: {n_passed}  Failed: {n_failed}  Skipped: {n_skipped}")
    for field, result in sorted(results.items()):
        if result is True:
            status = "PASS"
        elif result is False:
            status = "FAIL"
        else:
            status = "SKIP"
        print(f"    [{status}] {field}")
    print(f"{'=' * 60}\n")


def assert_financial_close(
    ticker: str,
    field_name: str,
    actual: float | None,
    expected: float,
    *,
    rel_tol: float = 0.05,
) -> None:
    """Assert a financial value is within relative tolerance.

    Raises pytest.fail on mismatch (imported lazily to avoid import
    overhead when helpers are loaded but not used in test context).
    """
    import pytest

    if actual is None:
        record(ticker, f"financials.{field_name}", passed=False)
        pytest.fail(
            f"{field_name}: extracted value is None, expected {expected}",
        )
    is_close = abs(actual - expected) <= rel_tol * abs(expected)
    record(ticker, f"financials.{field_name}", passed=is_close)
    pct_diff = abs(actual - expected) / abs(expected) * 100
    assert is_close, (
        f"{field_name}: {actual:,.0f} vs expected {expected:,.0f} "
        f"({pct_diff:.1f}% diff, tolerance {rel_tol * 100:.0f}%)"
    )


# ---------------------------------------------------------------------------
# .docx output validation helpers
# ---------------------------------------------------------------------------


def load_docx(ticker: str) -> Any | None:
    """Load the generated .docx worksheet for a ticker."""
    from docx import Document  # type: ignore[import-untyped]

    docx_path = OUTPUT_DIR / ticker / f"{ticker}_worksheet.docx"
    if not docx_path.exists():
        return None
    return cast(Any, Document(str(docx_path)))


def read_docx_tables(doc: Any) -> list[list[list[str]]]:
    """Read all tables from a docx Document as nested string lists."""
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)
    return tables


def read_docx_text(doc: Any) -> str:
    """Read all paragraph text from a docx Document."""
    return "\n".join(cast(str, p.text) for p in doc.paragraphs)


def find_in_tables(
    tables: list[list[list[str]]],
    label: str,
) -> list[tuple[int, int, list[str]]]:
    """Find all table rows where any cell contains label (case-insensitive)."""
    results: list[tuple[int, int, list[str]]] = []
    for ti, table in enumerate(tables):
        for ri, row in enumerate(table):
            if row and any(label.lower() in cell.lower() for cell in row):
                results.append((ti, ri, row))
    return results


def find_text_containing(doc: Any, substring: str) -> list[str]:
    """Find all paragraphs containing a substring (case-insensitive)."""
    results: list[str] = []
    for p in doc.paragraphs:
        if substring.lower() in cast(str, p.text).lower():
            results.append(cast(str, p.text))
    return results
