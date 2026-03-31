"""Tests for section renderers 3-4 (Financial Health, Market/Trading).

Covers rendering with full state data, None/empty state graceful
degradation, conditional formatting in financial tables, stock drop
tables, insider trading, and stock chart generation.

Split from test_render_sections_1_4.py for 500-line compliance.

Phase 60-01: Updated to pass context dict instead of raw state.
"""

from __future__ import annotations

from datetime import UTC, datetime
from typing import Any

from docx import Document  # type: ignore[import-untyped]

from do_uw.models.common import Confidence, SourcedValue
from do_uw.models.company import CompanyIdentity, CompanyProfile
from do_uw.models.density import DensityLevel, SectionDensity
from do_uw.models.financials import (
    AuditProfile,
    DistressIndicators,
    DistressResult,
    DistressZone,
    ExtractedFinancials,
    FinancialLineItem,
    FinancialStatement,
    FinancialStatements,
    PeerCompany,
    PeerGroup,
)
from do_uw.models.market import MarketSignals
from do_uw.models.market_events import (
    EarningsGuidanceAnalysis,
    InsiderTradingAnalysis,
    StockDropAnalysis,
    StockDropEvent,
)
from do_uw.models.state import AcquiredData, AnalysisResults, AnalysisState, ExtractedData
from do_uw.stages.render.design_system import (
    DesignSystem,
    configure_matplotlib_defaults,
    setup_styles,
)

# Configure matplotlib once for all tests
configure_matplotlib_defaults()

_NOW = datetime(2025, 6, 15, tzinfo=UTC)


def _sv(value: Any, source: str = "TEST") -> SourcedValue[Any]:
    """Create a test SourcedValue."""
    return SourcedValue(
        value=value,
        source=source,
        confidence=Confidence.HIGH,
        as_of=_NOW,
    )


def _make_doc() -> Any:
    """Create a document with custom styles."""
    doc: Any = Document()
    setup_styles(doc)
    return doc


def _make_rich_state() -> AnalysisState:
    """Create a state with representative data for sections 3-4."""
    identity = CompanyIdentity(
        ticker="TEST",
        legal_name=_sv("Test Corp Inc."),
        cik=_sv("0001234567"),
        sic_code=_sv("7372"),
        sic_description=_sv("Prepackaged Software"),
        exchange=_sv("NASDAQ"),
        state_of_incorporation=_sv("DE"),
    )
    company = CompanyProfile(
        identity=identity,
        business_description=_sv("Test Corp provides enterprise software."),
        market_cap=_sv(5_000_000_000.0),
        employee_count=_sv(10000),
    )
    # Financial data
    income = FinancialStatement(
        statement_type="income",
        periods=["FY2024", "FY2023"],
        line_items=[
            FinancialLineItem(
                label="Total Revenue",
                values={
                    "FY2024": _sv(2_000_000_000.0),
                    "FY2023": _sv(1_800_000_000.0),
                },
                yoy_change=11.1,
            ),
            FinancialLineItem(
                label="Net Income",
                values={
                    "FY2024": _sv(200_000_000.0),
                    "FY2023": _sv(250_000_000.0),
                },
                yoy_change=-20.0,
            ),
        ],
    )
    distress = DistressIndicators(
        altman_z_score=DistressResult(
            score=2.5,
            zone=DistressZone.GREY,
            model_variant="original",
        ),
        piotroski_f_score=DistressResult(
            score=6.0,
            zone=DistressZone.SAFE,
            model_variant="standard",
        ),
    )
    audit = AuditProfile(
        auditor_name=_sv("Deloitte & Touche LLP"),
        is_big4=_sv(True),
        tenure_years=_sv(10),
        opinion_type=_sv("unqualified"),
    )
    peer_group = PeerGroup(
        target_ticker="TEST",
        peers=[
            PeerCompany(
                ticker="PEER1",
                name="Peer One Inc.",
                market_cap=4_500_000_000.0,
                revenue=1_900_000_000.0,
                peer_score=85.0,
            ),
        ],
    )
    financials = ExtractedFinancials(
        statements=FinancialStatements(
            income_statement=income,
            periods_available=2,
        ),
        distress=distress,
        audit=audit,
        peer_group=peer_group,
        financial_health_narrative=_sv(
            "Revenue growing but margins under pressure."
        ),
    )
    # Market data
    market = MarketSignals(
        stock_drops=StockDropAnalysis(
            single_day_drops=[
                StockDropEvent(
                    date=_sv("2025-03-15"),
                    drop_pct=_sv(-12.3),
                    drop_type="SINGLE_DAY",
                ),
            ],
        ),
        insider_analysis=InsiderTradingAnalysis(
            net_buying_selling=_sv("NET_SELLING"),
        ),
        earnings_guidance=EarningsGuidanceAnalysis(
            consecutive_miss_count=2,
            philosophy="AGGRESSIVE",
        ),
    )
    # Price history in acquired data (column-oriented format from yfinance).
    acquired = AcquiredData(
        market_data={
            "history_1y": {
                "Date": [
                    "2024-06-15", "2024-07-15", "2024-08-15",
                    "2024-09-15", "2024-10-15", "2024-11-15",
                    "2024-12-15", "2025-01-15", "2025-02-15",
                    "2025-03-15", "2025-04-15", "2025-05-15",
                    "2025-06-15",
                ],
                "Close": [
                    100.0, 98.0, 97.0, 95.0, 92.0, 90.0,
                    88.0, 85.0, 80.0, 75.0, 72.0, 68.0, 65.0,
                ],
            },
            "sector_etf": "XLK",
            "sector_history_1y": {
                "Date": [
                    "2024-06-15", "2024-07-15", "2024-08-15",
                    "2024-09-15", "2024-10-15", "2024-11-15",
                    "2024-12-15", "2025-01-15", "2025-02-15",
                    "2025-03-15", "2025-04-15", "2025-05-15",
                    "2025-06-15",
                ],
                "Close": [
                    100.0, 101.0, 101.5, 102.0, 103.0, 103.5,
                    104.0, 105.0, 105.5, 106.0, 106.5, 107.0, 108.0,
                ],
            },
            "spy_history_1y": {
                "Date": [
                    "2024-06-15", "2024-07-15", "2024-08-15",
                    "2024-09-15", "2024-10-15", "2024-11-15",
                    "2024-12-15", "2025-01-15", "2025-02-15",
                    "2025-03-15", "2025-04-15", "2025-05-15",
                    "2025-06-15",
                ],
                "Close": [
                    450.0, 455.0, 458.0, 460.0, 462.0, 465.0,
                    468.0, 470.0, 472.0, 475.0, 478.0, 480.0, 482.0,
                ],
            },
        },
    )
    # Set up analysis with financial density = ELEVATED (grey zone distress)
    analysis = AnalysisResults()
    analysis.section_densities = {
        "financial": SectionDensity(
            level=DensityLevel.ELEVATED,
            concerns=["Altman Z-Score in GREY zone"],
        ),
    }
    return AnalysisState(
        ticker="TEST",
        company=company,
        extracted=ExtractedData(
            financials=financials,
            market=market,
        ),
        acquired_data=acquired,
        analysis=analysis,
    )


def _make_context(state: AnalysisState) -> dict[str, Any]:
    """Build a context dict from state for section renderers.

    Phase 60-01: Sections now receive context dict instead of raw state.
    The _state key provides backward-compat access to the full state.
    """
    from do_uw.stages.render.md_renderer import build_template_context

    context = build_template_context(state)
    context["_state"] = state
    return context


# --------------------------------------------------------------------------
# Section 3: Financial Health
# --------------------------------------------------------------------------


class TestSection3:
    """Tests for Section 3 (Financial Health) renderer."""

    def test_render_with_financial_data(self) -> None:
        from do_uw.stages.render.sections.sect3_financial import (
            render_section_3,
        )

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        ds = DesignSystem()

        render_section_3(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        assert "Financial Health" in text
        assert "revenue" in text.lower()
        assert "Distress Indicators" in text

    def test_render_with_none_financials(self) -> None:
        from do_uw.stages.render.sections.sect3_financial import (
            render_section_3,
        )

        doc = _make_doc()
        state = AnalysisState(ticker="EMPTY")
        context = _make_context(state)
        ds = DesignSystem()

        render_section_3(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        assert "Financial Health" in text
        assert "not available" in text.lower()

    def test_distress_zone_table(self) -> None:
        from do_uw.stages.render.sections.sect3_financial import (
            render_section_3,
        )

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        ds = DesignSystem()

        render_section_3(doc, context, ds)

        tables = doc.tables
        table_texts: list[str] = []
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    table_texts.append(cell.text)
        all_table_text = " ".join(table_texts)
        assert "Altman Z-Score" in all_table_text
        assert "Grey Zone" in all_table_text

    def test_conditional_formatting_applied(self) -> None:
        """Verify that YoY change cells get conditional shading."""
        from do_uw.stages.render.sections.sect3_tables import (
            _get_conditional_shading,
        )

        # Revenue up 15% = good = blue
        assert _get_conditional_shading(15.0, "Total Revenue") == "DCEEF8"
        # Net income down 20% = bad = red
        assert _get_conditional_shading(-20.0, "Net Income") == "FCE8E6"
        # Total debt up 15% = bad = red
        assert _get_conditional_shading(15.0, "Total Debt") == "FCE8E6"
        # Small change = amber
        assert _get_conditional_shading(3.0, "Total Revenue") == "FFF3CD"
        # Trivial change = None
        assert _get_conditional_shading(0.5, "Total Revenue") is None

    def test_no_green_in_formatting(self) -> None:
        """Verify NO green hex codes appear in conditional formatting."""
        from do_uw.stages.render.sections.sect3_tables import (
            _get_conditional_shading,
        )

        test_cases = [
            (50.0, "Total Revenue"),
            (-50.0, "Total Debt"),
            (20.0, "Net Income"),
            (5.0, "Operating Margin"),
        ]
        for change, label in test_cases:
            color = _get_conditional_shading(change, label)
            if color is not None:
                assert not color.startswith("00"), (
                    f"Green in {label}: {color}"
                )
                assert "28a745" not in color.lower()
                assert "00FF00" not in color.upper()

    def test_piotroski_trajectory_criteria_format(self) -> None:
        """Piotroski trajectory shows 'X/Y criteria met' not '?:1.0 -> ?:0.0'."""
        from do_uw.stages.render.sections.sect3_financial import (
            _format_trajectory,
        )

        criteria = [
            {"criterion": "positive_ni", "score": 1.0},
            {"criterion": "improving_roa", "score": 1.0},
            {"criterion": "positive_ocf", "score": 1.0},
            {"criterion": "ocf_exceeds_ni", "score": 0.0},
            {"criterion": "decreasing_leverage", "score": 1.0},
            {"criterion": "improving_current_ratio", "score": "N/A"},
            {"criterion": "no_dilution", "score": 1.0},
            {"criterion": "improving_gross_margin", "score": 0.0},
            {"criterion": "improving_asset_turnover", "score": 1.0},
        ]
        result = _format_trajectory(criteria)
        assert result == "6/9 criteria met"
        assert "?" not in result

    def test_altman_trajectory_period_format(self) -> None:
        """Altman trajectory shows period labels with scores and trend."""
        from do_uw.stages.render.sections.sect3_financial import (
            _format_trajectory,
        )

        trajectory = [
            {"period": "FY2023", "score": 2.1, "zone": "grey"},
            {"period": "FY2024", "score": 3.2, "zone": "safe"},
        ]
        result = _format_trajectory(trajectory)
        assert "FY2023: 2.1" in result
        assert "FY2024: 3.2" in result
        assert "(improving)" in result

    def test_trajectory_empty_returns_na(self) -> None:
        """Empty trajectory returns N/A."""
        from do_uw.stages.render.sections.sect3_financial import (
            _format_trajectory,
        )

        assert _format_trajectory([]) == "N/A"

    def test_distress_zone_humanized(self) -> None:
        """Zone labels should be humanized, not raw SCREAMING_SNAKE."""
        from do_uw.stages.render.sections.sect3_financial import (
            render_section_3,
        )

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        ds = DesignSystem()

        render_section_3(doc, context, ds)

        all_table_text = " ".join(
            cell.text
            for t in doc.tables
            for row in t.rows
            for cell in row.cells
        )
        # Should have humanized labels
        assert "Grey Zone" in all_table_text
        assert "Safe" in all_table_text
        # Should NOT have raw enum values
        assert "GREY" not in all_table_text.split()
        assert "SAFE" not in all_table_text.split()

    def test_peer_group_rendered(self) -> None:
        from do_uw.stages.render.sections.sect3_financial import (
            render_section_3,
        )

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        ds = DesignSystem()

        render_section_3(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        assert "Peer Group" in text

        tables = doc.tables
        all_table_text = " ".join(
            cell.text
            for t in tables
            for row in t.rows
            for cell in row.cells
        )
        assert "PEER1" in all_table_text

    def test_format_value_share_counts_no_dollar(self) -> None:
        """Shares Outstanding and similar non-currency items have no $ prefix."""
        from do_uw.stages.render.sections.sect3_tables import _format_value

        # Share count items should NOT have $ prefix
        assert _format_value(4_100_000_000, "Shares Outstanding") == "4.1B"
        assert _format_value(15_200_000_000, "Weighted Average Shares") == "15.2B"
        assert _format_value(15_400_000_000, "Diluted Shares Outstanding") == "15.4B"
        assert _format_value(500_000_000, "Basic Shares Outstanding") == "500.0M"

        # Currency items SHOULD have $ prefix
        assert _format_value(339_000_000_000, "Total Revenue").startswith("$")
        assert _format_value(200_000_000, "Net Income").startswith("$")
        assert _format_value(50_000_000_000, "Total Assets").startswith("$")

        # EPS should have $ prefix (per-share, not compact)
        assert _format_value(6.42, "EPS (Diluted)") == "$6.42"

        # Margins should be percentages
        assert "%" in _format_value(25.3, "Gross Margin")

    def test_format_value_shares_in_table_rendering(self) -> None:
        """End-to-end: shares line items render without $ in document tables."""
        from do_uw.stages.render.sections.sect3_financial import (
            render_section_3,
        )

        state = _make_rich_state()
        # Add shares line item to income statement
        income = state.extracted.financials.statements.income_statement  # type: ignore[union-attr]
        assert income is not None
        income.line_items.append(
            FinancialLineItem(
                label="Shares Outstanding",
                values={"FY2024": _sv(4_100_000_000.0), "FY2023": _sv(4_000_000_000.0)},
                yoy_change=2.5,
            )
        )
        doc = _make_doc()
        context = _make_context(state)
        ds = DesignSystem()
        render_section_3(doc, context, ds)

        all_table_text = " ".join(
            cell.text
            for t in doc.tables
            for row in t.rows
            for cell in row.cells
        )
        # Should find "4.1B" without $ prefix
        assert "4.1B" in all_table_text
        # Should NOT find "$4.1B" -- the bug we fixed
        assert "$4.1B" not in all_table_text


# --------------------------------------------------------------------------
# Section 4: Market & Trading
# --------------------------------------------------------------------------


class TestSection4:
    """Tests for Section 4 (Market/Trading) renderer."""

    def test_render_with_market_data(self) -> None:
        from do_uw.stages.render.sections.sect4_market import (
            render_section_4,
        )

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        ds = DesignSystem()

        render_section_4(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        assert "Market & Trading" in text
        assert "Stock Drop Analysis" in text

    def test_render_with_none_market(self) -> None:
        from do_uw.stages.render.sections.sect4_market import (
            render_section_4,
        )

        doc = _make_doc()
        state = AnalysisState(ticker="EMPTY")
        context = _make_context(state)
        ds = DesignSystem()

        render_section_4(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        assert "Market & Trading" in text
        assert "not available" in text.lower()

    def test_stock_drop_table(self) -> None:
        from do_uw.stages.render.sections.sect4_market import (
            render_section_4,
        )

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        ds = DesignSystem()

        render_section_4(doc, context, ds)

        tables = doc.tables
        all_table_text = " ".join(
            cell.text
            for t in tables
            for row in t.rows
            for cell in row.cells
        )
        assert "2025-03-15" in all_table_text
        assert "SINGLE_DAY" in all_table_text

    def test_insider_trading_section(self) -> None:
        from do_uw.stages.render.sections.sect4_market import (
            render_section_4,
        )

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        ds = DesignSystem()

        render_section_4(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        assert "Insider Trading" in text

    def test_earnings_guidance_d_and_o_context(self) -> None:
        from do_uw.stages.render.sections.sect4_market import (
            render_section_4,
        )

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        ds = DesignSystem()

        render_section_4(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        assert "Consecutive Earnings Misses" in text
        assert "scienter" in text


# --------------------------------------------------------------------------
# Stock Charts
# --------------------------------------------------------------------------


class TestStockCharts:
    """Tests for stock performance chart generation."""

    def test_creates_chart_with_price_data(self) -> None:
        from do_uw.stages.render.charts.stock_charts import (
            create_stock_performance_chart,
        )

        state = _make_rich_state()
        result = create_stock_performance_chart(state, period="1Y")

        assert result is not None
        data = result.read()
        assert len(data) > 0
        assert data[:8] == b"\x89PNG\r\n\x1a\n"

    def test_returns_none_for_no_data(self) -> None:
        from do_uw.stages.render.charts.stock_charts import (
            create_stock_performance_chart,
        )

        state = AnalysisState(ticker="EMPTY")
        result = create_stock_performance_chart(state)
        assert result is None

    def test_5y_chart_wrapper(self) -> None:
        from do_uw.stages.render.charts.stock_charts import (
            create_stock_performance_chart_5y,
        )

        state = _make_rich_state()
        result = create_stock_performance_chart_5y(state)
        assert result is None

    def test_chart_with_etf_data(self) -> None:
        from do_uw.stages.render.charts.stock_charts import (
            create_stock_performance_chart,
        )

        state = _make_rich_state()
        result = create_stock_performance_chart(state, period="1Y")
        assert result is not None
        data = result.read()
        assert len(data) > 1000
