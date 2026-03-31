"""Tests for section renderers 1-2 (Executive Summary, Company Profile).

Covers rendering with full state data, None/empty state graceful
degradation. Sections 3-4 and stock chart tests are in
test_render_sections_3_4.py.

Phase 60-01: Updated to pass context dict instead of raw state.
"""

from __future__ import annotations

from datetime import UTC, datetime
from typing import Any

from docx import Document  # type: ignore[import-untyped]

from do_uw.models.common import Confidence, SourcedValue
from do_uw.models.company import CompanyIdentity, CompanyProfile
from do_uw.models.executive_summary import (
    CompanySnapshot,
    ExecutiveSummary,
    InherentRiskBaseline,
    KeyFinding,
    KeyFindings,
    UnderwritingThesis,
)
from do_uw.models.financials import (
    AuditProfile,
    DistressIndicators,
    DistressResult,
    DistressZone,
    ExtractedFinancials,
    FinancialLineItem,
    FinancialStatement,
    FinancialStatements,
    PeerCompany,
    PeerGroup,
)
from do_uw.models.market import MarketSignals
from do_uw.models.market_events import (
    EarningsGuidanceAnalysis,
    InsiderTradingAnalysis,
    StockDropAnalysis,
    StockDropEvent,
)
from do_uw.models.scoring import ScoringResult, Tier, TierClassification
from do_uw.models.scoring_output import (
    ClaimProbability,
    ProbabilityBand,
    TowerPosition,
    TowerRecommendation,
)
from do_uw.models.state import AcquiredData, AnalysisState, ExtractedData
from do_uw.stages.render.design_system import (
    DesignSystem,
    configure_matplotlib_defaults,
    setup_styles,
)

# Configure matplotlib once for all tests
configure_matplotlib_defaults()

_NOW = datetime(2025, 6, 15, tzinfo=UTC)


def _sv(value: Any, source: str = "TEST") -> SourcedValue[Any]:
    """Create a test SourcedValue."""
    return SourcedValue(
        value=value,
        source=source,
        confidence=Confidence.HIGH,
        as_of=_NOW,
    )


def _make_doc() -> Any:
    """Create a document with custom styles."""
    doc: Any = Document()
    setup_styles(doc)
    return doc


def _make_rich_state() -> AnalysisState:
    """Create a state with representative data for all 4 sections."""
    identity = CompanyIdentity(
        ticker="TEST",
        legal_name=_sv("Test Corp Inc."),
        cik=_sv("0001234567"),
        sic_code=_sv("7372"),
        sic_description=_sv("Prepackaged Software"),
        exchange=_sv("NASDAQ"),
        state_of_incorporation=_sv("DE"),
    )
    company = CompanyProfile(
        identity=identity,
        business_description=_sv("Test Corp provides enterprise software."),
        market_cap=_sv(5_000_000_000.0),
        employee_count=_sv(10000),
        section_summary=_sv("Test Corp is a mid-cap software company."),
        do_exposure_factors=[
            _sv({"factor": "Market Cap", "level": "MODERATE", "rationale": "Mid-cap"}),
            _sv({"factor": "Industry", "level": "HIGH", "rationale": "Tech sector"}),
        ],
    )
    snapshot = CompanySnapshot(
        ticker="TEST",
        company_name="Test Corp Inc.",
        market_cap=_sv(5_000_000_000.0),
        revenue=_sv(2_000_000_000.0),
        employee_count=_sv(10000),
        industry="Software",
        exchange="NASDAQ",
    )
    thesis = UnderwritingThesis(
        narrative="Test Corp is a GUIDANCE_DEPENDENT risk profile.",
        risk_type_label="GUIDANCE_DEPENDENT",
        top_factor_summary="F1: +15, F2: +10",
    )
    inherent_risk = InherentRiskBaseline(
        sector_base_rate_pct=3.5,
        market_cap_multiplier=1.2,
        market_cap_adjusted_rate_pct=4.2,
        score_multiplier=1.1,
        company_adjusted_rate_pct=4.6,
        severity_range_25th=5.0,
        severity_range_50th=12.0,
        severity_range_75th=25.0,
        severity_range_95th=60.0,
        sector_name="Technology",
        market_cap_tier="LARGE",
        methodology_note="NEEDS CALIBRATION",
    )
    key_findings = KeyFindings(
        negatives=[
            KeyFinding(
                evidence_narrative="Stock declined 35% in 6 months",
                section_origin="SECT4",
                scoring_impact="F2: +15 points",
                theory_mapping="Theory A: Disclosure",
                ranking_score=0.9,
            ),
        ],
        positives=[
            KeyFinding(
                evidence_narrative="Big 4 auditor with 10-year tenure",
                section_origin="SECT3",
                scoring_impact="F3: -5 points",
                theory_mapping="Defense: Audit Quality",
                ranking_score=0.7,
            ),
        ],
    )
    exec_summary = ExecutiveSummary(
        snapshot=snapshot,
        thesis=thesis,
        inherent_risk=inherent_risk,
        key_findings=key_findings,
    )
    scoring = ScoringResult(
        composite_score=62.0,
        quality_score=55.0,
        tier=TierClassification(
            tier=Tier.WATCH,
            score_range_low=41,
            score_range_high=60,
            action="Monitor closely",
        ),
        claim_probability=ClaimProbability(
            band=ProbabilityBand.ELEVATED,
            range_low_pct=5.0,
            range_high_pct=10.0,
            industry_base_rate_pct=3.5,
        ),
        tower_recommendation=TowerRecommendation(
            recommended_position=TowerPosition.MID_EXCESS,
            minimum_attachment="$5M",
        ),
    )
    # Financial data
    income = FinancialStatement(
        statement_type="income",
        periods=["FY2024", "FY2023"],
        line_items=[
            FinancialLineItem(
                label="Total Revenue",
                values={
                    "FY2024": _sv(2_000_000_000.0),
                    "FY2023": _sv(1_800_000_000.0),
                },
                yoy_change=11.1,
            ),
            FinancialLineItem(
                label="Net Income",
                values={
                    "FY2024": _sv(200_000_000.0),
                    "FY2023": _sv(250_000_000.0),
                },
                yoy_change=-20.0,
            ),
        ],
    )
    distress = DistressIndicators(
        altman_z_score=DistressResult(
            score=2.5,
            zone=DistressZone.GREY,
            model_variant="original",
        ),
        piotroski_f_score=DistressResult(
            score=6.0,
            zone=DistressZone.SAFE,
            model_variant="standard",
        ),
    )
    audit = AuditProfile(
        auditor_name=_sv("Deloitte & Touche LLP"),
        is_big4=_sv(True),
        tenure_years=_sv(10),
        opinion_type=_sv("unqualified"),
    )
    peer_group = PeerGroup(
        target_ticker="TEST",
        peers=[
            PeerCompany(
                ticker="PEER1",
                name="Peer One Inc.",
                market_cap=4_500_000_000.0,
                revenue=1_900_000_000.0,
                peer_score=85.0,
            ),
        ],
    )
    financials = ExtractedFinancials(
        statements=FinancialStatements(
            income_statement=income,
            periods_available=2,
        ),
        distress=distress,
        audit=audit,
        peer_group=peer_group,
        financial_health_narrative=_sv("Revenue growing but margins under pressure."),
    )
    # Market data
    market = MarketSignals(
        stock_drops=StockDropAnalysis(
            single_day_drops=[
                StockDropEvent(
                    date=_sv("2025-03-15"),
                    drop_pct=_sv(-12.3),
                    drop_type="SINGLE_DAY",
                ),
            ],
        ),
        insider_analysis=InsiderTradingAnalysis(
            net_buying_selling=_sv("NET_SELLING"),
        ),
        earnings_guidance=EarningsGuidanceAnalysis(
            consecutive_miss_count=2,
            philosophy="AGGRESSIVE",
        ),
    )
    # Price history in acquired data
    acquired = AcquiredData(
        market_data={
            "price_history": {
                "1Y": [
                    {"date": "2024-06-15", "close": 100.0},
                    {"date": "2024-09-15", "close": 95.0},
                    {"date": "2025-01-15", "close": 85.0},
                    {"date": "2025-06-15", "close": 65.0},
                ],
            },
            "sector_etf": "XLK",
            "etf_history": {
                "1Y": [
                    {"date": "2024-06-15", "close": 100.0},
                    {"date": "2024-09-15", "close": 102.0},
                    {"date": "2025-01-15", "close": 105.0},
                    {"date": "2025-06-15", "close": 108.0},
                ],
            },
        },
    )
    return AnalysisState(
        ticker="TEST",
        company=company,
        extracted=ExtractedData(
            financials=financials,
            market=market,
        ),
        executive_summary=exec_summary,
        scoring=scoring,
        acquired_data=acquired,
    )


def _make_context(state: AnalysisState) -> dict[str, Any]:
    """Build a context dict from state for section renderers.

    Phase 60-01: Sections now receive context dict instead of raw state.
    The _state key provides backward-compat access to the full state.
    """
    from do_uw.stages.render.md_renderer import build_template_context

    context = build_template_context(state)
    context["_state"] = state
    return context


# --------------------------------------------------------------------------
# Section 1: Executive Summary
# --------------------------------------------------------------------------


class TestSection1:
    """Tests for Section 1 (Executive Summary) renderer."""

    def test_render_with_full_state(self) -> None:
        from do_uw.stages.render.sections.sect1_executive import render_section_1

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        ds = DesignSystem()

        render_section_1(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        assert "Executive Summary" in text
        assert "Decile 5 of 10" in text  # market cap decile in thesis
        assert "WATCH" in text
        assert "NEEDS CALIBRATION" in text

    def test_render_with_none_executive_summary(self) -> None:
        from do_uw.stages.render.sections.sect1_executive import render_section_1

        doc = _make_doc()
        state = AnalysisState(ticker="EMPTY")
        context = _make_context(state)
        ds = DesignSystem()

        render_section_1(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        assert "Executive Summary" in text
        # New narrative builder returns "The company." for empty state
        assert "Underwriting Thesis" in text
        assert "not available" in text.lower()

    def test_data_quality_notice_when_search_not_configured(self) -> None:
        """Data quality notice appears when no search API is configured."""
        from do_uw.stages.render.sections.sect1_executive import render_section_1

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        # Default state has no search_configured flag -- notice should appear
        ds = DesignSystem()

        render_section_1(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        assert "Data Quality Notice" in text
        assert "blind spot detection was not performed" in text

    def test_data_quality_notice_hidden_when_search_configured(self) -> None:
        """Data quality notice does NOT appear when search IS configured."""
        from do_uw.stages.render.sections.sect1_executive import render_section_1

        doc = _make_doc()
        state = _make_rich_state()
        # Mark search as configured
        state.acquired_data.blind_spot_results["search_configured"] = True
        context = _make_context(state)
        ds = DesignSystem()

        render_section_1(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        assert "Data Quality Notice" not in text

    def test_key_findings_rendered(self) -> None:
        from do_uw.stages.render.sections.sect1_executive import render_section_1

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        ds = DesignSystem()

        render_section_1(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        assert "volatility" in text.lower()  # enriched narrative
        assert "clean audit opinion" in text  # enriched narrative
        assert "Underwriting Headwinds" in text  # renamed section
        assert "Risk Mitigants" in text  # renamed section

    def test_cross_section_references(self) -> None:
        """Key findings include cross-references to downstream sections."""
        from do_uw.stages.render.sections.sect1_executive import render_section_1

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        ds = DesignSystem()

        render_section_1(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        # Negative finding references Market & Trading section
        assert "(see Market & Trading)" in text or "(see Section 4: Market & Trading)" in text
        # Positive finding references Financial Health section
        assert "(see Financial Health)" in text or "(see Section 3: Financial Health)" in text

    def test_tier_classification(self) -> None:
        from do_uw.stages.render.sections.sect1_executive import render_section_1

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        ds = DesignSystem()

        render_section_1(doc, context, ds)

        # Scores are in table cells, not paragraphs
        all_table_text = " ".join(
            cell.text for t in doc.tables for row in t.rows for cell in row.cells
        )
        assert "62.0/100" in all_table_text
        assert "55.0/100" in all_table_text
        assert "WATCH" in all_table_text


# --------------------------------------------------------------------------
# Section 2: Company Profile
# --------------------------------------------------------------------------


class TestSection2:
    """Tests for Section 2 (Company Profile) renderer."""

    def test_render_with_company_data(self) -> None:
        from do_uw.stages.render.sections.sect2_company import render_section_2

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        ds = DesignSystem()

        render_section_2(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        assert "Company Profile" in text
        assert "Test Corp" in text
        assert "mid-cap software" in text

    def test_render_with_none_company(self) -> None:
        from do_uw.stages.render.sections.sect2_company import render_section_2

        doc = _make_doc()
        state = AnalysisState(ticker="EMPTY")
        context = _make_context(state)
        ds = DesignSystem()

        render_section_2(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        assert "Company Profile" in text
        assert "not available" in text.lower()

    def test_sector_consistency(self) -> None:
        """Sector label is consistent across Section 1 and Section 2."""
        from do_uw.stages.render.sections.sect1_executive import render_section_1
        from do_uw.stages.render.sections.sect2_company import render_section_2

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        ds = DesignSystem()

        render_section_1(doc, context, ds)
        render_section_2(doc, context, ds)

        # Get all text from tables (where sector appears)
        all_table_text = " ".join(
            cell.text for t in doc.tables for row in t.rows for cell in row.cells
        )
        # Inherent risk sector_name is "Technology" (set in _make_rich_state)
        assert "Technology" in all_table_text
        # Both Section 1 snapshot and Section 2 identity should show "Technology"
        # Count occurrences to verify consistency
        sector_count = all_table_text.count("Technology")
        # At least 2: one in Section 1 snapshot Sector row, one in Section 2 identity
        assert sector_count >= 2, (
            f"Expected 'Technology' in at least 2 tables, found {sector_count}"
        )

    def test_exposure_factors_rendered(self) -> None:
        from do_uw.stages.render.sections.sect2_company import render_section_2

        doc = _make_doc()
        state = _make_rich_state()
        context = _make_context(state)
        ds = DesignSystem()

        render_section_2(doc, context, ds)

        text = " ".join(p.text for p in doc.paragraphs)
        assert "D&O Exposure Factors" in text
        assert "Market Cap" in text
        assert "Industry" in text
