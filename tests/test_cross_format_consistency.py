"""Cross-format consistency tests (SC-7).

Verifies that Markdown and Word output contain the same logical sections
and that key data points appear in both formats.

PDF/HTML consistency is guaranteed by architecture (same build_template_context).
"""

from __future__ import annotations

import re
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

import pytest

from do_uw.models.classification import ClassificationResult, MarketCapTier
from do_uw.models.common import Confidence, SourcedValue
from do_uw.models.company import CompanyIdentity, CompanyProfile
from do_uw.models.financials import (
    DistressIndicators,
    DistressResult,
    DistressZone,
    ExtractedFinancials,
    FinancialLineItem,
    FinancialStatement,
    FinancialStatements,
)
from do_uw.models.governance import (
    BoardProfile,
    CompensationFlags,
    GovernanceData,
)
from do_uw.models.litigation import (
    CaseDetail,
    LitigationLandscape,
    SECEnforcementPipeline,
)
from do_uw.models.market import MarketSignals, StockPerformance
from do_uw.models.scoring import (
    FactorScore,
    ScoringResult,
    Tier,
    TierClassification,
)
from do_uw.models.state import AnalysisResults, AnalysisState, ExtractedData
from do_uw.stages.render.design_system import DesignSystem

# ---------------------------------------------------------------------------
# Canonical section registry
# ---------------------------------------------------------------------------

# Maps canonical section names to patterns that identify them in each format.
# MD uses heading regex, Word uses paragraph style heading text.
SECTION_REGISTRY: dict[str, dict[str, str]] = {
    "Executive Summary": {
        "md": r"Section 1.*Executive Summary",
        "word": r"Executive Summary",
    },
    "Company Profile": {
        "md": r"Section 2.*Company Profile",
        "word": r"(Section 2.*)?Company Profile",
    },
    "Financial Health": {
        "md": r"Section 3.*Financial",
        "word": r"(Section 3.*)?Financial",
    },
    "Market & Trading": {
        "md": r"Section 4.*Market",
        "word": r"(Section 4.*)?Market",
    },
    "Governance & Leadership": {
        "md": r"Section 5.*Governance",
        "word": r"(Section 5.*)?Governance",
    },
    "Litigation & Regulatory": {
        "md": r"Section 6.*Litigation",
        "word": r"(Section 6.*)?Litigation",
    },
    "Scoring & Risk": {
        "md": r"Section 7.*Scoring",
        "word": r"(Section 7.*)?Scoring",
    },
    "AI Transformation Risk": {
        "md": r"Section 8.*AI",
        "word": r"(Section 8.*)?AI",
    },
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _extract_md_headings(md_text: str) -> list[str]:
    """Extract all Markdown section headings (##, ###)."""
    return re.findall(r"^#{1,3}\s+(.+)$", md_text, re.MULTILINE)


def _extract_word_text(docx_path: Path) -> tuple[str, list[str]]:
    """Extract full text and heading-style paragraphs from a Word doc.

    Returns (full_text, heading_texts).
    """
    from docx import Document  # type: ignore[import-untyped]

    doc = Document(str(docx_path))
    full_parts: list[str] = []
    headings: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        full_parts.append(text)
        # Heading styles start with "DOHeading" or "Heading"
        style_name = para.style.name if para.style else ""
        if "Heading" in style_name or "DOHeading" in style_name:
            headings.append(text)

    # Also extract table cell text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    full_parts.append(cell_text)

    return "\n".join(full_parts), headings


# ---------------------------------------------------------------------------
# Rich state fixture with ELEVATED density indicators
# ---------------------------------------------------------------------------


def _build_rich_state() -> AnalysisState:
    """Build a realistic AnalysisState with data in all domains.

    Includes ELEVATED density indicators for financials and litigation
    so density-gated render paths are exercised.
    """
    now = datetime.now(tz=UTC)

    def sv(val: object, source: str = "SEC 10-K") -> SourcedValue:
        return SourcedValue(
            value=val,
            source=source,
            confidence=Confidence.HIGH,
            as_of=now,
        )

    # Company identity
    identity = CompanyIdentity(
        ticker="AAPL",
        legal_name=sv("Apple Inc."),
        cik=sv("0000320193"),
        exchange=sv("NASDAQ"),
        sector=sv("TECH"),
    )
    company = CompanyProfile(
        identity=identity,
        market_cap=sv(3000000000000.0, "yfinance"),
        employee_count=sv(164000, "SEC 10-K"),
    )

    # Financial data with distress indicators
    financials = ExtractedFinancials(
        distress=DistressIndicators(
            altman_z_score=DistressResult(score=5.23, zone=DistressZone.SAFE),
        ),
        statements=FinancialStatements(
            income_statement=FinancialStatement(
                statement_type="income",
                periods=["FY2024", "FY2023"],
                line_items=[
                    FinancialLineItem(
                        label="Total Revenue",
                        values={
                            "FY2024": sv(391035000000.0),
                            "FY2023": sv(383285000000.0),
                        },
                    ),
                ],
            ),
        ),
    )

    # Governance
    governance = GovernanceData(
        board=BoardProfile(
            size=sv(8),
            independence_ratio=sv(0.875),
            ceo_chair_duality=sv(False),
        ),
        compensation=CompensationFlags(
            say_on_pay_support_pct=sv(95.2),
        ),
    )

    # Litigation with multiple entries to trigger ELEVATED density
    cases = [
        CaseDetail(
            case_name=sv("In re Apple Securities Litigation"),
            court=sv("N.D. Cal."),
            status=sv("ACTIVE"),
        ),
        CaseDetail(
            case_name=sv("Smith v. Apple Inc."),
            court=sv("S.D.N.Y."),
            status=sv("SETTLED"),
            settlement_amount=sv(50000000.0),
        ),
    ]
    litigation = LitigationLandscape(
        active_sca_count=sv(1),
        total_case_count=sv(2),
        cases=cases,
        sec_enforcement=SECEnforcementPipeline(
            pipeline_position=sv("COMMENT_LETTER"),
        ),
    )

    # Market
    market = MarketSignals(
        stock=StockPerformance(
            current_price=sv(178.50, "yfinance"),
            high_52w=sv(199.62, "yfinance"),
            low_52w=sv(124.17, "yfinance"),
        ),
    )

    # Scoring with factor scores
    scoring = ScoringResult(
        composite_score=82.0,
        quality_score=82.0,
        total_risk_points=18.0,
        factor_scores=[
            FactorScore(
                factor_name="Litigation History",
                factor_id="F1",
                max_points=15,
                points_deducted=5.0,
            ),
            FactorScore(
                factor_name="Stock Decline",
                factor_id="F2",
                max_points=15,
                points_deducted=3.0,
            ),
        ],
        tier=TierClassification(
            tier=Tier.WANT,
            score_range_low=76,
            score_range_high=85,
        ),
    )

    # Classification
    classification = ClassificationResult(
        market_cap_tier=MarketCapTier.MEGA,
        sector_code="TECH",
        sector_name="Technology",
        years_public=44,
        base_filing_rate_pct=3.2,
        severity_band_low_m=25.0,
        severity_band_high_m=150.0,
    )

    # Analysis with triggered checks (ELEVATED density for litigation)
    analysis = AnalysisResults(
        checks_executed=200,
        checks_passed=150,
        checks_failed=20,
        checks_skipped=30,
        signal_results={
            "LIT.01": {
                "signal_name": "Active SCA",
                "status": "TRIGGERED",
                "evidence": "1 active securities class action",
                "threshold_level": "WARNING",
                "section": 6,
                "data_status": "EVALUATED",
            },
            "FIN.03": {
                "signal_name": "Revenue Decline",
                "status": "TRIGGERED",
                "evidence": "Revenue declined 2%",
                "threshold_level": "WARNING",
                "section": 3,
                "data_status": "EVALUATED",
            },
        },
    )

    state = AnalysisState(
        ticker="AAPL",
        company=company,
        extracted=ExtractedData(
            financials=financials,
            governance=governance,
            litigation=litigation,
            market=market,
        ),
        scoring=scoring,
        classification=classification,
        analysis=analysis,
    )

    return state


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestCrossFormatSectionHeadings:
    """Verify both Markdown and Word formats contain all major sections."""

    @pytest.fixture(scope="class")
    def rendered_outputs(self, tmp_path_factory: pytest.TempPathFactory) -> dict[str, Any]:
        """Render state to both Markdown and Word, returning paths and text."""
        from do_uw.stages.render.md_renderer import render_markdown
        from do_uw.stages.render.word_renderer import render_word_document

        state = _build_rich_state()
        ds = DesignSystem()
        tmp = tmp_path_factory.mktemp("cross_format")

        md_path = tmp / "output.md"
        render_markdown(state, md_path, ds)
        md_text = md_path.read_text()

        word_path = tmp / "output.docx"
        render_word_document(state, word_path, ds)
        word_full_text, word_headings = _extract_word_text(word_path)

        return {
            "md_text": md_text,
            "md_headings": _extract_md_headings(md_text),
            "word_full_text": word_full_text,
            "word_headings": word_headings,
            "md_path": md_path,
            "word_path": word_path,
        }

    def test_md_and_word_share_section_headings(self, rendered_outputs: dict[str, Any]) -> None:
        """Both formats have all 8 major canonical sections."""
        md_headings_text = " ".join(rendered_outputs["md_headings"])
        word_headings_text = " ".join(rendered_outputs["word_headings"])

        missing_md: list[str] = []
        missing_word: list[str] = []

        for section_name, patterns in SECTION_REGISTRY.items():
            if not re.search(patterns["md"], md_headings_text, re.IGNORECASE):
                missing_md.append(section_name)
            if not re.search(patterns["word"], word_headings_text, re.IGNORECASE):
                missing_word.append(section_name)

        assert not missing_md, f"Sections missing in Markdown: {missing_md}"
        assert not missing_word, f"Sections missing in Word: {missing_word}"

    def test_key_values_in_both_formats(self, rendered_outputs: dict[str, Any]) -> None:
        """Ticker, company name, and composite score appear in both formats."""
        md = rendered_outputs["md_text"]
        word = rendered_outputs["word_full_text"]

        key_values = {
            "ticker": "AAPL",
            "company_name": "Apple Inc.",
            "composite_score": "82",
        }

        for label, value in key_values.items():
            assert value in md, f"{label}={value} missing in Markdown"
            assert value in word, f"{label}={value} missing in Word"


class TestKeyDataPointsInBothFormats:
    """Verify specific data domain values appear in both formats."""

    @pytest.fixture(scope="class")
    def rendered_outputs(self, tmp_path_factory: pytest.TempPathFactory) -> dict[str, Any]:
        """Render state to both Markdown and Word."""
        from do_uw.stages.render.md_renderer import render_markdown
        from do_uw.stages.render.word_renderer import render_word_document

        state = _build_rich_state()
        ds = DesignSystem()
        tmp = tmp_path_factory.mktemp("data_points")

        md_path = tmp / "output.md"
        render_markdown(state, md_path, ds)
        md_text = md_path.read_text()

        word_path = tmp / "output.docx"
        render_word_document(state, word_path, ds)
        word_full_text, _ = _extract_word_text(word_path)

        return {"md": md_text, "word": word_full_text}

    def test_market_cap_tier_in_both(self, rendered_outputs: dict[str, Any]) -> None:
        """Market cap tier (MEGA) appears in both formats."""
        assert "MEGA" in rendered_outputs["md"] or "Mega" in rendered_outputs["md"]
        assert "MEGA" in rendered_outputs["word"] or "Mega" in rendered_outputs["word"]

    def test_tier_classification_in_both(self, rendered_outputs: dict[str, Any]) -> None:
        """Scoring tier (WANT) appears in both formats."""
        assert "WANT" in rendered_outputs["md"]
        assert "WANT" in rendered_outputs["word"]
