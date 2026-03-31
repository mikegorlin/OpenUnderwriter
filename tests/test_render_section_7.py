"""Tests for Section 7 (Risk Scoring) renderer and radar chart.

Split from test_render_sections_5_7.py for 500-line compliance.
"""

from __future__ import annotations

import io
from datetime import UTC, datetime
from typing import Any

from docx import Document  # type: ignore[import-untyped]

from do_uw.models.common import Confidence, SourcedValue
from do_uw.models.governance import GovernanceData
from do_uw.models.governance_forensics import (
    CompensationAnalysis,
    GovernanceQualityScore,
    LeadershipForensicProfile,
    LeadershipStability,
    OwnershipAnalysis,
)
from do_uw.models.litigation import (
    CaseDetail,
    LitigationLandscape,
    SECEnforcementPipeline,
)
from do_uw.models.litigation_details import (
    DefenseAssessment,
    ForumProvisions,
    LitigationTimelineEvent,
)
from do_uw.models.scoring import (
    FactorScore,
    PatternMatch,
    RedFlagResult,
    ScoringResult,
    Tier,
    TierClassification,
)
from do_uw.models.scoring_output import (
    AllegationMapping,
    AllegationTheory,
    ClaimProbability,
    FlaggedItem,
    FlagSeverity,
    LayerAssessment,
    ProbabilityBand,
    RedFlagSummary,
    RiskType,
    RiskTypeClassification,
    SeverityScenario,
    SeverityScenarios,
    TheoryExposure,
    TowerPosition,
    TowerRecommendation,
)
from do_uw.models.state import AnalysisState, ExtractedData
from do_uw.stages.render.charts.radar_chart import create_radar_chart
from do_uw.stages.render.design_system import DesignSystem, setup_styles
from do_uw.stages.render.sections.sect5_governance import render_section_5
from do_uw.stages.render.sections.sect6_litigation import render_section_6
from do_uw.stages.render.sections.sect7_scoring import render_section_7
from do_uw.stages.render.sections.sect7_scoring_detail import (
    render_scoring_detail,
)

_NOW = datetime(2025, 1, 15, tzinfo=UTC)


def _make_context(state: AnalysisState) -> dict[str, Any]:
    """Wrap AnalysisState in context dict for section renderers.

    Phase 60-02: Sections now receive context dict from build_template_context().
    Tests use this minimal wrapper with _state escape hatch.
    """
    return {"_state": state, "company_name": state.ticker}


def _sv(value: Any, source: str = "test") -> SourcedValue[Any]:
    """Create a test SourcedValue."""
    return SourcedValue(
        value=value,
        source=source,
        confidence=Confidence.MEDIUM,
        as_of=_NOW,
    )


def _make_doc() -> Any:
    """Create a Document with custom styles."""
    doc: Any = Document()
    setup_styles(doc)
    return doc


def _make_ds() -> DesignSystem:
    """Create a DesignSystem instance."""
    return DesignSystem()


def _make_scoring_data() -> ScoringResult:
    """Create test scoring data."""
    factors = [
        FactorScore(
            factor_id=f"F{i}",
            factor_name=f"Factor {i}",
            max_points=15 if i <= 3 else 10,
            points_deducted=float(i),
            evidence=[f"Evidence for F{i}"],
            rules_triggered=[f"F{i}-001"],
        )
        for i in range(1, 11)
    ]

    tier = TierClassification(
        tier=Tier.WRITE,
        score_range_low=51,
        score_range_high=70,
        action="Write with conditions",
    )

    red_flag = RedFlagResult(
        flag_id="CRF-1",
        flag_name="Active SCA",
        triggered=True,
        ceiling_applied=50,
        max_tier="WATCH",
        evidence=["Active SCA filed 2024-03"],
    )

    pattern = PatternMatch(
        pattern_id="PATTERN.STOCK.EVENT_COLLAPSE",
        pattern_name="Event-Driven Collapse",
        detected=True,
        severity="HIGH",
        triggers_matched=["stock_drop_30pct", "insider_selling"],
        score_impact={"F2": 3.0, "F7": 2.0},
    )

    risk_type = RiskTypeClassification(
        primary=RiskType.GROWTH_DARLING,
        evidence=["High revenue growth", "Tech sector"],
    )

    allegation = AllegationMapping(
        theories=[
            TheoryExposure(
                theory=AllegationTheory.A_DISCLOSURE,
                exposure_level="HIGH",
                findings=["Revenue recognition concerns"],
                factor_sources=["F1", "F3"],
            ),
            TheoryExposure(
                theory=AllegationTheory.B_GUIDANCE,
                exposure_level="MODERATE",
                findings=["Guidance miss history"],
                factor_sources=["F5"],
            ),
        ],
        primary_exposure=AllegationTheory.A_DISCLOSURE,
        concentration_analysis="Disclosure theory dominates.",
    )

    severity = SeverityScenarios(
        market_cap=10_000_000_000.0,
        scenarios=[
            SeverityScenario(
                percentile=25,
                label="favorable",
                settlement_estimate=5_000_000.0,
                defense_cost_estimate=2_000_000.0,
                total_exposure=7_000_000.0,
            ),
            SeverityScenario(
                percentile=50,
                label="median",
                settlement_estimate=15_000_000.0,
                defense_cost_estimate=5_000_000.0,
                total_exposure=20_000_000.0,
            ),
            SeverityScenario(
                percentile=95,
                label="catastrophic",
                settlement_estimate=100_000_000.0,
                defense_cost_estimate=20_000_000.0,
                total_exposure=120_000_000.0,
            ),
        ],
    )

    red_flag_summary = RedFlagSummary(
        items=[
            FlaggedItem(
                description="Active securities class action",
                source="Stanford SCAC",
                severity=FlagSeverity.CRITICAL,
                scoring_impact="F1: +15 points",
                trajectory="WORSENING",
            ),
        ],
        critical_count=1,
        high_count=0,
        moderate_count=0,
        low_count=0,
    )

    claim_prob = ClaimProbability(
        band=ProbabilityBand.HIGH,
        range_low_pct=15.0,
        range_high_pct=25.0,
        industry_base_rate_pct=8.0,
        adjustment_narrative=(
            "Elevated due to active SCA and enforcement."
        ),
    )

    tower = TowerRecommendation(
        recommended_position=TowerPosition.LOW_EXCESS,
        minimum_attachment="$5M",
        layers=[
            LayerAssessment(
                position=TowerPosition.PRIMARY,
                risk_assessment="High risk, active SCA exposure",
                premium_guidance="Above market rate",
                attachment_range="$0 - $5M",
            ),
        ],
        side_a_assessment="DIC recommended given active litigation.",
    )

    return ScoringResult(
        composite_score=55.0,
        quality_score=50.0,
        total_risk_points=45.0,
        factor_scores=factors,
        red_flags=[red_flag],
        tier=tier,
        patterns_detected=[pattern],
        risk_type=risk_type,
        allegation_mapping=allegation,
        severity_scenarios=severity,
        red_flag_summary=red_flag_summary,
        calibration_notes=[
            "Risk type needs calibration",
            "Severity model needs calibration",
        ],
        binding_ceiling_id="CRF-1",
        claim_probability=claim_prob,
        tower_recommendation=tower,
    )


# ---------------------------------------------------------------------------
# Section 7 Tests
# ---------------------------------------------------------------------------


class TestRenderSection7:
    """Tests for Section 7 (Scoring) renderer."""

    def test_render_with_data(self) -> None:
        """Section 7 renders tier, breakdown, radar, red flags, detail."""
        doc = _make_doc()
        ds = _make_ds()
        scoring = _make_scoring_data()
        state = AnalysisState(
            ticker="ACME",
            scoring=scoring,
        )
        render_section_7(doc, _make_context(state), ds)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Section 7: Risk Scoring & Synthesis" in text
        assert "Tier Classification" in text
        assert "Composite Score Breakdown" in text
        assert "10-Factor Risk Profile" in text
        assert "Red Flag Gates" in text
        assert "Risk Type Classification" in text
        assert "Severity Scenarios" in text
        assert "Calibration Notes" in text
        # From scoring detail module
        assert "Factor Detail" in text
        assert "Pattern Detection Results" in text
        assert "Allegation Theory Mapping" in text

    def test_render_none_scoring(self) -> None:
        """Section 7 gracefully handles None scoring."""
        doc = _make_doc()
        ds = _make_ds()
        state = AnalysisState(ticker="ACME", scoring=None)
        render_section_7(doc, _make_context(state), ds)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "not available" in text.lower()


class TestScoringDetail:
    """Tests for the scoring detail (sect7_scoring_detail) renderer."""

    def test_render_detail(self) -> None:
        """Scoring detail renders factors, patterns, allegation mapping."""
        doc = _make_doc()
        ds = _make_ds()
        scoring = _make_scoring_data()
        state = AnalysisState(
            ticker="ACME",
            scoring=scoring,
        )
        render_scoring_detail(doc, _make_context(state), ds)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Factor Detail" in text
        assert "Pattern Detection Results" in text
        assert "Allegation Theory Mapping" in text
        assert "Claim Probability Detail" in text
        assert "Tower Position Recommendation" in text

    def test_render_detail_none(self) -> None:
        """Scoring detail handles None scoring."""
        doc = _make_doc()
        ds = _make_ds()
        state = AnalysisState(ticker="ACME", scoring=None)
        render_scoring_detail(doc, _make_context(state), ds)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Factor Detail" not in text


class TestRadarChart:
    """Tests for the radar/spider chart."""

    def test_returns_bytesio(self) -> None:
        """Radar chart returns BytesIO with PNG data."""
        factors = [
            FactorScore(
                factor_id=f"F{i}",
                factor_name=f"Factor {i}",
                max_points=10,
                points_deducted=float(i),
            )
            for i in range(1, 11)
        ]
        ds = _make_ds()
        result = create_radar_chart(factors, ds)
        assert result is not None
        assert isinstance(result, io.BytesIO)
        data = result.read()
        assert len(data) > 100
        assert data[:4] == b"\x89PNG"

    def test_returns_none_empty(self) -> None:
        """Radar chart returns None for empty list."""
        ds = _make_ds()
        result = create_radar_chart([], ds)
        assert result is None


# ---------------------------------------------------------------------------
# Integration Test
# ---------------------------------------------------------------------------


def _make_governance_data() -> GovernanceData:
    """Create minimal governance data for integration test."""
    exec1 = LeadershipForensicProfile(
        name=_sv("John Smith"),
        title=_sv("CEO"),
        tenure_years=5.0,
    )
    leadership = LeadershipStability(
        executives=[exec1],
        avg_tenure_years=_sv(5.0),
        stability_score=_sv(72.0),
    )
    comp = CompensationAnalysis(
        ceo_total_comp=_sv(12_500_000.0),
    )
    ownership = OwnershipAnalysis(
        institutional_pct=_sv(75.0),
        insider_pct=_sv(5.0),
    )
    gov_score = GovernanceQualityScore(total_score=_sv(68.0))
    return GovernanceData(
        governance_summary=_sv("Governance adequate."),
        leadership=leadership,
        comp_analysis=comp,
        ownership=ownership,
        governance_score=gov_score,
    )


def _make_litigation_data() -> LitigationLandscape:
    """Create minimal litigation data for integration test."""
    from datetime import date

    sca = CaseDetail(
        case_name=_sv("In re Acme Corp Securities Lit."),
        filing_date=_sv(date(2024, 3, 15)),
        status=_sv("ACTIVE"),
    )
    enforcement = SECEnforcementPipeline(
        highest_confirmed_stage=_sv("COMMENT_LETTER"),
    )
    event = LitigationTimelineEvent(
        event_date=date(2024, 3, 15),
        event_type=_sv("case_filing"),
        description=_sv("SCA filed"),
        severity=_sv("HIGH"),
    )
    defense = DefenseAssessment(
        forum_provisions=ForumProvisions(
            has_federal_forum=_sv(True),
        ),
        overall_defense_strength=_sv("MODERATE"),
    )
    return LitigationLandscape(
        litigation_summary=_sv("Active SCA pending."),
        securities_class_actions=[sca],
        sec_enforcement=enforcement,
        litigation_timeline_events=[event],
        defense=defense,
    )


class TestAllSectionsIntegration:
    """Integration test: all sections 5-7 render into a single document."""

    def test_all_sections_render(self) -> None:
        """All sections 5-7 render without errors."""
        doc = _make_doc()
        ds = _make_ds()
        gov = _make_governance_data()
        lit = _make_litigation_data()
        scoring = _make_scoring_data()

        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(
                governance=gov,
                litigation=lit,
            ),
            scoring=scoring,
        )

        ctx = _make_context(state)
        render_section_5(doc, ctx, ds)
        render_section_6(doc, ctx, ds)
        render_section_7(doc, ctx, ds)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Section 5" in text
        assert "Section 6" in text
        assert "Section 7" in text

        assert len(doc.paragraphs) > 10
