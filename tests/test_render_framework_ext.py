"""Tests for RENDER stage framework: docx helpers, chart helpers,
word renderer, RenderStage integration, and pipeline integration.

Split from test_render_framework.py for 500-line compliance.
"""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document  # type: ignore[import-untyped]

from do_uw.models.common import StageStatus
from do_uw.models.state import AnalysisState
from do_uw.stages.render.chart_helpers import (
    create_figure,
    create_radar_chart,
    save_chart_to_bytes,
)
from do_uw.stages.render.design_system import (
    DesignSystem,
    configure_matplotlib_defaults,
    setup_styles,
)
from do_uw.stages.render.docx_helpers import (
    add_page_number,
    add_styled_table,
    add_toc_field,
    set_cell_shading,
)

# ---- Docx helpers tests ----


class TestSetCellShading:
    """Tests for set_cell_shading function."""

    def test_applies_shading_element(self) -> None:
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        set_cell_shading(cell, "1A1446")
        tc_pr = cell._tc.get_or_add_tcPr()
        ns = (
            "{http://schemas.openxmlformats.org/"
            "wordprocessingml/2006/main}"
        )
        shading_elements = tc_pr.findall(f"{ns}shd")
        assert len(shading_elements) == 1
        assert shading_elements[0].get(f"{ns}fill") == "1A1446"


class TestAddStyledTable:
    """Tests for add_styled_table function."""

    def test_creates_table_with_correct_dimensions(self) -> None:
        doc = Document()
        setup_styles(doc)
        ds = DesignSystem()
        headers = ["Name", "Value", "Status"]
        rows = [["Alpha", "100", "OK"], ["Beta", "200", "WARN"]]
        table = add_styled_table(doc, headers, rows, ds)
        assert len(table.rows) == 3  # 1 header + 2 data
        assert len(table.columns) == 3

    def test_header_row_has_shading(self) -> None:
        doc = Document()
        setup_styles(doc)
        ds = DesignSystem()
        table = add_styled_table(doc, ["Col"], [["val"]], ds)
        header_cell = table.rows[0].cells[0]
        tc_pr = header_cell._tc.get_or_add_tcPr()
        ns = (
            "{http://schemas.openxmlformats.org/"
            "wordprocessingml/2006/main}"
        )
        shading = tc_pr.findall(f"{ns}shd")
        assert len(shading) >= 1


class TestAddPageNumber:
    """Tests for add_page_number function."""

    def test_adds_field_elements(self) -> None:
        doc = Document()
        para = doc.add_paragraph()
        add_page_number(para)
        xml = para._p.xml
        assert "fldChar" in xml
        assert "PAGE" in xml


class TestAddTocField:
    """Tests for add_toc_field function."""

    def test_adds_toc_instruction(self) -> None:
        doc = Document()
        add_toc_field(doc)
        body_xml = doc.element.body.xml
        assert "TOC" in body_xml
        assert "fldChar" in body_xml


# ---- Chart helpers tests ----


class TestSaveChartToBytes:
    """Tests for save_chart_to_bytes function."""

    def test_returns_png_data(self) -> None:
        configure_matplotlib_defaults()
        fig, ax = create_figure()
        ax.plot([1, 2, 3], [1, 4, 9])
        buf = save_chart_to_bytes(fig)
        assert isinstance(buf, io.BytesIO)
        data = buf.read()
        assert data[:4] == b"\x89PNG"


class TestCreateRadarChart:
    """Tests for create_radar_chart function."""

    def test_returns_bytesio_with_data(self) -> None:
        configure_matplotlib_defaults()
        ds = DesignSystem()
        categories = ["F1", "F2", "F3", "F4", "F5"]
        values = [7.0, 5.0, 8.0, 3.0, 6.0]
        buf = create_radar_chart(
            categories, values, 10.0, "Test Radar", ds
        )
        assert isinstance(buf, io.BytesIO)
        data = buf.read()
        assert data[:4] == b"\x89PNG"
        assert len(data) > 1000

    def test_empty_categories(self) -> None:
        configure_matplotlib_defaults()
        ds = DesignSystem()
        buf = create_radar_chart([], [], 10.0, "", ds)
        data = buf.read()
        assert data[:4] == b"\x89PNG"


# ---- Word renderer tests ----


class TestWordRenderer:
    """Tests for word_renderer module."""

    def test_creates_valid_docx(self, tmp_path: Path) -> None:
        configure_matplotlib_defaults()
        ds = DesignSystem()
        state = AnalysisState(ticker="TEST")
        output_path = tmp_path / "test.docx"

        from do_uw.stages.render.word_renderer import render_word_document

        result = render_word_document(state, output_path, ds)
        assert result == output_path
        assert output_path.exists()
        doc = Document(str(output_path))
        assert len(doc.paragraphs) > 0

    def test_has_expected_placeholder_sections(
        self, tmp_path: Path
    ) -> None:
        configure_matplotlib_defaults()
        ds = DesignSystem()
        state = AnalysisState(ticker="AAPL")
        output_path = tmp_path / "aapl.docx"

        from do_uw.stages.render.word_renderer import render_word_document

        render_word_document(state, output_path, ds)
        doc = Document(str(output_path))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Executive Summary" in all_text
        assert "Section 2: Company Profile" in all_text
        assert "Section 7:" in all_text
        # Meeting prep is now embedded in Scoring section (HTML), not a standalone Word section


# ---- RenderStage tests ----


class TestRenderStage:
    """Tests for RenderStage integration."""

    def _make_completed_state(self) -> AnalysisState:
        """Create a state with all pre-render stages completed."""
        state = AnalysisState(ticker="TEST")
        for stage_name in [
            "resolve",
            "acquire",
            "extract",
            "analyze",
            "score",
            "benchmark",
        ]:
            state.mark_stage_running(stage_name)
            state.mark_stage_completed(stage_name)
        return state

    def test_run_creates_docx(self, tmp_path: Path) -> None:
        from do_uw.stages.render import RenderStage

        state = self._make_completed_state()
        stage = RenderStage(output_dir=tmp_path)
        stage.run(state)

        docx_path = tmp_path / "TEST_worksheet.docx"
        assert docx_path.exists()
        assert state.stages["render"].status == StageStatus.COMPLETED

    def test_validate_input_requires_benchmark(self) -> None:
        from do_uw.stages.render import RenderStage

        state = AnalysisState(ticker="TEST")
        stage = RenderStage()
        try:
            stage.validate_input(state)
            raise AssertionError("Should have raised ValueError")
        except ValueError:
            pass

    def test_validate_input_passes_with_benchmark(self) -> None:
        from do_uw.stages.render import RenderStage

        state = self._make_completed_state()
        stage = RenderStage()
        stage.validate_input(state)  # Should not raise


# ---- Pipeline integration test ----


class TestPipelineRenderIntegration:
    """Test that Pipeline passes output_dir to RenderStage."""

    def test_build_default_stages_includes_render(self) -> None:
        from do_uw.pipeline import _build_default_stages

        stages = _build_default_stages()
        names = [s.name for s in stages]
        assert "render" in names

    def test_render_stage_receives_output_dir(self) -> None:
        from do_uw.pipeline import _build_default_stages

        stages = _build_default_stages(
            output_dir=Path("/tmp/test")  # noqa: S108
        )
        render = next(s for s in stages if s.name == "render")
        assert render._output_dir == Path("/tmp/test")  # noqa: S108
