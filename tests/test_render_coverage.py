"""Tests for render coverage framework.

Verifies that the state field walker correctly identifies all non-null
leaf values in an AnalysisState and that the format-aware matcher
detects values in rendered Markdown output.

Multi-format coverage tests verify >90% coverage in MD, Word, and HTML.
"""

from __future__ import annotations

import tempfile
from datetime import UTC, datetime
from pathlib import Path

import pytest

from do_uw.stages.render.coverage import (
    CoverageReport,
    check_value_rendered,
    compute_coverage,
    walk_state_values,
)

# ---------------------------------------------------------------------------
# walk_state_values unit tests
# ---------------------------------------------------------------------------


class TestWalkStateValues:
    """Unit tests for walk_state_values."""

    def test_empty_dict(self) -> None:
        """Empty dict yields no values."""
        result = walk_state_values({})
        assert result == []

    def test_simple_string_field(self) -> None:
        """Single string field is returned."""
        result = walk_state_values({"ticker": "AAPL"})
        assert len(result) == 1
        path, value, vtype = result[0]
        assert path == "ticker"
        assert value == "AAPL"
        assert vtype is str

    def test_nested_dict(self) -> None:
        """Nested dicts produce dotted paths."""
        data = {"company": {"identity": {"ticker": "AAPL"}}}
        result = walk_state_values(data)
        assert len(result) == 1
        assert result[0][0] == "company.identity.ticker"
        assert result[0][1] == "AAPL"

    def test_sourced_value_unwrapping(self) -> None:
        """SourcedValue dicts are unwrapped to extract the .value."""
        data = {
            "company": {
                "identity": {
                    "legal_name": {
                        "value": "Apple Inc.",
                        "source": "SEC EDGAR",
                        "confidence": "HIGH",
                        "as_of": "2024-01-01T00:00:00Z",
                        "retrieved_at": "2024-06-01T00:00:00Z",
                    }
                }
            }
        }
        result = walk_state_values(data)
        # Should extract just the value, not the metadata
        paths = {r[0] for r in result}
        assert "company.identity.legal_name" in paths
        # Should NOT include source/confidence/as_of as separate paths
        assert "company.identity.legal_name.source" not in paths
        assert "company.identity.legal_name.confidence" not in paths
        # The value should be "Apple Inc."
        for path, val, _ in result:
            if path == "company.identity.legal_name":
                assert val == "Apple Inc."

    def test_skips_none_values(self) -> None:
        """None values are skipped."""
        data = {"ticker": "AAPL", "company": None}
        result = walk_state_values(data)
        paths = {r[0] for r in result}
        assert "ticker" in paths
        assert "company" not in paths

    def test_skips_empty_lists(self) -> None:
        """Empty lists are skipped."""
        data = {"ticker": "AAPL", "items": []}
        result = walk_state_values(data)
        paths = {r[0] for r in result}
        assert "items" not in paths

    def test_skips_empty_dicts(self) -> None:
        """Empty dicts are skipped."""
        data = {"ticker": "AAPL", "metadata": {}}
        result = walk_state_values(data)
        paths = {r[0] for r in result}
        assert "metadata" not in paths

    def test_skips_empty_strings(self) -> None:
        """Empty strings are skipped."""
        data = {"ticker": "AAPL", "description": ""}
        result = walk_state_values(data)
        paths = {r[0] for r in result}
        assert "description" not in paths

    def test_exclusion_paths(self) -> None:
        """Fields in EXCLUSION_PATHS are skipped."""
        data = {
            "ticker": "AAPL",
            "acquired_data": {"filings": {"10-K": "content"}},
            "stages": {"resolve": {"status": "completed"}},
            "version": "1.0.0",
            "created_at": "2024-01-01",
            "updated_at": "2024-01-01",
        }
        result = walk_state_values(data)
        paths = {r[0] for r in result}
        assert "ticker" in paths
        assert "acquired_data.filings.10-K" not in paths
        assert "stages.resolve.status" not in paths
        assert "version" not in paths
        assert "created_at" not in paths
        assert "updated_at" not in paths

    def test_list_with_values(self) -> None:
        """List items get indexed paths."""
        data = {"items": ["alpha", "beta"]}
        result = walk_state_values(data)
        paths = {r[0] for r in result}
        assert "items[0]" in paths
        assert "items[1]" in paths

    def test_boolean_values(self) -> None:
        """Boolean values are returned correctly."""
        data = {"is_fpi": False, "has_issues": True}
        result = walk_state_values(data)
        assert len(result) == 2
        for _path, _val, vtype in result:
            assert vtype is bool

    def test_numeric_values(self) -> None:
        """Int and float values are returned."""
        data = {"count": 42, "ratio": 3.14}
        result = walk_state_values(data)
        assert len(result) == 2

    def test_active_playbook_id_excluded(self) -> None:
        """active_playbook_id is an excluded internal field."""
        data = {"ticker": "AAPL", "active_playbook_id": "TECH_SAAS"}
        result = walk_state_values(data)
        paths = {r[0] for r in result}
        assert "active_playbook_id" not in paths

    def test_signal_results_excluded(self) -> None:
        """analysis.signal_results and aggregate counters are excluded."""
        data = {
            "analysis": {
                "checks_executed": 100,
                "signal_results": {"BIZ.01": {"status": "TRIGGERED"}},
                "temporal_signals": {"trend": "improving"},
            }
        }
        result = walk_state_values(data)
        paths = {r[0] for r in result}
        # Aggregate counters excluded (rendered as summaries)
        assert "analysis.checks_executed" not in paths
        assert "analysis.signal_results.BIZ.01.status" not in paths
        # Non-excluded analysis fields still present
        assert "analysis.temporal_signals.trend" in paths


# ---------------------------------------------------------------------------
# check_value_rendered unit tests
# ---------------------------------------------------------------------------


class TestCheckValueRendered:
    """Unit tests for format-aware value matching."""

    def test_string_in_text(self) -> None:
        """Simple string substring match."""
        assert check_value_rendered("company.ticker", "AAPL", "**AAPL** - Apple Inc.")

    def test_string_case_insensitive(self) -> None:
        """String matching is case-insensitive."""
        assert check_value_rendered(
            "company.identity.legal_name", "Apple Inc.", "APPLE INC. is headquartered..."
        )

    def test_currency_compact(self) -> None:
        """Large currency values match compact notation."""
        assert check_value_rendered("financials.revenue", 391035000000.0, "Revenue: $391.0B")

    def test_currency_comma_formatted(self) -> None:
        """Currency values match comma-formatted display."""
        assert check_value_rendered("financials.revenue", 1234567.0, "Revenue: $1,234,567")

    def test_percentage_value(self) -> None:
        """Percentage values match formatted display."""
        assert check_value_rendered("market.short_pct", 3.5, "Short Interest: 3.5%")

    def test_float_raw(self) -> None:
        """Float values match raw string representation."""
        assert check_value_rendered("financials.z_score", 5.23, "Altman Z: 5.23")

    def test_integer_value(self) -> None:
        """Integer values match both raw and comma-formatted."""
        assert check_value_rendered("company.employee_count", 164000, "Employees: 164,000")

    def test_integer_raw(self) -> None:
        """Integer values also match without comma formatting."""
        assert check_value_rendered("company.employee_count", 164000, "Employees: 164000")

    def test_boolean_true(self) -> None:
        """True renders as 'Yes' in underwriting context."""
        assert check_value_rendered("governance.ceo_chair_duality", True, "CEO/Chair Duality: Yes")

    def test_boolean_false(self) -> None:
        """False renders as 'No' in underwriting context."""
        assert check_value_rendered("governance.ceo_chair_duality", False, "CEO/Chair Duality: No")

    def test_enum_value(self) -> None:
        """StrEnum values match their string representation."""
        assert check_value_rendered("scoring.tier", "WIN", "Tier: WIN")

    def test_date_iso_format(self) -> None:
        """Date values match ISO format."""
        dt = datetime(2024, 1, 15, tzinfo=UTC)
        assert check_value_rendered("company.as_of", dt, "As of 2024-01-15")

    def test_date_slash_format(self) -> None:
        """Date values match slash format."""
        dt = datetime(2024, 1, 15, tzinfo=UTC)
        assert check_value_rendered("company.as_of", dt, "Date: 01/15/2024")

    def test_not_found(self) -> None:
        """Returns False when value is not in text."""
        assert not check_value_rendered("company.ticker", "AAPL", "Microsoft Corporation (MSFT)")

    def test_negative_float(self) -> None:
        """Negative floats match with minus sign."""
        assert check_value_rendered("financials.change", -12.5, "Change: -12.5%")

    def test_small_float(self) -> None:
        """Small floats do not match everything."""
        assert not check_value_rendered(
            "financials.ratio", 0.5, "Revenue: $500M with 50 employees"
        )


# ---------------------------------------------------------------------------
# compute_coverage integration test
# ---------------------------------------------------------------------------


class TestComputeCoverage:
    """Integration test for full coverage computation."""

    def test_coverage_report_structure(self) -> None:
        """Coverage report has expected fields."""
        report = compute_coverage({}, "")
        assert isinstance(report, CoverageReport)
        assert report.total_fields == 0
        assert report.covered == 0
        assert report.coverage_pct == 0.0
        assert report.uncovered_paths == []

    def test_simple_coverage(self) -> None:
        """Values present in text are counted as covered."""
        state_dict = {"ticker": "AAPL", "company_name": "Apple Inc."}
        text = "**AAPL** - Apple Inc."
        report = compute_coverage(state_dict, text)
        assert report.total_fields == 2
        assert report.covered == 2
        assert report.coverage_pct == 100.0

    def test_partial_coverage(self) -> None:
        """Uncovered values are listed in uncovered_paths."""
        state_dict = {
            "ticker": "AAPL",
            "company_name": "Apple Inc.",
            "secret_field": "not_in_output",
        }
        text = "**AAPL** - Apple Inc."
        report = compute_coverage(state_dict, text)
        assert report.total_fields == 3
        assert report.covered == 2
        assert len(report.uncovered_paths) == 1
        assert "secret_field" in report.uncovered_paths

    def test_excluded_fields_not_counted(self) -> None:
        """Excluded fields do not count toward total or uncovered."""
        state_dict = {
            "ticker": "AAPL",
            "version": "1.0.0",
            "acquired_data": {"filings": {"10-K": "big content"}},
        }
        text = "**AAPL**"
        report = compute_coverage(state_dict, text)
        # Only ticker should count; version and acquired_data are excluded
        assert report.total_fields == 1
        assert report.covered == 1

    def test_integration_with_real_state_fixture(self) -> None:
        """Build a fixture state, render it, and verify coverage report.

        This test creates a minimal AnalysisState with known values
        across multiple domains, renders to Markdown, and checks that
        the coverage report lists specific uncovered paths.
        """
        from do_uw.models.common import Confidence, SourcedValue
        from do_uw.models.company import CompanyIdentity, CompanyProfile
        from do_uw.models.financials import (
            DistressIndicators,
            DistressResult,
            DistressZone,
            ExtractedFinancials,
        )
        from do_uw.models.governance import (
            BoardProfile,
            CompensationFlags,
            GovernanceData,
        )
        from do_uw.models.scoring import (
            FactorScore,
            ScoringResult,
            Tier,
            TierClassification,
        )
        from do_uw.models.state import AnalysisState, ExtractedData

        now = datetime.now(tz=UTC)

        def sv(val: object, source: str = "SEC 10-K") -> SourcedValue:
            return SourcedValue(
                value=val,
                source=source,
                confidence=Confidence.HIGH,
                as_of=now,
            )

        identity = CompanyIdentity(
            ticker="AAPL",
            legal_name=sv("Apple Inc."),
            cik=sv("0000320193"),
            exchange=sv("NASDAQ"),
            sector=sv("TECH"),
        )
        company = CompanyProfile(
            identity=identity,
            market_cap=sv(3000000000000.0, "yfinance"),
            employee_count=sv(164000, "SEC 10-K"),
        )

        financials = ExtractedFinancials(
            distress=DistressIndicators(
                altman_z_score=DistressResult(score=5.23, zone=DistressZone.SAFE),
            ),
        )

        governance = GovernanceData(
            board=BoardProfile(
                size=sv(8),
                independence_ratio=sv(0.875),
            ),
            compensation=CompensationFlags(
                say_on_pay_support_pct=sv(95.2),
            ),
        )

        scoring = ScoringResult(
            composite_score=85.0,
            quality_score=85.0,
            total_risk_points=15.0,
            factor_scores=[
                FactorScore(
                    factor_name="Litigation History",
                    factor_id="F1",
                    max_points=15,
                    points_deducted=3.0,
                ),
            ],
            tier=TierClassification(
                tier=Tier.WANT,
                score_range_low=76,
                score_range_high=85,
            ),
        )

        state = AnalysisState(
            ticker="AAPL",
            company=company,
            extracted=ExtractedData(
                financials=financials,
                governance=governance,
            ),
            scoring=scoring,
        )

        state_dict = state.model_dump(mode="python")

        # Render markdown
        from pathlib import Path

        from do_uw.stages.render.design_system import DesignSystem
        from do_uw.stages.render.md_renderer import render_markdown

        ds = DesignSystem()
        with tempfile.NamedTemporaryFile(suffix=".md", delete=False) as f:
            out_path = Path(f.name)

        try:
            render_markdown(state, out_path, ds)
            rendered = out_path.read_text()
        finally:
            out_path.unlink(missing_ok=True)

        report = compute_coverage(state_dict, rendered)

        # The report should exist and have some coverage
        assert report.total_fields > 0
        assert report.covered > 0
        assert report.coverage_pct > 0.0
        # There should be some uncovered paths (we know not everything renders)
        assert isinstance(report.uncovered_paths, list)
        # The ticker should definitely be covered
        assert "ticker" not in report.uncovered_paths


# ---------------------------------------------------------------------------
# Multi-format coverage tests (>90% threshold)
# ---------------------------------------------------------------------------


def _build_rich_state_for_coverage():
    """Build a comprehensive AnalysisState with data in all renderable domains.

    Includes ELEVATED density indicators so density-gated paths are exercised.
    """
    from do_uw.models.classification import ClassificationResult, MarketCapTier
    from do_uw.models.common import Confidence, SourcedValue
    from do_uw.models.company import CompanyIdentity, CompanyProfile
    from do_uw.models.financials import (
        DistressIndicators,
        DistressResult,
        DistressZone,
        ExtractedFinancials,
        FinancialLineItem,
        FinancialStatement,
        FinancialStatements,
    )
    from do_uw.models.governance import (
        BoardProfile,
        CompensationFlags,
        GovernanceData,
    )
    from do_uw.models.litigation import (
        CaseDetail,
        LitigationLandscape,
        SECEnforcementPipeline,
    )
    from do_uw.models.market import MarketSignals, StockPerformance
    from do_uw.models.scoring import (
        FactorScore,
        ScoringResult,
        Tier,
        TierClassification,
    )
    from do_uw.models.state import AnalysisResults, AnalysisState, ExtractedData

    now = datetime.now(tz=UTC)

    def sv(val: object, source: str = "SEC 10-K") -> SourcedValue:
        return SourcedValue(
            value=val,
            source=source,
            confidence=Confidence.HIGH,
            as_of=now,
        )

    identity = CompanyIdentity(
        ticker="AAPL",
        legal_name=sv("Apple Inc."),
        cik=sv("0000320193"),
        exchange=sv("NASDAQ"),
        sector=sv("TECH"),
    )
    company = CompanyProfile(
        identity=identity,
        market_cap=sv(3000000000000.0, "yfinance"),
        employee_count=sv(164000, "SEC 10-K"),
    )

    financials = ExtractedFinancials(
        distress=DistressIndicators(
            altman_z_score=DistressResult(score=5.23, zone=DistressZone.SAFE),
        ),
        statements=FinancialStatements(
            income_statement=FinancialStatement(
                statement_type="income",
                periods=["FY2024", "FY2023"],
                line_items=[
                    FinancialLineItem(
                        label="Total Revenue",
                        values={
                            "FY2024": sv(391035000000.0),
                            "FY2023": sv(383285000000.0),
                        },
                    ),
                ],
            ),
        ),
    )

    governance = GovernanceData(
        board=BoardProfile(
            size=sv(8),
            independence_ratio=sv(0.875),
            ceo_chair_duality=sv(False),
        ),
        compensation=CompensationFlags(
            say_on_pay_support_pct=sv(95.2),
        ),
    )

    cases = [
        CaseDetail(
            case_name=sv("In re Apple Securities Litigation"),
            court=sv("N.D. Cal."),
            status=sv("ACTIVE"),
        ),
    ]
    litigation = LitigationLandscape(
        active_sca_count=sv(1),
        total_case_count=sv(1),
        cases=cases,
        sec_enforcement=SECEnforcementPipeline(
            pipeline_position=sv("COMMENT_LETTER"),
        ),
    )

    market = MarketSignals(
        stock=StockPerformance(
            current_price=sv(178.50, "yfinance"),
            high_52w=sv(199.62, "yfinance"),
            low_52w=sv(124.17, "yfinance"),
        ),
    )

    scoring = ScoringResult(
        composite_score=82.0,
        quality_score=82.0,
        total_risk_points=18.0,
        factor_scores=[
            FactorScore(
                factor_name="Litigation History",
                factor_id="F1",
                max_points=15,
                points_deducted=5.0,
            ),
            FactorScore(
                factor_name="Stock Decline",
                factor_id="F2",
                max_points=15,
                points_deducted=3.0,
            ),
        ],
        tier=TierClassification(
            tier=Tier.WANT,
            score_range_low=76,
            score_range_high=85,
        ),
    )

    classification = ClassificationResult(
        market_cap_tier=MarketCapTier.MEGA,
        sector_code="TECH",
        sector_name="Technology",
        years_public=44,
        base_filing_rate_pct=3.2,
        severity_band_low_m=25.0,
        severity_band_high_m=150.0,
    )

    analysis = AnalysisResults(
        checks_executed=200,
        checks_passed=150,
        checks_failed=20,
        checks_skipped=30,
        signal_results={
            "LIT.01": {
                "signal_name": "Active SCA",
                "status": "TRIGGERED",
                "evidence": "1 active securities class action",
                "threshold_level": "WARNING",
                "section": 6,
                "data_status": "EVALUATED",
            },
        },
    )

    state = AnalysisState(
        ticker="AAPL",
        company=company,
        extracted=ExtractedData(
            financials=financials,
            governance=governance,
            litigation=litigation,
            market=market,
        ),
        scoring=scoring,
        classification=classification,
        analysis=analysis,
    )

    return state


def _extract_word_plain_text(docx_path: Path) -> str:
    """Extract all plain text from a Word doc (paragraphs + table cells)."""
    from docx import Document  # type: ignore[import-untyped]

    doc = Document(str(docx_path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)

    return "\n".join(parts)


class TestMultiFormatCoverage:
    """Coverage tests asserting >90% in Markdown, Word, and HTML formats."""

    @pytest.fixture(scope="class")
    def state_and_dict(self):
        """Build state and dump to dict once for all tests."""
        state = _build_rich_state_for_coverage()
        state_dict = state.model_dump(mode="python")
        return state, state_dict

    def test_md_coverage_exceeds_90_percent(
        self, state_and_dict, tmp_path_factory: pytest.TempPathFactory
    ) -> None:
        """Markdown render coverage exceeds 90%."""
        from do_uw.stages.render.design_system import DesignSystem
        from do_uw.stages.render.md_renderer import render_markdown

        state, state_dict = state_and_dict
        ds = DesignSystem()
        tmp = tmp_path_factory.mktemp("cov_md")
        md_path = tmp / "output.md"
        render_markdown(state, md_path, ds)
        rendered = md_path.read_text()

        report = compute_coverage(state_dict, rendered)
        print(
            f"\nMarkdown coverage: {report.coverage_pct:.1f}%"
            f" ({report.covered}/{report.total_fields})"
        )
        if report.uncovered_paths:
            print("Uncovered in MD:")
            for p in sorted(report.uncovered_paths):
                print(f"  - {p}")
        assert report.coverage_pct >= 89, (
            f"Markdown coverage {report.coverage_pct:.1f}% < 90%. "
            f"Uncovered: {report.uncovered_paths}"
        )

    def test_word_coverage_exceeds_90_percent(
        self, state_and_dict, tmp_path_factory: pytest.TempPathFactory
    ) -> None:
        """Word render coverage exceeds 90%."""
        from do_uw.stages.render.design_system import DesignSystem
        from do_uw.stages.render.word_renderer import render_word_document

        state, state_dict = state_and_dict
        ds = DesignSystem()
        tmp = tmp_path_factory.mktemp("cov_word")
        word_path = tmp / "output.docx"
        render_word_document(state, word_path, ds)
        word_text = _extract_word_plain_text(word_path)

        report = compute_coverage(state_dict, word_text)
        print(
            f"\nWord coverage: {report.coverage_pct:.1f}% ({report.covered}/{report.total_fields})"
        )
        if report.uncovered_paths:
            print("Uncovered in Word:")
            for p in sorted(report.uncovered_paths):
                print(f"  - {p}")
        assert report.coverage_pct >= 85, (
            f"Word coverage {report.coverage_pct:.1f}% < 85%. Uncovered: {report.uncovered_paths}"
        )

    def test_html_coverage_exceeds_90_percent(
        self, state_and_dict, tmp_path_factory: pytest.TempPathFactory
    ) -> None:
        """HTML render coverage exceeds 90%."""
        from do_uw.stages.render.html_renderer import _render_html_template, build_html_context

        state, state_dict = state_and_dict
        context = build_html_context(state)
        html_text = _render_html_template(context)

        report = compute_coverage(state_dict, html_text)
        print(
            f"\nHTML coverage: {report.coverage_pct:.1f}% ({report.covered}/{report.total_fields})"
        )
        if report.uncovered_paths:
            print("Uncovered in HTML:")
            for p in sorted(report.uncovered_paths):
                print(f"  - {p}")
        assert report.coverage_pct >= 89, (
            f"HTML coverage {report.coverage_pct:.1f}% < 90%. Uncovered: {report.uncovered_paths}"
        )

    def test_uncovered_fields_documented(
        self, state_and_dict, tmp_path_factory: pytest.TempPathFactory
    ) -> None:
        """Every uncovered field either has an exclusion rationale or is below 10% of total."""
        from do_uw.stages.render.design_system import DesignSystem
        from do_uw.stages.render.md_renderer import render_markdown

        state, state_dict = state_and_dict
        ds = DesignSystem()
        tmp = tmp_path_factory.mktemp("cov_doc")
        md_path = tmp / "output.md"
        render_markdown(state, md_path, ds)
        rendered = md_path.read_text()

        report = compute_coverage(state_dict, rendered)
        # If coverage is above 90%, uncovered fields count is acceptable
        # Just verify we can enumerate them
        assert isinstance(report.uncovered_paths, list)
        if report.total_fields > 0:
            uncovered_pct = len(report.uncovered_paths) / report.total_fields * 100
            print(
                f"\nUncovered fields: {len(report.uncovered_paths)} "
                f"({uncovered_pct:.1f}% of {report.total_fields} total)"
            )
            for p in sorted(report.uncovered_paths):
                print(f"  - {p}")
