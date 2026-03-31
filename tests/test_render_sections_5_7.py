"""Tests for Section 5-6 renderers and associated charts.

Tests cover:
- render_section_5 with governance data and None
- create_ownership_chart with data and None
- render_section_6 with litigation data and None
- create_litigation_timeline with data and None
- Defense assessment and litigation details

Section 7, radar chart, and integration tests are in
test_render_section_7.py.

Phase 60-02: Updated to pass context dict instead of raw state.
"""

from __future__ import annotations

import io
from datetime import UTC, date, datetime
from typing import Any

from docx import Document  # type: ignore[import-untyped]

from do_uw.models.common import Confidence, SourcedValue
from do_uw.models.governance import GovernanceData
from do_uw.models.governance_forensics import (
    BoardForensicProfile,
    CompensationAnalysis,
    GovernanceQualityScore,
    LeadershipForensicProfile,
    LeadershipStability,
    NarrativeCoherence,
    OwnershipAnalysis,
    SentimentProfile,
)
from do_uw.models.litigation import (
    CaseDetail,
    LitigationLandscape,
    SECEnforcementPipeline,
)
from do_uw.models.litigation_details import (
    ContingentLiability,
    DefenseAssessment,
    ForumProvisions,
    IndustryClaimPattern,
    LitigationTimelineEvent,
    SOLWindow,
)
from do_uw.models.density import DensityLevel, SectionDensity
from do_uw.models.financials import AuditProfile, ExtractedFinancials
from do_uw.models.state import AcquiredData, AnalysisResults, AnalysisState, ExtractedData
from do_uw.stages.render.charts.ownership_chart import create_ownership_chart
from do_uw.stages.render.charts.timeline_chart import (
    create_litigation_timeline,
)
from do_uw.stages.render.design_system import DesignSystem, setup_styles
from do_uw.stages.render.sections.sect3_audit import render_audit_risk
from do_uw.stages.render.sections.sect5_governance import render_section_5
from do_uw.stages.render.sections.sect5_governance_board import (
    clean_board_name as _clean_board_name,
)
from do_uw.stages.render.sections.sect6_defense import (
    render_defense_assessment,
)
from do_uw.stages.render.sections.sect6_litigation import render_section_6
from do_uw.stages.render.sections.sect6_timeline import (
    render_litigation_details,
)

_NOW = datetime(2025, 1, 15, tzinfo=UTC)


def _sv(value: Any, source: str = "test") -> SourcedValue[Any]:
    """Create a test SourcedValue."""
    return SourcedValue(
        value=value,
        source=source,
        confidence=Confidence.MEDIUM,
        as_of=_NOW,
    )


def _make_doc() -> Any:
    """Create a Document with custom styles."""
    doc: Any = Document()
    setup_styles(doc)
    return doc


def _make_ds() -> DesignSystem:
    """Create a DesignSystem instance."""
    return DesignSystem()


def _make_context(state: AnalysisState) -> dict[str, Any]:
    """Wrap AnalysisState in context dict for section renderers.

    Phase 60-02: Sections now receive context dict from build_template_context().
    Tests use this minimal wrapper with _state escape hatch.
    """
    return {"_state": state, "company_name": state.ticker}


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


def _make_governance_data() -> GovernanceData:
    """Create test governance data."""
    exec1 = LeadershipForensicProfile(
        name=_sv("John Smith"),
        title=_sv("CEO"),
        tenure_years=5.0,
        departure_type="ACTIVE",
        prior_litigation=[_sv("In re OldCorp Securities Lit.")],
    )
    exec2 = LeadershipForensicProfile(
        name=_sv("Jane Doe"),
        title=_sv("CFO"),
        tenure_years=2.5,
        departure_type="ACTIVE",
    )

    leadership = LeadershipStability(
        executives=[exec1, exec2],
        avg_tenure_years=_sv(3.75),
        stability_score=_sv(72.0),
    )

    board_member = BoardForensicProfile(
        name=_sv("Bob Director"),
        tenure_years=_sv(8.0),
        is_independent=_sv(True),
        committees=["Audit", "Compensation"],
        is_overboarded=False,
    )

    comp = CompensationAnalysis(
        ceo_total_comp=_sv(12_500_000.0),
        ceo_pay_ratio=_sv(250.0),
        say_on_pay_pct=_sv(85.0),
        has_clawback=_sv(True),
    )

    ownership = OwnershipAnalysis(
        institutional_pct=_sv(75.0),
        insider_pct=_sv(5.0),
        top_holders=[
            _sv({"name": "Vanguard", "pct": 8.5}),
            _sv({"name": "BlackRock", "pct": 7.2}),
        ],
    )

    sentiment = SentimentProfile(
        management_tone_trajectory=_sv("STABLE"),
        hedging_language_trend=_sv("INCREASING"),
    )

    coherence = NarrativeCoherence(
        overall_assessment=_sv("MINOR_GAPS"),
    )

    gov_score = GovernanceQualityScore(
        total_score=_sv(68.0),
    )

    return GovernanceData(
        governance_summary=_sv("Governance is adequate with minor concerns."),
        leadership=leadership,
        board_forensics=[board_member],
        comp_analysis=comp,
        ownership=ownership,
        sentiment=sentiment,
        narrative_coherence=coherence,
        governance_score=gov_score,
    )


def _make_litigation_data() -> LitigationLandscape:
    """Create test litigation data."""
    sca = CaseDetail(
        case_name=_sv("In re Acme Corp Securities Lit."),
        filing_date=_sv(date(2024, 3, 15)),
        status=_sv("ACTIVE"),
        allegations=[_sv("10b-5"), _sv("Section 11")],
        lead_counsel=_sv("Robbins Geller"),
        lead_counsel_tier=_sv(1),
    )

    enforcement = SECEnforcementPipeline(
        highest_confirmed_stage=_sv("COMMENT_LETTER"),
        comment_letter_count=_sv(3),
        comment_letter_topics=[_sv("Revenue recognition")],
    )

    timeline_event = LitigationTimelineEvent(
        event_date=date(2024, 3, 15),
        event_type=_sv("case_filing"),
        description=_sv("SCA filed alleging 10b-5 violations"),
        severity=_sv("HIGH"),
    )
    timeline_event_2 = LitigationTimelineEvent(
        event_date=date(2024, 6, 1),
        event_type=_sv("regulatory"),
        description=_sv("SEC comment letter received"),
        severity=_sv("MEDIUM"),
    )
    timeline_event_3 = LitigationTimelineEvent(
        event_date=date(2024, 9, 10),
        event_type=_sv("settlement"),
        description=_sv("Partial settlement reached"),
        severity=_sv("MODERATE"),
    )

    defense = DefenseAssessment(
        forum_provisions=ForumProvisions(
            has_federal_forum=_sv(True),
            has_exclusive_forum=_sv(True),
        ),
        pslra_safe_harbor_usage=_sv("MODERATE"),
        overall_defense_strength=_sv("MODERATE"),
    )

    sol = SOLWindow(
        claim_type="10b-5",
        trigger_date=date(2024, 3, 15),
        sol_years=2,
        repose_years=5,
        sol_expiry=date(2026, 3, 15),
        repose_expiry=date(2029, 3, 15),
        window_open=True,
    )

    pattern = IndustryClaimPattern(
        legal_theory=_sv("RULE_10B5"),
        description=_sv("Revenue recognition claims in tech sector"),
        this_company_exposed=_sv(True),
    )

    contingency = ContingentLiability(
        description=_sv("Patent infringement"),
        asc_450_classification=_sv("reasonably_possible"),
        range_low=_sv(5_000_000.0),
        range_high=_sv(15_000_000.0),
    )

    return LitigationLandscape(
        litigation_summary=_sv("Active SCA pending with moderate defense."),
        securities_class_actions=[sca],
        sec_enforcement=enforcement,
        litigation_timeline_events=[timeline_event, timeline_event_2, timeline_event_3],
        defense=defense,
        sol_map=[sol],
        industry_patterns=[pattern],
        contingent_liabilities=[contingency],
        active_matter_count=_sv(1),
        historical_matter_count=_sv(3),
        total_litigation_reserve=_sv(10_000_000.0),
    )


# ---------------------------------------------------------------------------
# Section 5 Tests
# ---------------------------------------------------------------------------


class TestRenderSection5:
    """Tests for Section 5 (Governance) renderer."""

    def test_render_with_data(self) -> None:
        """Section 5 renders leadership, board, compensation, ownership."""
        doc = _make_doc()
        ds = _make_ds()
        gov = _make_governance_data()
        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(governance=gov),
        )
        render_section_5(doc, _make_context(state), ds)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Section 5: Governance & Leadership" in text
        assert "Leadership Stability" in text
        assert "John Smith" in text
        assert "Board Composition" in text
        assert "Compensation Analysis" in text
        assert "Ownership Structure" in text
        assert "Sentiment" in text

    def test_render_none_governance(self) -> None:
        """Section 5 gracefully handles None governance."""
        doc = _make_doc()
        ds = _make_ds()
        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(governance=None),
        )
        render_section_5(doc, _make_context(state), ds)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "not available" in text.lower()

    def test_render_none_extracted(self) -> None:
        """Section 5 handles None extracted data."""
        doc = _make_doc()
        ds = _make_ds()
        state = AnalysisState(ticker="ACME", extracted=None)
        render_section_5(doc, _make_context(state), ds)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "not available" in text.lower()


class TestOwnershipChart:
    """Tests for the ownership donut chart."""

    def test_returns_bytesio(self) -> None:
        """Chart returns BytesIO with PNG data."""
        ownership = OwnershipAnalysis(
            institutional_pct=_sv(70.0),
            insider_pct=_sv(10.0),
        )
        ds = _make_ds()
        result = create_ownership_chart(ownership, ds)
        assert result is not None
        assert isinstance(result, io.BytesIO)
        data = result.read()
        assert len(data) > 100
        # PNG magic bytes
        assert data[:4] == b"\x89PNG"

    def test_returns_none_for_none(self) -> None:
        """Chart returns None for None input."""
        ds = _make_ds()
        result = create_ownership_chart(None, ds)
        assert result is None

    def test_returns_none_for_empty(self) -> None:
        """Chart returns None when no ownership percentages."""
        ownership = OwnershipAnalysis()
        ds = _make_ds()
        result = create_ownership_chart(ownership, ds)
        assert result is None


class TestCleanBoardName:
    """Tests for _clean_board_name helper."""

    def test_clean_name_no_change(self) -> None:
        """Normal name passes through unchanged."""
        assert _clean_board_name("John Smith") == "John Smith"

    def test_strips_inc_suffix(self) -> None:
        """Removes ', Inc.' company suffix from concatenated names."""
        assert _clean_board_name("John Smith, Apple Inc.") == "John Smith"

    def test_strips_corp_suffix(self) -> None:
        """Removes ', Corp.' company suffix."""
        assert _clean_board_name("Jane Doe, Acme Corp.") == "Jane Doe"

    def test_strips_llc_suffix(self) -> None:
        """Removes ', LLC' company suffix."""
        assert _clean_board_name("Bob Jones, Widgets LLC") == "Bob Jones"

    def test_strips_independent_parenthetical(self) -> None:
        """Removes (Independent) role annotation."""
        assert _clean_board_name("Alice Brown (Independent)") == "Alice Brown"

    def test_strips_non_independent_parenthetical(self) -> None:
        """Removes (Non-Independent) role annotation."""
        result = _clean_board_name("Alice Brown (Non-Independent)")
        assert result == "Alice Brown"

    def test_preserves_na(self) -> None:
        """N/A passes through."""
        assert _clean_board_name("N/A") == "N/A"

    def test_preserves_empty(self) -> None:
        """Empty string passes through."""
        assert _clean_board_name("") == ""

    def test_board_name_used_in_render(self) -> None:
        """Board member names with company artifacts are cleaned in render."""
        doc = _make_doc()
        ds = _make_ds()
        board_member = BoardForensicProfile(
            name=_sv("Bob Director, Acme Inc."),
            tenure_years=_sv(8.0),
            is_independent=_sv(True),
            committees=["Audit"],
            is_overboarded=False,
        )
        gov = GovernanceData(
            board_forensics=[board_member],
        )
        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(governance=gov),
        )
        render_section_5(doc, _make_context(state), ds)
        # Board names are in tables, not paragraphs
        table_text = ""
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    table_text += cell.text + " "
        # Cleaned name should appear, not the original with company
        assert "Bob Director" in table_text
        assert "Acme Inc." not in table_text


# ---------------------------------------------------------------------------
# Section 6 Tests
# ---------------------------------------------------------------------------


class TestRenderSection6:
    """Tests for Section 6 (Litigation) renderer."""

    def test_render_with_data(self) -> None:
        """Section 6 renders SCA table, enforcement, timeline."""
        doc = _make_doc()
        ds = _make_ds()
        lit = _make_litigation_data()
        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(litigation=lit),
        )
        render_section_6(doc, _make_context(state), ds)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Section 6: Litigation & Regulatory Exposure" in text
        assert "Securities Class Actions" in text
        assert "SEC Enforcement Pipeline" in text
        assert "Defense Strength Assessment" in text

    def test_render_none_litigation(self) -> None:
        """Section 6 gracefully handles None litigation."""
        doc = _make_doc()
        ds = _make_ds()
        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(litigation=None),
        )
        render_section_6(doc, _make_context(state), ds)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "not available" in text.lower()


class TestLitigationDetails:
    """Tests for the litigation details (sect6_timeline) renderer."""

    def test_render_details(self) -> None:
        """Litigation details renders derivative, regulatory, SOL, patterns."""
        doc = _make_doc()
        ds = _make_ds()
        lit = _make_litigation_data()
        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(litigation=lit),
        )
        render_litigation_details(doc, _make_context(state), ds)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Derivative Suits" in text
        assert "Regulatory Proceedings" in text
        assert "Industry Claim Patterns" in text
        assert "Statute of Limitations Map" in text


class TestDefenseAssessment:
    """Tests for the defense assessment (sect6_defense) renderer."""

    def test_render_defense(self) -> None:
        """Defense assessment renders strength, contingencies, whistleblower."""
        doc = _make_doc()
        ds = _make_ds()
        lit = _make_litigation_data()
        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(litigation=lit),
        )
        render_defense_assessment(doc, _make_context(state), ds)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Defense Strength Assessment" in text
        assert "Contingent Liabilities" in text
        assert "Whistleblower Indicators" in text

    def test_render_defense_none(self) -> None:
        """Defense assessment handles None litigation."""
        doc = _make_doc()
        ds = _make_ds()
        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(litigation=None),
        )
        render_defense_assessment(doc, _make_context(state), ds)
        # No crash, no paragraphs added
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Defense Strength" not in text


class TestLitigationTimeline:
    """Tests for the litigation timeline chart."""

    def test_returns_bytesio(self) -> None:
        """Timeline chart returns BytesIO with PNG data."""
        lit = _make_litigation_data()
        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(litigation=lit),
        )
        ds = _make_ds()
        result = create_litigation_timeline(state, ds)
        assert result is not None
        assert isinstance(result, io.BytesIO)
        data = result.read()
        assert data[:4] == b"\x89PNG"

    def test_returns_none_no_events(self) -> None:
        """Timeline returns None when no events."""
        lit = LitigationLandscape()
        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(litigation=lit),
        )
        ds = _make_ds()
        result = create_litigation_timeline(state, ds)
        assert result is None

    def test_returns_none_sparse_events(self) -> None:
        """Timeline returns None when <=2 events (sparse gate)."""
        evt = LitigationTimelineEvent(
            event_date=date(2024, 3, 15),
            event_type=_sv("case_filing"),
            description=_sv("SCA filed"),
            severity=_sv("HIGH"),
        )
        lit = LitigationLandscape(litigation_timeline_events=[evt])
        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(litigation=lit),
        )
        ds = _make_ds()
        result = create_litigation_timeline(state, ds)
        assert result is None, "Chart should not render for <=2 events"


# ---------------------------------------------------------------------------
# Auditor Display Tests
# ---------------------------------------------------------------------------


class TestAuditorDisplay:
    """Tests for auditor identity display in sect3_audit."""

    def test_auditor_name_shown_when_available(self) -> None:
        """Auditor name displays correctly when populated.

        Clean audits (no material weaknesses, going concern, restatements)
        render concisely, so auditor name appears in paragraph text.
        """
        doc = _make_doc()
        ds = _make_ds()
        audit = AuditProfile(auditor_name=_sv("BDO USA, LLP"))
        financials = ExtractedFinancials(audit=audit)
        state = AnalysisState(
            ticker="SMCI",
            extracted=ExtractedData(financials=financials),
        )
        render_audit_risk(doc, _make_context(state), ds)
        # Clean audits render concisely (paragraph, not table)
        all_text = ""
        for para in doc.paragraphs:
            all_text += para.text + " "
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        assert "BDO USA, LLP" in all_text

    def test_auditor_helpful_fallback_when_none(self) -> None:
        """Shows helpful message instead of raw N/A when auditor unknown.

        Clean audits render concisely, so the fallback appears in
        paragraph text rather than a table.
        """
        doc = _make_doc()
        ds = _make_ds()
        audit = AuditProfile()  # No auditor name
        financials = ExtractedFinancials(audit=audit)
        state = AnalysisState(
            ticker="SMCI",
            extracted=ExtractedData(financials=financials),
        )
        render_audit_risk(doc, _make_context(state), ds)
        all_text = ""
        for para in doc.paragraphs:
            all_text += para.text + " "
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        # Should NOT show raw "N/A" for auditor
        assert "Not identified" in all_text


# ---------------------------------------------------------------------------
# Litigation Data Source Notice Tests
# ---------------------------------------------------------------------------


class TestLitigationDataSourceNotice:
    """Tests for the web search notice in litigation section."""

    def test_no_blind_spot_shows_notice(self) -> None:
        """When blind spot results empty, shows data source notice."""
        doc = _make_doc()
        ds = _make_ds()
        lit = _make_litigation_data()
        acquired = AcquiredData(
            blind_spot_results={"pre_structured": {}, "post_structured": {}},
        )
        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(litigation=lit),
            acquired_data=acquired,
        )
        render_section_6(doc, _make_context(state), ds)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "blind spot detection was not performed" in text.lower()

    def test_with_blind_spot_no_notice(self) -> None:
        """When blind spot results present, no data source notice."""
        doc = _make_doc()
        ds = _make_ds()
        lit = _make_litigation_data()
        acquired = AcquiredData(
            blind_spot_results={
                "pre_structured": {"litigation": [{"title": "test"}]},
                "post_structured": {},
            },
        )
        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(litigation=lit),
            acquired_data=acquired,
        )
        render_section_6(doc, _make_context(state), ds)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "blind spot detection was not performed" not in text.lower()


class TestEmptySCAMessage:
    """Tests for improved empty SCA section message."""

    def test_no_sca_shows_helpful_message(self) -> None:
        """Clean litigation renders concise summary instead of full SCA table.

        When no active litigation exists, density gating produces a brief
        summary rather than empty forensic tables.
        Requires section_densities to be populated by ANALYZE stage.
        """
        doc = _make_doc()
        ds = _make_ds()
        lit = LitigationLandscape(
            sec_enforcement=SECEnforcementPipeline(),
        )
        analysis = AnalysisResults()
        analysis.section_densities = {
            "litigation": SectionDensity(level=DensityLevel.CLEAN),
        }
        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(litigation=lit),
            analysis=analysis,
        )
        render_section_6(doc, _make_context(state), ds)
        text = "\n".join(p.text for p in doc.paragraphs)
        # Clean litigation now renders concisely
        assert "no litigation exposure identified" in text.lower()

    def test_active_sca_shows_full_detail(self) -> None:
        """Active litigation renders full forensic detail with SCA table."""
        doc = _make_doc()
        ds = _make_ds()
        lit = _make_litigation_data()  # Has active SCA
        state = AnalysisState(
            ticker="ACME",
            extracted=ExtractedData(litigation=lit),
        )
        render_section_6(doc, _make_context(state), ds)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "securities class actions" in text.lower()
        assert "sec enforcement pipeline" in text.lower()


