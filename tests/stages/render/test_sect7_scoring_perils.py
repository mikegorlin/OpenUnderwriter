"""Tests for Section 7 peril-organized scoring renderer.

Verifies peril summary table, deep dive sections, chain narratives,
and graceful handling of missing/empty peril data.
"""

from __future__ import annotations

from typing import Any

from docx import Document  # type: ignore[import-untyped]

from do_uw.stages.render.design_system import DesignSystem, setup_styles
from do_uw.stages.render.sections.sect7_scoring_perils import (
    render_peril_deep_dives,
    render_peril_summary,
)


def _make_ds() -> DesignSystem:
    return DesignSystem()


def _make_doc() -> Any:
    doc = Document()
    setup_styles(doc)
    return doc


def _get_all_text(doc: Any) -> str:
    """Extract all text from a document."""
    texts: list[str] = []
    for para in doc.paragraphs:
        texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)


def _make_peril_data() -> dict[str, Any]:
    """Create full peril scoring data matching extract_peril_scoring output."""
    return {
        "perils": [
            {
                "peril_id": "P1",
                "name": "Securities Class Action",
                "description": "Shareholder lawsuits for securities fraud",
                "risk_level": "HIGH",
                "active_chain_count": 2,
                "total_chain_count": 3,
                "frequency": "common",
                "severity": "high",
                "typical_settlement_range": "$5M-$500M",
                "key_drivers": ["stock_drop", "restatement"],
                "key_evidence": ["Stock drop >15%", "Revenue restatement"],
                "chains": [
                    {
                        "chain_id": "C1",
                        "name": "Disclosure Fraud Chain",
                        "peril_id": "P1",
                        "description": "Material misstatement leads to corrective disclosure and stock drop",
                        "active": True,
                        "risk_level": "HIGH",
                        "triggered_triggers": ["FIN.ACCT.restatement", "MKT.STOCK.drop"],
                        "active_amplifiers": ["GOV.INSIDER.trading"],
                        "active_mitigators": [],
                        "total_triggers": 3,
                        "total_amplifiers": 2,
                        "total_mitigators": 1,
                        "evidence_summary": ["Revenue restatement in Q3"],
                        "evidence_map": {
                            "FIN.ACCT.restatement": "Q3 10-Q/A revenue restatement",
                            "MKT.STOCK.drop": "18% decline on disclosure day",
                        },
                        "historical_filing_rate": 0.15,
                        "median_severity_usd": 25_000_000,
                    },
                    {
                        "chain_id": "C2",
                        "name": "Earnings Guidance Chain",
                        "peril_id": "P1",
                        "description": "Misleading guidance followed by earnings miss",
                        "active": True,
                        "risk_level": "MODERATE",
                        "triggered_triggers": ["FIN.TEMPORAL.earnings_miss"],
                        "active_amplifiers": [],
                        "active_mitigators": ["GOV.BOARD.independence"],
                        "total_triggers": 2,
                        "total_amplifiers": 1,
                        "total_mitigators": 1,
                        "evidence_summary": ["EPS missed by 20%"],
                        "evidence_map": {},
                        "historical_filing_rate": 0.08,
                        "median_severity_usd": 10_000_000,
                    },
                    {
                        "chain_id": "C3",
                        "name": "Insider Trading Chain",
                        "peril_id": "P1",
                        "description": "Suspicious insider sales before bad news",
                        "active": False,
                        "risk_level": "LOW",
                        "triggered_triggers": [],
                        "active_amplifiers": [],
                        "active_mitigators": [],
                        "total_triggers": 2,
                        "total_amplifiers": 1,
                        "total_mitigators": 0,
                        "evidence_summary": [],
                        "evidence_map": {},
                        "historical_filing_rate": None,
                        "median_severity_usd": None,
                    },
                ],
            },
            {
                "peril_id": "P5",
                "name": "Regulatory Enforcement",
                "description": "Government enforcement actions",
                "risk_level": "MODERATE",
                "active_chain_count": 1,
                "total_chain_count": 2,
                "frequency": "occasional",
                "severity": "moderate",
                "typical_settlement_range": "$1M-$50M",
                "key_drivers": ["sec_investigation"],
                "key_evidence": ["Active SEC investigation"],
                "chains": [
                    {
                        "chain_id": "C10",
                        "name": "SEC Investigation Chain",
                        "peril_id": "P5",
                        "description": "SEC opens formal investigation into disclosures",
                        "active": True,
                        "risk_level": "MODERATE",
                        "triggered_triggers": ["GOV.ENFORCE.sec_investigation"],
                        "active_amplifiers": [],
                        "active_mitigators": [],
                        "total_triggers": 2,
                        "total_amplifiers": 1,
                        "total_mitigators": 1,
                        "evidence_summary": ["SEC investigation disclosed in 10-K"],
                        "evidence_map": {},
                        "historical_filing_rate": None,
                        "median_severity_usd": None,
                    },
                ],
            },
        ],
        "all_perils": [
            {
                "peril_id": "P1",
                "name": "Securities Class Action",
                "risk_level": "HIGH",
                "active_chain_count": 2,
                "total_chain_count": 3,
                "key_evidence": ["Stock drop >15%", "Revenue restatement"],
            },
            {
                "peril_id": "P2",
                "name": "Derivative Suit",
                "risk_level": "LOW",
                "active_chain_count": 0,
                "total_chain_count": 2,
                "key_evidence": [],
            },
            {
                "peril_id": "P3",
                "name": "ERISA Litigation",
                "risk_level": "LOW",
                "active_chain_count": 0,
                "total_chain_count": 1,
                "key_evidence": [],
            },
            {
                "peril_id": "P4",
                "name": "Bankruptcy Claims",
                "risk_level": "LOW",
                "active_chain_count": 0,
                "total_chain_count": 2,
                "key_evidence": [],
            },
            {
                "peril_id": "P5",
                "name": "Regulatory Enforcement",
                "risk_level": "MODERATE",
                "active_chain_count": 1,
                "total_chain_count": 2,
                "key_evidence": ["Active SEC investigation"],
            },
            {
                "peril_id": "P6",
                "name": "Employment Practices",
                "risk_level": "LOW",
                "active_chain_count": 0,
                "total_chain_count": 2,
                "key_evidence": [],
            },
            {
                "peril_id": "P7",
                "name": "Antitrust",
                "risk_level": "LOW",
                "active_chain_count": 0,
                "total_chain_count": 1,
                "key_evidence": [],
            },
            {
                "peril_id": "P8",
                "name": "Cyber/Privacy",
                "risk_level": "LOW",
                "active_chain_count": 0,
                "total_chain_count": 2,
                "key_evidence": [],
            },
        ],
        "active_count": 2,
        "highest_peril": "P1",
    }


def _make_empty_peril_data() -> dict[str, Any]:
    """Create peril data with no active perils."""
    return {
        "perils": [],
        "all_perils": [
            {
                "peril_id": "P1",
                "name": "Securities Class Action",
                "risk_level": "LOW",
                "active_chain_count": 0,
                "total_chain_count": 3,
                "key_evidence": [],
            },
        ],
        "active_count": 0,
        "highest_peril": None,
    }


# -----------------------------------------------------------------------
# Tests: Peril Summary Table
# -----------------------------------------------------------------------


class TestPerilSummary:
    """Test the peril summary table renderer."""

    def test_renders_heading(self) -> None:
        """Summary should have D&O Claim Peril Assessment heading."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_summary(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "D&O Claim Peril Assessment" in text

    def test_renders_active_count(self) -> None:
        """Should show count of active perils."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_summary(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "2 of 8" in text

    def test_table_has_correct_row_count(self) -> None:
        """Table should have 8 data rows + 1 header = 9 rows."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_summary(doc, peril_data, ds)

        tables = doc.tables
        assert len(tables) >= 1
        summary_table = tables[0]
        assert len(summary_table.rows) == 9  # 1 header + 8 perils

    def test_table_has_four_columns(self) -> None:
        """Table should have 4 columns: Peril, Risk Level, Active Chains, Key Evidence."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_summary(doc, peril_data, ds)

        summary_table = doc.tables[0]
        assert len(summary_table.rows[0].cells) == 4

    def test_contains_all_peril_names(self) -> None:
        """All 8 peril names should appear in the table."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_summary(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Securities Class Action" in text
        assert "Derivative Suit" in text
        assert "ERISA Litigation" in text
        assert "Bankruptcy Claims" in text
        assert "Regulatory Enforcement" in text
        assert "Employment Practices" in text
        assert "Antitrust" in text
        assert "Cyber/Privacy" in text

    def test_contains_risk_levels(self) -> None:
        """Risk levels should appear in the table."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_summary(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "HIGH" in text
        assert "MODERATE" in text
        assert "LOW" in text

    def test_contains_chain_counts(self) -> None:
        """Active chain counts should appear."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_summary(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "2/3" in text  # SCA: 2 of 3 active
        assert "1/2" in text  # Regulatory: 1 of 2 active

    def test_contains_key_evidence(self) -> None:
        """Key evidence snippets should appear."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_summary(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Stock drop >15%" in text
        assert "Active SEC investigation" in text

    def test_empty_all_perils_shows_not_available(self) -> None:
        """Empty all_perils should show not available message."""
        doc = _make_doc()
        ds = _make_ds()

        render_peril_summary(doc, {"all_perils": []}, ds)

        text = _get_all_text(doc)
        assert "Peril assessment not available" in text

    def test_long_evidence_not_truncated(self) -> None:
        """Evidence longer than 60 chars should NOT be truncated — full text preserved."""
        doc = _make_doc()
        ds = _make_ds()
        long_evidence = "This is a very long evidence string that should definitely not be truncated because analytical content must be preserved in full"
        peril_data = {
            "all_perils": [
                {
                    "peril_id": "P1",
                    "name": "Test Peril",
                    "risk_level": "HIGH",
                    "active_chain_count": 1,
                    "total_chain_count": 1,
                    "key_evidence": [long_evidence],
                },
            ],
            "active_count": 1,
        }

        render_peril_summary(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert long_evidence in text


# -----------------------------------------------------------------------
# Tests: Peril Deep Dives
# -----------------------------------------------------------------------


class TestPerilDeepDives:
    """Test per-peril deep dive rendering."""

    def test_renders_active_peril_analysis_heading(self) -> None:
        """Should show Active Peril Analysis heading."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_deep_dives(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Active Peril Analysis" in text

    def test_renders_peril_names_with_risk(self) -> None:
        """Each active peril should show name and risk level."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_deep_dives(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Securities Class Action: HIGH" in text
        assert "Regulatory Enforcement: MODERATE" in text

    def test_renders_frequency_severity(self) -> None:
        """Frequency and severity context should appear."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_deep_dives(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Frequency: common" in text
        assert "Severity: high" in text
        assert "Typical range: $5M-$500M" in text

    def test_renders_chain_names(self) -> None:
        """Active chain names should appear."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_deep_dives(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Chain: Disclosure Fraud Chain" in text
        assert "Chain: Earnings Guidance Chain" in text
        assert "Chain: SEC Investigation Chain" in text

    def test_inactive_chains_not_rendered(self) -> None:
        """Inactive chains should not get deep dive rendering."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_deep_dives(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Chain: Insider Trading Chain" not in text

    def test_renders_triggers(self) -> None:
        """Triggered checks should appear under Triggers."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_deep_dives(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Triggers:" in text
        assert "FIN.ACCT.restatement" in text
        assert "MKT.STOCK.drop" in text

    def test_renders_trigger_evidence(self) -> None:
        """Trigger evidence from evidence_map should appear."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_deep_dives(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Q3 10-Q/A revenue restatement" in text
        assert "18% decline on disclosure day" in text

    def test_renders_amplifiers(self) -> None:
        """Active amplifiers should appear."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_deep_dives(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Amplifiers:" in text
        assert "GOV.INSIDER.trading" in text

    def test_renders_mitigators(self) -> None:
        """Active mitigators should appear."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_deep_dives(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Mitigators:" in text
        assert "GOV.BOARD.independence" in text

    def test_renders_historical_context(self) -> None:
        """Historical filing rate and median severity should appear."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_deep_dives(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Historical filing rate: 15%" in text
        assert "Median severity: $25M" in text

    def test_renders_chain_description(self) -> None:
        """Chain description should appear in italic."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_deep_dives(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Material misstatement leads to corrective disclosure" in text

    def test_active_chain_count_summary(self) -> None:
        """Should show active chain count per peril."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_peril_data()

        render_peril_deep_dives(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "2 of 3 causal chains active" in text
        # Regulatory has 1 chain in perils (only active shown with its total)

    def test_no_active_perils_renders_nothing(self) -> None:
        """No active perils should produce no output."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = _make_empty_peril_data()

        render_peril_deep_dives(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Active Peril Analysis" not in text


# -----------------------------------------------------------------------
# Tests: Edge cases
# -----------------------------------------------------------------------


class TestEdgeCases:
    """Test edge cases and graceful fallbacks."""

    def test_empty_peril_data_dict(self) -> None:
        """Empty dict should show not available."""
        doc = _make_doc()
        ds = _make_ds()

        render_peril_summary(doc, {}, ds)

        text = _get_all_text(doc)
        assert "Peril assessment not available" in text

    def test_peril_without_chains(self) -> None:
        """Peril with no chains should still render in summary."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = {
            "all_perils": [
                {
                    "peril_id": "P1",
                    "name": "Test Peril",
                    "risk_level": "LOW",
                    "active_chain_count": 0,
                    "total_chain_count": 0,
                    "key_evidence": [],
                },
            ],
            "perils": [],
            "active_count": 0,
        }

        render_peril_summary(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Test Peril" in text
        assert "0/0" in text

    def test_chain_without_historical_data(self) -> None:
        """Chain with None filing_rate/severity should not render historical."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = {
            "perils": [
                {
                    "peril_id": "P1",
                    "name": "Test",
                    "risk_level": "MODERATE",
                    "active_chain_count": 1,
                    "total_chain_count": 1,
                    "frequency": "rare",
                    "severity": "low",
                    "typical_settlement_range": "",
                    "chains": [
                        {
                            "chain_id": "C1",
                            "name": "Test Chain",
                            "active": True,
                            "risk_level": "MODERATE",
                            "description": "",
                            "triggered_triggers": ["CHECK.1"],
                            "active_amplifiers": [],
                            "active_mitigators": [],
                            "evidence_map": {},
                            "historical_filing_rate": None,
                            "median_severity_usd": None,
                        },
                    ],
                },
            ],
            "all_perils": [],
            "active_count": 1,
        }

        render_peril_deep_dives(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "Historical filing rate" not in text
        assert "Median severity" not in text

    def test_evidence_em_dash_for_no_evidence(self) -> None:
        """Peril with no key evidence should show em-dash in table."""
        doc = _make_doc()
        ds = _make_ds()
        peril_data = {
            "all_perils": [
                {
                    "peril_id": "P1",
                    "name": "Clean Peril",
                    "risk_level": "LOW",
                    "active_chain_count": 0,
                    "total_chain_count": 1,
                    "key_evidence": [],
                },
            ],
            "active_count": 0,
        }

        render_peril_summary(doc, peril_data, ds)

        text = _get_all_text(doc)
        assert "\u2014" in text  # em-dash
