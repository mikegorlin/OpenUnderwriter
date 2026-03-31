"""Tests for Section 7 peril map renderer.

Verifies heat map grid, bear case rendering, settlement prediction,
tower characterization, and graceful handling of missing data.
"""

from __future__ import annotations

from typing import Any

from docx import Document  # type: ignore[import-untyped]

from do_uw.models.peril import (
    BearCase,
    EvidenceItem,
    PerilMap,
    PlaintiffAssessment,
)
from do_uw.models.state import AnalysisResults, AnalysisState
from do_uw.stages.render.design_system import DesignSystem, setup_styles
from do_uw.stages.render.sections.sect7_peril_map import render_peril_map


def _make_context(state: AnalysisState) -> dict[str, Any]:
    """Wrap AnalysisState in context dict for section renderers.

    Phase 60-02: Sections now receive context dict from build_template_context().
    Tests use this minimal wrapper with _state escape hatch.
    """
    return {"_state": state, "company_name": state.ticker}


def _make_ds() -> DesignSystem:
    return DesignSystem()


def _make_doc() -> Any:
    doc = Document()
    setup_styles(doc)
    return doc


def _make_assessments() -> list[PlaintiffAssessment]:
    """Create 7 plaintiff assessments for testing."""
    lenses = [
        ("SHAREHOLDERS", "HIGH", "SEVERE", 5, 20, 18, ["Stock drop >15%"]),
        ("REGULATORS", "MODERATE", "SIGNIFICANT", 3, 15, 12, ["SEC investigation"]),
        ("CUSTOMERS", "LOW", "MINOR", 1, 10, 8, ["Product recall"]),
        ("EMPLOYEES", "VERY_LOW", "NUISANCE", 0, 8, 6, []),
        ("CREDITORS", "VERY_LOW", "NUISANCE", 0, 5, 4, []),
        ("COMPETITORS", "LOW", "MINOR", 1, 5, 4, ["Patent dispute"]),
        ("GOVERNMENT", "ELEVATED", "MODERATE", 2, 10, 9, ["Compliance gaps"]),
    ]
    return [
        PlaintiffAssessment(
            plaintiff_type=pt,
            probability_band=prob,
            severity_band=sev,
            triggered_signal_count=triggered,
            total_signal_count=total,
            evaluated_signal_count=evaluated,
            key_findings=findings,
            modeling_depth="FULL" if pt in ("SHAREHOLDERS", "REGULATORS") else "PROPORTIONAL",
        )
        for pt, prob, sev, triggered, total, evaluated, findings in lenses
    ]


def _make_bear_cases() -> list[BearCase]:
    """Create test bear cases."""
    return [
        BearCase(
            theory="A_DISCLOSURE",
            plaintiff_type="SHAREHOLDERS",
            committee_summary=(
                "Company failed to disclose material revenue recognition "
                "changes in Q3 10-Q. Stock dropped 18% on restatement. "
                "Class period spans 6 months."
            ),
            evidence_chain=[
                EvidenceItem(
                    signal_id="FIN.ACCT.restatement",
                    description="Revenue restatement in Q3 10-Q/A",
                    source="SEC EDGAR 10-Q/A filing",
                    severity="CRITICAL",
                    data_status="EVALUATED",
                ),
                EvidenceItem(
                    signal_id="MKT.STOCK.drop",
                    description="Stock declined 18% on disclosure day",
                    source="Yahoo Finance",
                    severity="HIGH",
                    data_status="EVALUATED",
                ),
            ],
            severity_estimate="SEVERE",
            defense_assessment="Strong scienter argument given CFO departure timing",
            probability_band="HIGH",
            supporting_signal_count=5,
        ),
        BearCase(
            theory="C_REGULATORY",
            plaintiff_type="REGULATORS",
            committee_summary=(
                "SEC investigation into accounting practices opened. "
                "Potential enforcement action for internal controls failure."
            ),
            evidence_chain=[
                EvidenceItem(
                    signal_id="GOV.ENFORCE.sec_investigation",
                    description="Active SEC investigation disclosed",
                    source="10-K Item 3",
                    severity="HIGH",
                    data_status="EVALUATED",
                ),
            ],
            severity_estimate="SIGNIFICANT",
            defense_assessment=None,
            probability_band="MODERATE",
            supporting_signal_count=3,
        ),
    ]


def _make_peril_map() -> PerilMap:
    """Create a full peril map for testing."""
    return PerilMap(
        assessments=_make_assessments(),
        bear_cases=_make_bear_cases(),
        overall_peril_rating="HIGH",
        coverage_gaps=["FIN.CASH.burn_rate: Missing quarterly cash flow data"],
    )


def _make_state_with_peril(
    peril_map: PerilMap | None = None,
    settlement_pred: dict[str, Any] | None = None,
) -> AnalysisState:
    """Create AnalysisState with peril map data."""
    pm_dict = peril_map.model_dump() if peril_map else None
    analysis = AnalysisResults(
        peril_map=pm_dict,
        settlement_prediction=settlement_pred,
    )
    return AnalysisState(ticker="TEST", analysis=analysis)


def _get_all_text(doc: Any) -> str:
    """Extract all text from a document."""
    texts: list[str] = []
    for para in doc.paragraphs:
        texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)


# -----------------------------------------------------------------------
# Tests: render_peril_map with full data
# -----------------------------------------------------------------------


class TestRenderPerilMap:
    """Test the main render_peril_map entry point."""

    def test_renders_with_full_peril_map(self) -> None:
        """Full peril map should produce heat map, bear cases, etc."""
        doc = _make_doc()
        ds = _make_ds()
        pm = _make_peril_map()
        state = _make_state_with_peril(pm)

        render_peril_map(doc, _make_context(state), ds)

        text = _get_all_text(doc)
        assert "Peril Assessment" in text
        assert "Plaintiff Peril Heat Map" in text
        assert "Bear Cases" in text

    def test_renders_placeholder_when_no_peril_map(self) -> None:
        """Missing peril map should show placeholder."""
        doc = _make_doc()
        ds = _make_ds()
        state = _make_state_with_peril(None)

        render_peril_map(doc, _make_context(state), ds)

        text = _get_all_text(doc)
        assert "Peril assessment not available" in text


# -----------------------------------------------------------------------
# Tests: Heat map table
# -----------------------------------------------------------------------


class TestPlaintiffHeatMap:
    """Test the 7x2 heat map grid."""

    def test_heat_map_has_correct_row_count(self) -> None:
        """Table should have 7 data rows + 1 header = 8 rows."""
        doc = _make_doc()
        ds = _make_ds()
        pm = _make_peril_map()
        state = _make_state_with_peril(pm)

        render_peril_map(doc, _make_context(state), ds)

        # Find the heat map table (first table in doc)
        tables = doc.tables
        assert len(tables) >= 1
        heat_map_table = tables[0]
        assert len(heat_map_table.rows) == 8  # 1 header + 7 data

    def test_heat_map_has_three_columns(self) -> None:
        """Table should have 3 columns: Type, Probability, Severity."""
        doc = _make_doc()
        ds = _make_ds()
        pm = _make_peril_map()
        state = _make_state_with_peril(pm)

        render_peril_map(doc, _make_context(state), ds)

        heat_map_table = doc.tables[0]
        assert len(heat_map_table.rows[0].cells) == 3

    def test_heat_map_contains_plaintiff_types(self) -> None:
        """All 7 plaintiff types should appear in the table."""
        doc = _make_doc()
        ds = _make_ds()
        pm = _make_peril_map()
        state = _make_state_with_peril(pm)

        render_peril_map(doc, _make_context(state), ds)

        text = _get_all_text(doc)
        assert "Shareholders" in text
        assert "Regulators" in text
        assert "Customers" in text
        assert "Employees" in text
        assert "Creditors" in text
        assert "Competitors" in text
        assert "Government" in text

    def test_elevated_lenses_get_notes(self) -> None:
        """Lenses with probability >= MODERATE get summary notes below table."""
        doc = _make_doc()
        ds = _make_ds()
        pm = _make_peril_map()
        state = _make_state_with_peril(pm)

        render_peril_map(doc, _make_context(state), ds)

        text = _get_all_text(doc)
        # Shareholders is HIGH, Regulators is MODERATE, Government is ELEVATED
        assert "Shareholders: 5 triggered checks" in text
        assert "Regulators: 3 triggered checks" in text
        assert "Government: 2 triggered checks" in text


# -----------------------------------------------------------------------
# Tests: Bear cases
# -----------------------------------------------------------------------


class TestBearCases:
    """Test bear case rendering."""

    def test_bear_cases_include_committee_summary(self) -> None:
        """Each bear case should show the committee summary."""
        doc = _make_doc()
        ds = _make_ds()
        pm = _make_peril_map()
        state = _make_state_with_peril(pm)

        render_peril_map(doc, _make_context(state), ds)

        text = _get_all_text(doc)
        assert "failed to disclose material revenue recognition" in text
        assert "SEC investigation into accounting practices" in text

    def test_bear_cases_include_evidence_chain(self) -> None:
        """Bear cases should render numbered evidence items."""
        doc = _make_doc()
        ds = _make_ds()
        pm = _make_peril_map()
        state = _make_state_with_peril(pm)

        render_peril_map(doc, _make_context(state), ds)

        text = _get_all_text(doc)
        assert "Evidence Chain" in text
        assert "FIN.ACCT.restatement" in text
        assert "MKT.STOCK.drop" in text

    def test_bear_cases_include_defense_assessment(self) -> None:
        """Bear cases with defense assessment should display it."""
        doc = _make_doc()
        ds = _make_ds()
        pm = _make_peril_map()
        state = _make_state_with_peril(pm)

        render_peril_map(doc, _make_context(state), ds)

        text = _get_all_text(doc)
        assert "Strong scienter argument" in text

    def test_empty_bear_cases_show_clean_message(self) -> None:
        """No bear cases should produce the 'no bear cases' message."""
        doc = _make_doc()
        ds = _make_ds()
        pm = PerilMap(
            assessments=_make_assessments(),
            bear_cases=[],
            overall_peril_rating="LOW",
        )
        state = _make_state_with_peril(pm)

        render_peril_map(doc, _make_context(state), ds)

        text = _get_all_text(doc)
        assert "No bear cases identified" in text

    def test_bear_case_footer_has_probability_severity(self) -> None:
        """Each bear case should end with probability and severity."""
        doc = _make_doc()
        ds = _make_ds()
        pm = _make_peril_map()
        state = _make_state_with_peril(pm)

        render_peril_map(doc, _make_context(state), ds)

        text = _get_all_text(doc)
        assert "Probability: HIGH" in text
        assert "Severity: SEVERE" in text


# -----------------------------------------------------------------------
# Tests: Settlement summary
# -----------------------------------------------------------------------


class TestSettlementSummary:
    """Test settlement prediction rendering."""

    def test_settlement_with_ddl_prediction(self) -> None:
        """Phase 27 settlement prediction should render DDL table."""
        doc = _make_doc()
        ds = _make_ds()
        pm = _make_peril_map()
        settlement = {
            "market_cap": 5_000_000_000,
            "scenarios": [
                {
                    "percentile": 25,
                    "label": "Favorable",
                    "ddl_amount": 50_000_000,
                    "settlement_estimate": 15_000_000,
                    "defense_cost_estimate": 5_000_000,
                    "total_exposure": 20_000_000,
                },
                {
                    "percentile": 50,
                    "label": "Median",
                    "ddl_amount": 200_000_000,
                    "settlement_estimate": 60_000_000,
                    "defense_cost_estimate": 15_000_000,
                    "total_exposure": 75_000_000,
                },
                {
                    "percentile": 75,
                    "label": "Adverse",
                    "ddl_amount": 500_000_000,
                    "settlement_estimate": 150_000_000,
                    "defense_cost_estimate": 30_000_000,
                    "total_exposure": 180_000_000,
                },
                {
                    "percentile": 95,
                    "label": "Catastrophic",
                    "ddl_amount": 1_500_000_000,
                    "settlement_estimate": 450_000_000,
                    "defense_cost_estimate": 60_000_000,
                    "total_exposure": 510_000_000,
                },
            ],
        }
        state = _make_state_with_peril(pm, settlement)

        render_peril_map(doc, _make_context(state), ds)

        text = _get_all_text(doc)
        assert "Settlement Prediction" in text
        assert "Favorable" in text
        assert "Catastrophic" in text

    def test_settlement_table_has_4_scenarios(self) -> None:
        """Settlement table should have 4 data rows + 1 header."""
        doc = _make_doc()
        ds = _make_ds()
        pm = _make_peril_map()
        settlement = {
            "scenarios": [
                {"percentile": 25, "label": "Favorable", "ddl_amount": 50e6,
                 "settlement_estimate": 15e6, "defense_cost_estimate": 5e6,
                 "total_exposure": 20e6},
                {"percentile": 50, "label": "Median", "ddl_amount": 200e6,
                 "settlement_estimate": 60e6, "defense_cost_estimate": 15e6,
                 "total_exposure": 75e6},
                {"percentile": 75, "label": "Adverse", "ddl_amount": 500e6,
                 "settlement_estimate": 150e6, "defense_cost_estimate": 30e6,
                 "total_exposure": 180e6},
                {"percentile": 95, "label": "Catastrophic", "ddl_amount": 1.5e9,
                 "settlement_estimate": 450e6, "defense_cost_estimate": 60e6,
                 "total_exposure": 510e6},
            ]
        }
        state = _make_state_with_peril(pm, settlement)

        render_peril_map(doc, _make_context(state), ds)

        # Find settlement table -- it's the second table (after heat map)
        settlement_table = doc.tables[1]
        assert len(settlement_table.rows) == 5  # 1 header + 4 scenarios

    def test_settlement_not_available(self) -> None:
        """No settlement data should show placeholder."""
        doc = _make_doc()
        ds = _make_ds()
        pm = _make_peril_map()
        state = _make_state_with_peril(pm, None)

        render_peril_map(doc, _make_context(state), ds)

        text = _get_all_text(doc)
        assert "Settlement prediction not available" in text


# -----------------------------------------------------------------------
# Tests: Tower characterization
# -----------------------------------------------------------------------


class TestTowerCharacterization:
    """Test tower risk characterization rendering."""

    def test_tower_renders_layers(self) -> None:
        """Tower characterization should render per-layer risk table."""
        doc = _make_doc()
        ds = _make_ds()
        pm = _make_peril_map()
        settlement = {
            "scenarios": [],
            "tower_risk_characterization": {
                "layers": [
                    {
                        "layer_type": "primary",
                        "expected_loss_share_pct": 65,
                        "description": "Primary layer carries majority of expected loss",
                    },
                    {
                        "layer_type": "first_excess",
                        "expected_loss_share_pct": 25,
                        "description": "First excess has moderate exposure",
                    },
                    {
                        "layer_type": "second_excess",
                        "expected_loss_share_pct": 10,
                        "description": "Upper excess has tail risk only",
                    },
                ],
            },
        }
        state = _make_state_with_peril(pm, settlement)

        render_peril_map(doc, _make_context(state), ds)

        text = _get_all_text(doc)
        assert "Tower Risk Characterization" in text
        assert "Primary" in text
        assert "65%" in text
        assert "Primary layer carries 65% of expected loss exposure" in text

    def test_tower_not_available(self) -> None:
        """Missing tower data should show placeholder."""
        doc = _make_doc()
        ds = _make_ds()
        pm = _make_peril_map()
        state = _make_state_with_peril(pm, None)

        render_peril_map(doc, _make_context(state), ds)

        text = _get_all_text(doc)
        assert "Tower characterization not available" in text


# -----------------------------------------------------------------------
# Tests: Graceful handling
# -----------------------------------------------------------------------


class TestGracefulHandling:
    """Test edge cases and missing data handling."""

    def test_none_analysis(self) -> None:
        """State with None analysis should show placeholder."""
        doc = _make_doc()
        ds = _make_ds()
        state = AnalysisState(ticker="TEST", analysis=None)

        render_peril_map(doc, _make_context(state), ds)

        text = _get_all_text(doc)
        assert "Peril assessment not available" in text

    def test_empty_peril_map_dict(self) -> None:
        """Empty peril map dict should still render (with defaults)."""
        doc = _make_doc()
        ds = _make_ds()
        analysis = AnalysisResults(peril_map={})
        state = AnalysisState(ticker="TEST", analysis=analysis)

        render_peril_map(doc, _make_context(state), ds)

        text = _get_all_text(doc)
        # PerilMap with defaults has empty assessments
        assert "Peril Assessment" in text
