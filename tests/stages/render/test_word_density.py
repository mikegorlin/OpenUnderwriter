"""Tests for Word renderer density-tier indicators and pre-computed narratives.

Phase 35-06: Verifies that the Word renderer correctly renders density
indicators (CLEAN/ELEVATED/CRITICAL) and pre-computed narrative paragraphs.
"""

from __future__ import annotations

from datetime import UTC, datetime
from typing import Any
from unittest.mock import MagicMock

import pytest
from docx import Document  # type: ignore[import-untyped]

from do_uw.models.density import (
    DensityLevel,
    PreComputedNarratives,
    SectionDensity,
)
from do_uw.models.state import AnalysisResults, AnalysisState
from do_uw.stages.render.design_system import setup_styles
from do_uw.stages.render.word_renderer import (
    _add_density_indicator,
    _add_section_narrative,
)


def _make_state(
    densities: dict[str, SectionDensity] | None = None,
    narratives: PreComputedNarratives | None = None,
) -> AnalysisState:
    """Build a minimal AnalysisState with density/narrative data."""
    state = AnalysisState(ticker="TEST")
    state.analysis = AnalysisResults()
    if densities is not None:
        state.analysis.section_densities = densities
    if narratives is not None:
        state.analysis.pre_computed_narratives = narratives
    return state


def _make_doc() -> Any:
    """Create a Document with custom styles for testing."""
    doc: Any = Document()
    setup_styles(doc)
    return doc


def _get_paragraph_texts(doc: Any) -> list[str]:
    """Extract all paragraph text from a document."""
    return [p.text for p in doc.paragraphs]


class TestWordDensityIndicatorClean:
    """CLEAN density should add NO indicator paragraph."""

    def test_clean_no_indicator(self) -> None:
        doc = _make_doc()
        state = _make_state(
            densities={"financial": SectionDensity(level=DensityLevel.CLEAN)}
        )
        initial_count = len(doc.paragraphs)
        _add_density_indicator(doc, "Section 3: Financial Health", state)
        assert len(doc.paragraphs) == initial_count

    def test_clean_no_indicator_all_sections(self) -> None:
        """All clean sections should produce zero indicator paragraphs."""
        doc = _make_doc()
        state = _make_state(
            densities={
                "governance": SectionDensity(level=DensityLevel.CLEAN),
                "financial": SectionDensity(level=DensityLevel.CLEAN),
                "market": SectionDensity(level=DensityLevel.CLEAN),
                "litigation": SectionDensity(level=DensityLevel.CLEAN),
            }
        )
        initial_count = len(doc.paragraphs)
        for name in [
            "Section 3: Financial Health",
            "Section 4: Market & Trading",
            "Section 5: Governance & Leadership",
            "Section 6: Litigation & Regulatory",
        ]:
            _add_density_indicator(doc, name, state)
        assert len(doc.paragraphs) == initial_count


class TestWordDensityIndicatorElevated:
    """ELEVATED density should add an amber 'ELEVATED CONCERN' paragraph."""

    def test_elevated_adds_paragraph(self) -> None:
        doc = _make_doc()
        state = _make_state(
            densities={"financial": SectionDensity(level=DensityLevel.ELEVATED)}
        )
        initial_count = len(doc.paragraphs)
        _add_density_indicator(doc, "Section 3: Financial Health", state)
        assert len(doc.paragraphs) == initial_count + 1

    def test_elevated_text_content(self) -> None:
        doc = _make_doc()
        state = _make_state(
            densities={"financial": SectionDensity(level=DensityLevel.ELEVATED)}
        )
        _add_density_indicator(doc, "Section 3: Financial Health", state)
        texts = _get_paragraph_texts(doc)
        assert "ELEVATED CONCERN" in texts[-1]

    def test_elevated_amber_color(self) -> None:
        doc = _make_doc()
        state = _make_state(
            densities={"governance": SectionDensity(level=DensityLevel.ELEVATED)}
        )
        _add_density_indicator(doc, "Section 5: Governance & Leadership", state)
        last_para = doc.paragraphs[-1]
        run = last_para.runs[0]
        # Amber color: RGB(255, 184, 0)
        assert run.font.color.rgb is not None
        assert run.bold is True


class TestWordDensityIndicatorCritical:
    """CRITICAL density should add a red 'CRITICAL RISK' paragraph."""

    def test_critical_adds_paragraph(self) -> None:
        doc = _make_doc()
        state = _make_state(
            densities={"litigation": SectionDensity(level=DensityLevel.CRITICAL)}
        )
        initial_count = len(doc.paragraphs)
        _add_density_indicator(doc, "Section 6: Litigation & Regulatory", state)
        assert len(doc.paragraphs) == initial_count + 1

    def test_critical_text_content(self) -> None:
        doc = _make_doc()
        state = _make_state(
            densities={"litigation": SectionDensity(level=DensityLevel.CRITICAL)}
        )
        _add_density_indicator(doc, "Section 6: Litigation & Regulatory", state)
        texts = _get_paragraph_texts(doc)
        assert "CRITICAL RISK" in texts[-1]

    def test_critical_red_color(self) -> None:
        doc = _make_doc()
        state = _make_state(
            densities={"market": SectionDensity(level=DensityLevel.CRITICAL)}
        )
        _add_density_indicator(doc, "Section 4: Market & Trading", state)
        last_para = doc.paragraphs[-1]
        run = last_para.runs[0]
        assert run.font.color.rgb is not None
        assert run.bold is True


class TestWordNarrativeFromPrecomputed:
    """Pre-computed narratives should render with AI Assessment label."""

    def test_narrative_adds_paragraph(self) -> None:
        doc = _make_doc()
        state = _make_state(
            narratives=PreComputedNarratives(
                financial="Strong balance sheet with declining debt levels."
            )
        )
        initial_count = len(doc.paragraphs)
        _add_section_narrative(doc, "Section 3: Financial Health", state)
        assert len(doc.paragraphs) == initial_count + 1

    def test_narrative_contains_content_without_ai_label(self) -> None:
        """Narrative renders content directly without AI Assessment prefix."""
        doc = _make_doc()
        state = _make_state(
            narratives=PreComputedNarratives(
                governance="Board is well-structured with majority independence."
            )
        )
        _add_section_narrative(doc, "Section 5: Governance & Leadership", state)
        text = doc.paragraphs[-1].text
        assert "[AI Assessment]" not in text
        assert "Board is well-structured" in text

    def test_narrative_body_is_first_run(self) -> None:
        """Narrative body is rendered as the first (and only) run, not prefixed."""
        doc = _make_doc()
        state = _make_state(
            narratives=PreComputedNarratives(
                market="Moderate volatility with limited short interest."
            )
        )
        _add_section_narrative(doc, "Section 4: Market & Trading", state)
        last_para = doc.paragraphs[-1]
        body_run = last_para.runs[0]
        assert "Moderate volatility" in body_run.text

    def test_narrative_body_not_italic(self) -> None:
        doc = _make_doc()
        state = _make_state(
            narratives=PreComputedNarratives(
                litigation="No active securities class actions."
            )
        )
        _add_section_narrative(doc, "Section 6: Litigation & Regulatory", state)
        last_para = doc.paragraphs[-1]
        # Body run is the first run (no prefix)
        body_run = last_para.runs[0]
        assert "No active securities" in body_run.text
        # Body should not be italic (default is None or False)
        assert body_run.italic is not True


class TestWordNarrativeFallback:
    """When pre_computed_narratives is None, no paragraph should be added."""

    def test_no_narratives_no_paragraph(self) -> None:
        doc = _make_doc()
        state = _make_state(narratives=None)
        initial_count = len(doc.paragraphs)
        _add_section_narrative(doc, "Section 3: Financial Health", state)
        assert len(doc.paragraphs) == initial_count

    def test_empty_narrative_field_no_paragraph(self) -> None:
        """Narrative field is None for a specific section."""
        doc = _make_doc()
        state = _make_state(
            narratives=PreComputedNarratives(financial=None, governance="Has content")
        )
        initial_count = len(doc.paragraphs)
        _add_section_narrative(doc, "Section 3: Financial Health", state)
        assert len(doc.paragraphs) == initial_count

    def test_no_analysis_no_paragraph(self) -> None:
        """State has no analysis at all."""
        doc = _make_doc()
        state = AnalysisState(ticker="TEST")
        state.analysis = None
        initial_count = len(doc.paragraphs)
        _add_section_narrative(doc, "Section 3: Financial Health", state)
        assert len(doc.paragraphs) == initial_count


class TestWordDensityEdgeCases:
    """Edge cases: unknown section names, missing analysis, etc."""

    def test_unknown_section_name_no_indicator(self) -> None:
        doc = _make_doc()
        state = _make_state(
            densities={"financial": SectionDensity(level=DensityLevel.CRITICAL)}
        )
        initial_count = len(doc.paragraphs)
        _add_density_indicator(doc, "Calibration Notes", state)
        assert len(doc.paragraphs) == initial_count

    def test_no_analysis_no_indicator(self) -> None:
        doc = _make_doc()
        state = AnalysisState(ticker="TEST")
        state.analysis = None
        initial_count = len(doc.paragraphs)
        _add_density_indicator(doc, "Section 3: Financial Health", state)
        assert len(doc.paragraphs) == initial_count

    def test_missing_section_density_no_indicator(self) -> None:
        """Section exists in map but not in state.section_densities."""
        doc = _make_doc()
        state = _make_state(densities={})  # Empty densities
        initial_count = len(doc.paragraphs)
        _add_density_indicator(doc, "Section 3: Financial Health", state)
        assert len(doc.paragraphs) == initial_count

    def test_appendix_section_no_indicator(self) -> None:
        """Meeting Prep appendix should not get density indicators."""
        doc = _make_doc()
        state = _make_state(
            densities={"scoring": SectionDensity(level=DensityLevel.CRITICAL)}
        )
        initial_count = len(doc.paragraphs)
        _add_density_indicator(doc, "Appendix: Meeting Preparation", state)
        assert len(doc.paragraphs) == initial_count
