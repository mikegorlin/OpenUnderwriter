"""End-to-end render integration tests.

Verifies that RenderStage produces all three output formats (Word,
Markdown, PDF) from a comprehensive test fixture state with density
indicators and pre-computed narratives.
"""

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from do_uw.models.common import Confidence, SourcedValue, StageResult, StageStatus
from do_uw.models.company import CompanyIdentity, CompanyProfile
from do_uw.models.density import (
    DensityLevel,
    PreComputedNarratives,
    SectionDensity,
)
from do_uw.models.financials import (
    AuditProfile,
    DistressIndicators,
    DistressResult,
    DistressZone,
    ExtractedFinancials,
)
from do_uw.models.governance import BoardProfile, GovernanceData
from do_uw.models.governance_forensics import (
    LeadershipStability,
    OwnershipAnalysis,
)
from do_uw.models.litigation import LitigationLandscape
from do_uw.models.market import MarketSignals
from do_uw.models.scoring import ScoringResult
from do_uw.models.state import (
    AnalysisResults,
    AnalysisState,
    ExtractedData,
)

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


_NOW = datetime.now(tz=UTC)


def _sv(val: object, confidence: str = "HIGH") -> SourcedValue:  # type: ignore[type-arg]
    """Shorthand for creating SourcedValue."""
    return SourcedValue(
        value=val,
        source="test",
        confidence=Confidence(confidence),
        as_of=_NOW,
    )


def _build_fixture_state() -> AnalysisState:
    """Build a comprehensive fixture state for render integration tests.

    Includes:
    - Company profile with identity
    - Extracted data (financials, governance, litigation, market)
    - Analysis results with section densities and check results
    - Scoring results
    - Pre-computed narratives
    - Stage completion markers for all pre-render stages
    """
    state = AnalysisState(ticker="TEST")

    # Company profile
    state.company = CompanyProfile(
        identity=CompanyIdentity(
            ticker="TEST",
            legal_name=_sv("Test Company Inc."),
            cik=_sv("0001234567"),
            sic_code=_sv("3674"),
            sic_description=_sv("Semiconductors"),
        ),
    )

    # Extracted data
    state.extracted = ExtractedData(
        financials=ExtractedFinancials(
            distress=DistressIndicators(
                altman_z_score=DistressResult(score=3.5, zone=DistressZone.SAFE),
                beneish_m_score=DistressResult(score=-2.5, zone=DistressZone.SAFE),
            ),
            audit=AuditProfile(
                auditor_name=_sv("Deloitte & Touche LLP"),
                going_concern=_sv(False),
            ),
        ),
        governance=GovernanceData(
            board=BoardProfile(
                independence_ratio=_sv(0.85),
                ceo_chair_duality=_sv(False),
            ),
            ownership=OwnershipAnalysis(),
            leadership=LeadershipStability(),
        ),
        litigation=LitigationLandscape(),
        market=MarketSignals(),
    )

    # Analysis results
    state.analysis = AnalysisResults()
    state.analysis.section_densities = {
        "executive": SectionDensity(level=DensityLevel.CLEAN),
        "company": SectionDensity(level=DensityLevel.CLEAN),
        "financial": SectionDensity(level=DensityLevel.CRITICAL,
                                    concerns=["Distress models not computed"],
                                    critical_evidence=["Going concern opinion"]),
        "market": SectionDensity(level=DensityLevel.ELEVATED,
                                  concerns=["Short interest elevated"]),
        "governance": SectionDensity(level=DensityLevel.CLEAN),
        "litigation": SectionDensity(level=DensityLevel.CLEAN),
    }
    state.analysis.pre_computed_narratives = PreComputedNarratives(
        executive_summary="Test Company exhibits stable financial profile.",
        company="Test Company operates in the semiconductor industry.",
        financial="Financial health is satisfactory with no distress signals.",
        governance="Governance structure meets best practice standards.",
        litigation="No active litigation matters identified.",
        scoring="Overall risk profile is low.",
    )
    state.analysis.signal_results = {
        "FIN.LIQ.position": {
            "signal_name": "Liquidity Position",
            "status": "CLEAR",
            "value": 2.1,
            "evidence": "Current ratio above 1.5",
            "content_type": "EVALUATIVE_CHECK",
            "data_status": "EVALUATED",
        },
        "GOV.BOARD.independence": {
            "signal_name": "Board Independence",
            "status": "CLEAR",
            "value": 0.85,
            "evidence": "85% independent directors",
            "content_type": "MANAGEMENT_DISPLAY",
            "data_status": "EVALUATED",
        },
    }

    # Scoring
    state.scoring = ScoringResult()

    # Mark all pre-render stages as completed
    for stage_name in ["resolve", "acquire", "extract", "analyze", "score", "benchmark"]:
        state.stages[stage_name] = StageResult(
            stage=stage_name,
            status=StageStatus.COMPLETED,
            started_at=_NOW,
            completed_at=_NOW,
        )

    return state


# ---------------------------------------------------------------------------
# Test: All formats render
# ---------------------------------------------------------------------------


class TestAllFormatsRender:
    """End-to-end render produces output files."""

    def test_all_formats_render(self, tmp_path: Path) -> None:
        """RenderStage.run() produces Word and Markdown files.

        PDF generation may fail if Playwright is not installed -- that
        is acceptable as it falls back gracefully.
        """
        from do_uw.stages.render import RenderStage

        state = _build_fixture_state()
        stage = RenderStage(output_dir=tmp_path)
        stage.run(state)

        # Word output must exist and be non-empty
        docx_path = tmp_path / "TEST_worksheet.docx"
        assert docx_path.exists(), "Word document not generated"
        assert docx_path.stat().st_size > 0, "Word document is empty"

        # Markdown output must exist and contain section headings
        md_path = tmp_path / "TEST_worksheet.md"
        assert md_path.exists(), "Markdown file not generated"
        md_content = md_path.read_text(encoding="utf-8")
        assert len(md_content) > 100, "Markdown file too small"
        # Should have at least some structure
        assert "#" in md_content, "Markdown file has no headings"

    def test_word_is_valid_docx(self, tmp_path: Path) -> None:
        """Generated Word file can be opened by python-docx."""
        from docx import Document  # type: ignore[import-untyped]

        from do_uw.stages.render import RenderStage

        state = _build_fixture_state()
        stage = RenderStage(output_dir=tmp_path)
        stage.run(state)

        docx_path = tmp_path / "TEST_worksheet.docx"
        doc: Any = Document(str(docx_path))
        # Should have paragraphs
        assert len(doc.paragraphs) > 0, "Word document has no paragraphs"


# ---------------------------------------------------------------------------
# Test: Word density indicators
# ---------------------------------------------------------------------------


class TestWordDensityIndicators:
    """Word output reflects density indicators."""

    def test_word_has_density_indicators(self, tmp_path: Path) -> None:
        """Render Word doc with CRITICAL financial section.

        The CRITICAL density should produce a visible indicator in the
        document text.
        """
        from docx import Document  # type: ignore[import-untyped]

        from do_uw.stages.render import RenderStage

        state = _build_fixture_state()
        # Ensure financial is CRITICAL
        assert state.analysis is not None
        fin_density = state.analysis.section_densities.get("financial")
        assert fin_density is not None
        assert fin_density.level == DensityLevel.CRITICAL

        stage = RenderStage(output_dir=tmp_path)
        stage.run(state)

        docx_path = tmp_path / "TEST_worksheet.docx"
        doc: Any = Document(str(docx_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)

        # CRITICAL RISK indicator should appear somewhere in the document
        assert "CRITICAL" in all_text.upper(), (
            "Expected CRITICAL risk indicator in Word document"
        )


# ---------------------------------------------------------------------------
# Test: Markdown narratives
# ---------------------------------------------------------------------------


class TestMarkdownNarratives:
    """Markdown output includes pre-computed narratives."""

    def test_markdown_has_narratives(self, tmp_path: Path) -> None:
        """Render Markdown from fixture with pre_computed_narratives.

        Pre-computed narrative text should appear in the output.
        """
        from do_uw.stages.render import RenderStage

        state = _build_fixture_state()
        stage = RenderStage(output_dir=tmp_path)
        stage.run(state)

        md_path = tmp_path / "TEST_worksheet.md"
        md_content = md_path.read_text(encoding="utf-8")

        # Check that at least one narrative appears in the output
        narratives = state.analysis.pre_computed_narratives  # type: ignore[union-attr]
        narrative_found = False
        for field_name in ["executive_summary", "financial", "governance", "litigation"]:
            narrative_text = getattr(narratives, field_name, None)
            if narrative_text and narrative_text in md_content:
                narrative_found = True
                break

        assert narrative_found, (
            "No pre-computed narrative text found in Markdown output"
        )


# ---------------------------------------------------------------------------
# Test: HTML context
# ---------------------------------------------------------------------------


class TestHTMLContext:
    """HTML context builder produces expected keys."""

    def test_html_context_has_density_and_narratives(self) -> None:
        """build_html_context returns densities and narratives keys."""
        from do_uw.stages.render.html_renderer import build_html_context

        state = _build_fixture_state()
        context = build_html_context(state)

        assert "densities" in context, "HTML context missing 'densities' key"
        assert "narratives" in context, "HTML context missing 'narratives' key"

    def test_html_context_density_levels(self) -> None:
        """Density levels in HTML context match state."""
        from do_uw.stages.render.html_renderer import build_html_context

        state = _build_fixture_state()
        context = build_html_context(state)

        densities = context["densities"]
        assert "financial" in densities, "Missing financial density"
        assert densities["financial"].level == DensityLevel.CRITICAL

    def test_html_context_narrative_content(self) -> None:
        """Narrative content in HTML context matches state."""
        from do_uw.stages.render.html_renderer import build_html_context

        state = _build_fixture_state()
        context = build_html_context(state)

        narratives = context["narratives"]
        assert narratives.financial == "Financial health is satisfactory with no distress signals."
