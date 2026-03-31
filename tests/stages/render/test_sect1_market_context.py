"""Tests for Section 1-08: Market Context renderer.

Verifies rendering with populated MarketIntelligence data,
graceful handling of missing data, and mispricing alert display.

Phase 60-01: Updated to pass context dict instead of raw state.
"""

from __future__ import annotations

from typing import Any

from docx import Document  # type: ignore[import-untyped]

from do_uw.models.executive_summary import (
    DealContext,
    ExecutiveSummary,
    MarketIntelligence,
)
from do_uw.models.state import AnalysisState
from do_uw.stages.render.design_system import DesignSystem, setup_styles
from do_uw.stages.render.sections.sect1_market_context import (
    render_market_context,
)


def _make_ds() -> DesignSystem:
    return DesignSystem()


def _make_doc() -> Any:
    doc = Document()
    setup_styles(doc)
    return doc


def _make_state_with_mi(
    mi: MarketIntelligence | None,
) -> AnalysisState:
    """Create an AnalysisState with the given MarketIntelligence."""
    deal = DealContext(market_intelligence=mi)
    es = ExecutiveSummary(deal_context=deal)
    return AnalysisState(ticker="TEST", executive_summary=es)


def _make_context(state: AnalysisState) -> dict[str, Any]:
    """Build a context dict from state for section renderers.

    Phase 60-01: Sections now receive context dict instead of raw state.
    """
    from do_uw.stages.render.md_renderer import build_template_context

    context = build_template_context(state)
    context["_state"] = state
    return context


def _extract_text(doc: Any) -> str:
    """Extract all paragraph and table text from a document."""
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


class TestMarketContextNoData:
    """Tests for rendering when pricing data is unavailable."""

    def test_has_data_false_shows_no_data_message(self) -> None:
        """When has_data is False, render 'No market pricing data available'."""
        doc = _make_doc()
        mi = MarketIntelligence(has_data=False)
        state = _make_state_with_mi(mi)
        context = _make_context(state)

        render_market_context(doc, context, _make_ds())

        text = _extract_text(doc)
        assert "Market Context" in text
        assert "No market pricing data available" in text

    def test_market_intelligence_none_shows_no_data(self) -> None:
        """When market_intelligence is None, render no-data message."""
        doc = _make_doc()
        state = _make_state_with_mi(None)
        context = _make_context(state)

        render_market_context(doc, context, _make_ds())

        text = _extract_text(doc)
        assert "Market Context" in text
        assert "No market pricing data available" in text

    def test_executive_summary_none_shows_no_data(self) -> None:
        """When executive_summary is None, render no-data gracefully."""
        doc = _make_doc()
        state = AnalysisState(ticker="TEST")
        context = _make_context(state)

        render_market_context(doc, context, _make_ds())

        text = _extract_text(doc)
        assert "Market Context" in text
        assert "No market pricing data available" in text

    def test_insufficient_data_with_peers_shows_count(self) -> None:
        """When has_data is False but peers exist, show peer count."""
        doc = _make_doc()
        mi = MarketIntelligence(
            has_data=False,
            peer_count=3,
            confidence_level="INSUFFICIENT",
        )
        state = _make_state_with_mi(mi)
        context = _make_context(state)

        render_market_context(doc, context, _make_ds())

        text = _extract_text(doc)
        assert "3 peers found" in text
        assert "INSUFFICIENT" in text


class TestMarketContextWithData:
    """Tests for rendering when pricing data is available."""

    def test_full_data_renders_table_and_summary(self) -> None:
        """When has_data is True with full fields, render table and summary."""
        doc = _make_doc()
        mi = MarketIntelligence(
            has_data=True,
            peer_count=15,
            confidence_level="HIGH",
            median_rate_on_line=0.0350,
            ci_low=0.0280,
            ci_high=0.0420,
            trend_direction="HARDENING",
            trend_magnitude_pct=5.2,
            data_window="2025-Q1 to 2025-Q4",
            segment_label="LARGE / TECH",
        )
        state = _make_state_with_mi(mi)
        context = _make_context(state)

        render_market_context(doc, context, _make_ds())

        text = _extract_text(doc)
        assert "Market Context" in text
        assert "15 comparable quotes" in text
        assert "HIGH" in text
        assert "0.0350" in text
        assert "0.0280 - 0.0420" in text
        assert "HARDENING" in text
        assert "LARGE / TECH" in text

    def test_minimal_data_renders_core_metrics(self) -> None:
        """When has_data is True with minimal fields, render core metrics."""
        doc = _make_doc()
        mi = MarketIntelligence(
            has_data=True,
            peer_count=5,
            confidence_level="LOW",
            median_rate_on_line=0.0400,
        )
        state = _make_state_with_mi(mi)
        context = _make_context(state)

        render_market_context(doc, context, _make_ds())

        text = _extract_text(doc)
        assert "Market Context" in text
        assert "5 comparable quotes" in text
        assert "0.0400" in text
        # No CI or trend should appear
        assert "Confidence Interval" not in text

    def test_trend_direction_rendered(self) -> None:
        """Trend direction and magnitude appear in summary and table."""
        doc = _make_doc()
        mi = MarketIntelligence(
            has_data=True,
            peer_count=10,
            confidence_level="MODERATE",
            median_rate_on_line=0.0300,
            trend_direction="SOFTENING",
            trend_magnitude_pct=-3.1,
        )
        state = _make_state_with_mi(mi)
        context = _make_context(state)

        render_market_context(doc, context, _make_ds())

        text = _extract_text(doc)
        assert "softening" in text
        assert "SOFTENING" in text

    def test_no_table_rows_without_median(self) -> None:
        """Table still renders peer count even without median ROL."""
        doc = _make_doc()
        mi = MarketIntelligence(
            has_data=True,
            peer_count=8,
            confidence_level="LOW",
        )
        state = _make_state_with_mi(mi)
        context = _make_context(state)

        render_market_context(doc, context, _make_ds())

        text = _extract_text(doc)
        # Heading and summary should appear
        assert "Market Context" in text
        assert "8 comparable quotes" in text


class TestMarketContextAlerts:
    """Tests for mispricing alert rendering."""

    def test_mispricing_alert_rendered(self) -> None:
        """Mispricing alert text appears in output."""
        doc = _make_doc()
        mi = MarketIntelligence(
            has_data=True,
            peer_count=12,
            confidence_level="HIGH",
            median_rate_on_line=0.0350,
            mispricing_alert=(
                "OVERPRICED vs market: current ROL 0.0450 is "
                "28.6% above median 0.0350 (n=12)"
            ),
        )
        state = _make_state_with_mi(mi)
        context = _make_context(state)

        render_market_context(doc, context, _make_ds())

        text = _extract_text(doc)
        assert "Mispricing Signals" in text
        assert "OVERPRICED" in text
        assert "28.6%" in text

    def test_model_vs_market_alert_rendered(self) -> None:
        """Model-vs-market alert text appears in output."""
        doc = _make_doc()
        mi = MarketIntelligence(
            has_data=True,
            peer_count=10,
            confidence_level="MODERATE",
            median_rate_on_line=0.0400,
            model_vs_market_alert=(
                "MODEL SUGGESTS UNDERPRICED BY MARKET: model indicated "
                "ROL 0.0550 is 37.1% above market median 0.0401"
            ),
        )
        state = _make_state_with_mi(mi)
        context = _make_context(state)

        render_market_context(doc, context, _make_ds())

        text = _extract_text(doc)
        assert "Mispricing Signals" in text
        assert "MODEL SUGGESTS UNDERPRICED" in text

    def test_both_alerts_rendered(self) -> None:
        """Both mispricing and model-vs-market alerts render."""
        doc = _make_doc()
        mi = MarketIntelligence(
            has_data=True,
            peer_count=15,
            confidence_level="HIGH",
            median_rate_on_line=0.0350,
            mispricing_alert="OVERPRICED vs market",
            model_vs_market_alert="MODEL SUGGESTS UNDERPRICED BY MARKET",
        )
        state = _make_state_with_mi(mi)
        context = _make_context(state)

        render_market_context(doc, context, _make_ds())

        text = _extract_text(doc)
        assert "OVERPRICED" in text
        assert "MODEL SUGGESTS UNDERPRICED" in text

    def test_no_alerts_when_absent(self) -> None:
        """No alert section when no mispricing signals."""
        doc = _make_doc()
        mi = MarketIntelligence(
            has_data=True,
            peer_count=10,
            confidence_level="MODERATE",
            median_rate_on_line=0.0350,
        )
        state = _make_state_with_mi(mi)
        context = _make_context(state)

        render_market_context(doc, context, _make_ds())

        text = _extract_text(doc)
        assert "Mispricing Signals" not in text
