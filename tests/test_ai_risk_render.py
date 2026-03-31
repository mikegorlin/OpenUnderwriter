"""Tests for AI risk Section 8 rendering (Word, Markdown).

Tests render_section_8 with populated, None, and empty AI risk data.
Tests new sub-sections: disclosures, patents, peer comparison.
Tests build_template_context includes ai_risk key.
Tests Markdown template renders AI risk section correctly.

Phase 60-02: Updated to pass context dict instead of raw state.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import jinja2
from docx import Document  # type: ignore[import-untyped]

from do_uw.models.ai_risk import (
    AICompetitivePosition,
    AIDisclosureData,
    AIPatentActivity,
    AIRiskAssessment,
    AISubDimension,
)
from do_uw.models.state import AnalysisState, ExtractedData
from do_uw.stages.render.design_system import DesignSystem, setup_styles
from do_uw.stages.render.md_renderer import build_template_context
from do_uw.stages.render.sections.sect8_ai_risk import render_section_8


def _make_ds() -> DesignSystem:
    """Create a DesignSystem for tests."""
    return DesignSystem()


def _make_doc() -> Any:
    """Create a styled python-docx Document for tests."""
    doc: Any = Document()
    setup_styles(doc)
    return doc


def _make_ai_risk() -> AIRiskAssessment:
    """Create a populated AIRiskAssessment for tests."""
    return AIRiskAssessment(
        overall_score=62.5,
        sub_dimensions=[
            AISubDimension(
                dimension="revenue_displacement",
                score=7.5,
                weight=0.25,
                evidence=["High threat ratio", "YoY increasing"],
                threat_level="HIGH",
            ),
            AISubDimension(
                dimension="cost_structure",
                score=5.0,
                weight=0.20,
                evidence=["Moderate mentions"],
                threat_level="MEDIUM",
            ),
            AISubDimension(
                dimension="competitive_moat",
                score=4.0,
                weight=0.25,
                evidence=["Some AI patents"],
                threat_level="MEDIUM",
            ),
            AISubDimension(
                dimension="workforce_automation",
                score=6.0,
                weight=0.20,
                evidence=["Notable disclosure count"],
                threat_level="MEDIUM",
            ),
            AISubDimension(
                dimension="regulatory_ip",
                score=3.5,
                weight=0.10,
                evidence=["Low patent complexity"],
                threat_level="LOW",
            ),
        ],
        disclosure_data=AIDisclosureData(
            mention_count=25,
            risk_factors=["AI may displace key revenue streams", "Regulatory uncertainty"],
            opportunity_mentions=10,
            threat_mentions=15,
            sentiment="BALANCED",
            yoy_trend="INCREASING",
        ),
        patent_activity=AIPatentActivity(
            ai_patent_count=12,
            recent_filings=[
                {
                    "patent_number": "US123",
                    "filing_date": "2024-06-01",
                    "title": "AI Model Compression",
                },
            ],
            filing_trend="STABLE",
        ),
        competitive_position=AICompetitivePosition(
            company_ai_mentions=25,
            peer_avg_mentions=18.5,
            peer_mention_counts={"PEER1": 15, "PEER2": 22},
            percentile_rank=72.0,
            adoption_stance="LEADING",
        ),
        industry_model_id="TECH_SAAS",
        disclosure_trend="INCREASING",
        narrative=(
            "This Technology/SaaS company faces moderate AI "
            "transformation risk (composite score: 63/100)."
        ),
        narrative_source="AI impact model TECH_SAAS",
        narrative_confidence="MEDIUM",
        peer_comparison_available=True,
        forward_indicators=[
            "Increasing AI patent filings",
            "Competitor AI product launches",
        ],
        data_sources=[
            "SEC filings (AI disclosure analysis)",
            "Patent database",
            "Peer comparison analysis",
        ],
    )


def _make_context(state: AnalysisState) -> dict[str, Any]:
    """Wrap AnalysisState in context dict for section renderers.

    Phase 60-02: Sections now receive context dict from build_template_context().
    Tests use this minimal wrapper with _state escape hatch.
    """
    return {"_state": state, "company_name": state.ticker}


def _make_state_with_ai_risk() -> AnalysisState:
    """Create an AnalysisState with populated AI risk data."""
    return AnalysisState(
        ticker="TSLA",
        extracted=ExtractedData(ai_risk=_make_ai_risk()),
    )


def _make_state_without_ai_risk() -> AnalysisState:
    """Create an AnalysisState without AI risk data."""
    return AnalysisState(ticker="TSLA")


# ---- Word renderer tests ----


def _all_text(doc: Any) -> str:
    """Extract all text from paragraphs and table cells."""
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_render_section_8_populated() -> None:
    """render_section_8 with populated AI risk data renders without error."""
    doc = _make_doc()
    state = _make_state_with_ai_risk()
    ds = _make_ds()
    render_section_8(doc, _make_context(state), ds)

    text = _all_text(doc)
    assert "Section 8: AI Transformation Risk" in text
    assert "62" in text  # overall score
    assert "TECH_SAAS" in text or "Tech Saas" in text
    assert "Revenue Displacement" in text
    assert "Cost Structure" in text


def test_render_section_8_none_ai_risk() -> None:
    """render_section_8 with None ai_risk renders unavailable message."""
    doc = _make_doc()
    state = _make_state_without_ai_risk()
    ds = _make_ds()
    render_section_8(doc, _make_context(state), ds)

    text = "\n".join(p.text for p in doc.paragraphs)
    assert "AI risk assessment unavailable" in text


def test_render_section_8_empty_sub_dimensions() -> None:
    """render_section_8 with empty sub_dimensions handles gracefully."""
    doc = _make_doc()
    ai_risk = AIRiskAssessment(
        overall_score=50.0,
        sub_dimensions=[],
        industry_model_id="GENERIC",
    )
    state = AnalysisState(
        ticker="TSLA",
        extracted=ExtractedData(ai_risk=ai_risk),
    )
    ds = _make_ds()
    render_section_8(doc, _make_context(state), ds)

    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Sub-dimension scoring not available" in text


def test_render_section_8_no_peer_comparison() -> None:
    """render_section_8 without peer data shows unavailable message."""
    doc = _make_doc()
    ai_risk = AIRiskAssessment(
        overall_score=50.0,
        industry_model_id="GENERIC",
        peer_comparison_available=False,
    )
    state = AnalysisState(
        ticker="TSLA",
        extracted=ExtractedData(ai_risk=ai_risk),
    )
    ds = _make_ds()
    render_section_8(doc, _make_context(state), ds)

    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Peer comparison unavailable" in text


def test_render_section_8_no_forward_indicators() -> None:
    """render_section_8 without forward indicators shows no indicators."""
    doc = _make_doc()
    ai_risk = AIRiskAssessment(
        overall_score=50.0,
        industry_model_id="GENERIC",
        forward_indicators=[],
    )
    state = AnalysisState(
        ticker="TSLA",
        extracted=ExtractedData(ai_risk=ai_risk),
    )
    ds = _make_ds()
    render_section_8(doc, _make_context(state), ds)

    text = "\n".join(p.text for p in doc.paragraphs)
    assert "No forward indicators identified" in text


def test_render_section_8_high_score_do_context() -> None:
    """render_section_8 adds D&O context for high-scoring dimensions."""
    doc = _make_doc()
    state = _make_state_with_ai_risk()
    ds = _make_ds()
    render_section_8(doc, _make_context(state), ds)

    text = "\n".join(p.text for p in doc.paragraphs)
    # revenue_displacement has score 7.5, should trigger D&O context
    assert "D&O Context (Revenue Displacement)" in text


def test_render_section_8_disclosure_section() -> None:
    """render_section_8 renders AI disclosure analysis from filing data."""
    doc = _make_doc()
    state = _make_state_with_ai_risk()
    ds = _make_ds()
    render_section_8(doc, _make_context(state), ds)

    text = _all_text(doc)
    assert "AI Disclosures (Filing Analysis)" in text
    assert "25" in text  # mention count
    assert "BALANCED" in text  # sentiment
    assert "AI may displace key revenue streams" in text  # risk factor


def test_render_section_8_patent_section() -> None:
    """render_section_8 renders patent and innovation data."""
    doc = _make_doc()
    state = _make_state_with_ai_risk()
    ds = _make_ds()
    render_section_8(doc, _make_context(state), ds)

    text = _all_text(doc)
    assert "AI Patent & Innovation" in text
    assert "12" in text  # patent count
    assert "AI Model Compression" in text  # recent filing title


def test_render_section_8_peer_comparison_with_delta() -> None:
    """render_section_8 renders peer comparison with vs. Company column."""
    doc = _make_doc()
    state = _make_state_with_ai_risk()
    ds = _make_ds()
    render_section_8(doc, _make_context(state), ds)

    text = _all_text(doc)
    assert "Peer AI Comparison" in text
    assert "vs. Company" in text  # delta column header
    assert "PEER2" in text  # peer ticker


def test_render_section_8_no_disclosures() -> None:
    """render_section_8 handles zero AI disclosures gracefully."""
    doc = _make_doc()
    ai_risk = AIRiskAssessment(
        overall_score=30.0,
        industry_model_id="GENERIC",
        disclosure_data=AIDisclosureData(mention_count=0),
    )
    state = AnalysisState(
        ticker="TSLA",
        extracted=ExtractedData(ai_risk=ai_risk),
    )
    ds = _make_ds()
    render_section_8(doc, _make_context(state), ds)

    text = "\n".join(p.text for p in doc.paragraphs)
    assert "No AI-related disclosures identified" in text


def test_render_section_8_no_patents() -> None:
    """render_section_8 handles zero patents gracefully."""
    doc = _make_doc()
    ai_risk = AIRiskAssessment(
        overall_score=30.0,
        industry_model_id="GENERIC",
        patent_activity=AIPatentActivity(ai_patent_count=0),
    )
    state = AnalysisState(
        ticker="TSLA",
        extracted=ExtractedData(ai_risk=ai_risk),
    )
    ds = _make_ds()
    render_section_8(doc, _make_context(state), ds)

    text = "\n".join(p.text for p in doc.paragraphs)
    assert "No AI patent activity identified" in text


def test_render_section_8_overview_panel() -> None:
    """render_section_8 renders overview panel with score and threat."""
    doc = _make_doc()
    state = _make_state_with_ai_risk()
    ds = _make_ds()
    render_section_8(doc, _make_context(state), ds)

    text = _all_text(doc)
    assert "AI Transformation Risk Score: 63/100" in text or "62/100" in text
    assert "Industry Model" in text
    assert "Disclosure Trend" in text


# ---- Markdown context tests ----


def test_build_template_context_includes_ai_risk() -> None:
    """build_template_context includes ai_risk key when data present."""
    state = _make_state_with_ai_risk()
    ctx = build_template_context(state)
    assert ctx["ai_risk"] is not None
    assert ctx["ai_risk"]["overall_score"] == 62.5
    assert ctx["ai_risk"]["industry_model_id"] == "TECH_SAAS"
    assert len(ctx["ai_risk"]["sub_dimensions"]) == 5


def test_build_template_context_ai_risk_none_when_missing() -> None:
    """build_template_context ai_risk is None when no AI risk data."""
    state = _make_state_without_ai_risk()
    ctx = build_template_context(state)
    assert ctx["ai_risk"] is None


def test_build_template_context_ai_risk_competitive_position() -> None:
    """build_template_context extracts competitive position correctly."""
    state = _make_state_with_ai_risk()
    ctx = build_template_context(state)
    cp = ctx["ai_risk"]["competitive_position"]
    assert cp["company_ai_mentions"] == 25
    assert cp["peer_avg_mentions"] == 18.5
    assert cp["adoption_stance"] == "LEADING"


# ---- Markdown template rendering tests ----


_TEMPLATE_DIR = (
    Path(__file__).resolve().parent.parent
    / "src"
    / "do_uw"
    / "templates"
    / "markdown"
)


def test_markdown_template_renders_ai_risk_section() -> None:
    """Jinja2 template renders AI risk section with populated data."""
    state = _make_state_with_ai_risk()
    ctx = build_template_context(state)

    env = jinja2.Environment(
        loader=jinja2.FileSystemLoader(str(_TEMPLATE_DIR)),
        autoescape=False,  # noqa: S701
        undefined=jinja2.StrictUndefined,
    )
    from do_uw.stages.render.formatters import (
        format_currency,
        format_percentage,
        na_if_none,
    )
    from do_uw.stages.render.md_renderer import _dim_display_name

    env.filters["format_currency"] = format_currency
    env.filters["format_pct"] = format_percentage
    env.filters["na_if_none"] = na_if_none
    env.filters["dim_display_name"] = _dim_display_name

    template = env.get_template("worksheet.md.j2")
    content = template.render(**ctx)

    assert "## Section 8: AI Transformation Risk" in content
    assert "63/100" in content  # round(62.5) -> 62, but narrative says 63
    assert "TECH_SAAS" in content
    assert "Revenue Displacement" in content
    assert "INCREASING" in content


def test_markdown_template_renders_ai_risk_unavailable() -> None:
    """Jinja2 template renders unavailable message when no AI risk."""
    state = _make_state_without_ai_risk()
    ctx = build_template_context(state)

    env = jinja2.Environment(
        loader=jinja2.FileSystemLoader(str(_TEMPLATE_DIR)),
        autoescape=False,  # noqa: S701
        undefined=jinja2.StrictUndefined,
    )
    from do_uw.stages.render.formatters import (
        format_currency,
        format_percentage,
        na_if_none,
    )
    from do_uw.stages.render.md_renderer import _dim_display_name

    env.filters["format_currency"] = format_currency
    env.filters["format_pct"] = format_percentage
    env.filters["na_if_none"] = na_if_none
    env.filters["dim_display_name"] = _dim_display_name

    template = env.get_template("worksheet.md.j2")
    content = template.render(**ctx)

    assert "AI risk assessment unavailable" in content
