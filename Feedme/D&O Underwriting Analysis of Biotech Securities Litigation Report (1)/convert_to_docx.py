#!/usr/bin/env python3
"""Convert Markdown files to professionally formatted DOCX documents."""

import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_table(doc, headers, rows):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "1F4E79")
    
    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, "E8F0FE")
    
    return table

def parse_md_table(lines):
    """Parse a markdown table into headers and rows."""
    headers = []
    rows = []
    for i, line in enumerate(lines):
        line = line.strip()
        if not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if i == 0 or (not headers):
            headers = cells
        elif all(set(c.strip()) <= set('-| :') for c in cells):
            continue  # separator row
        else:
            rows.append(cells)
    return headers, rows

def process_markdown_to_docx(md_path, docx_path, title_text):
    """Convert a markdown file to a styled DOCX document."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(51, 51, 51)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    with open(md_path, 'r') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            i += 1
            continue
        
        # Horizontal rule
        if stripped == '---':
            doc.add_paragraph('_' * 60)
            i += 1
            continue
        
        # Headers
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            text = stripped.lstrip('#').strip()
            # Clean markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            
            if level == 1:
                p = doc.add_heading(text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                doc.add_heading(text, level=1)
            elif level == 3:
                doc.add_heading(text, level=2)
            elif level == 4:
                doc.add_heading(text, level=3)
            i += 1
            continue
        
        # Table detection
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            headers, rows = parse_md_table(table_lines)
            if headers and rows:
                # Clean markdown formatting from cells
                clean_headers = [re.sub(r'\*\*(.*?)\*\*', r'\1', h) for h in headers]
                clean_rows = [[re.sub(r'\*\*(.*?)\*\*', r'\1', c) for c in row] for row in rows]
                add_styled_table(doc, clean_headers, clean_rows)
                doc.add_paragraph('')  # spacing
            continue
        
        # Blockquote
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('>').strip()
            quote_text = re.sub(r'\*(.*?)\*', r'\1', quote_text)
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.right_indent = Cm(1.27)
            run = p.add_run(quote_text)
            run.italic = True
            run.font.color.rgb = RGBColor(89, 89, 89)
            i += 1
            continue
        
        # Reference-style links at end
        if re.match(r'^\[\d+\]:', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            # Extract reference number and URL
            match = re.match(r'\[(\d+)\]:\s*(https?://\S+)\s*"?(.*?)"?$', stripped)
            if match:
                ref_num, url, title = match.groups()
                run = p.add_run(f'[{ref_num}] ')
                run.font.size = Pt(9)
                run.bold = True
                if title:
                    run2 = p.add_run(f'{title}. ')
                    run2.font.size = Pt(9)
                run3 = p.add_run(url)
                run3.font.size = Pt(8)
                run3.font.color.rgb = RGBColor(0, 102, 204)
            else:
                run = p.add_run(stripped)
                run.font.size = Pt(9)
            i += 1
            continue
        
        # Bullet points
        if stripped.startswith('*   ') or stripped.startswith('- ') or stripped.startswith('* '):
            text = re.sub(r'^[\*\-]\s+', '', stripped)
            # Handle bold text
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Handle italic
                    italic_parts = re.split(r'(\*.*?\*)', part)
                    for ip in italic_parts:
                        if ip.startswith('*') and ip.endswith('*') and not ip.startswith('**'):
                            run = p.add_run(ip[1:-1])
                            run.italic = True
                        else:
                            p.add_run(ip)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        # Process inline formatting
        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                # Handle italic
                italic_parts = re.split(r'(\*.*?\*)', part)
                for ip in italic_parts:
                    if ip.startswith('*') and ip.endswith('*') and not ip.startswith('**'):
                        run = p.add_run(ip[1:-1])
                        run.italic = True
                    else:
                        # Handle inline code
                        code_parts = re.split(r'(`.*?`)', ip)
                        for cp in code_parts:
                            if cp.startswith('`') and cp.endswith('`'):
                                run = p.add_run(cp[1:-1])
                                run.font.name = 'Consolas'
                                run.font.size = Pt(9.5)
                                run.font.color.rgb = RGBColor(0, 102, 0)
                            else:
                                # Handle citation links
                                cp_clean = re.sub(r'\[(\d+)\]', r'[\1]', cp)
                                p.add_run(cp_clean)
        i += 1
    
    doc.save(docx_path)
    print(f"Saved: {docx_path}")

# Convert both documents
process_markdown_to_docx(
    '/home/ubuntu/output/Swiss_Re_BioTech_Critical_Analysis.md',
    '/home/ubuntu/output/Swiss_Re_BioTech_Critical_Analysis.docx',
    'Critical D&O Underwriting Analysis: Swiss Re Biotech Presentation'
)

process_markdown_to_docx(
    '/home/ubuntu/output/Biotech_DO_Underwriting_Guide.md',
    '/home/ubuntu/output/Biotech_DO_Underwriting_Guide.docx',
    'Modern Biotechnology D&O Underwriting Guide'
)

print("Both documents converted successfully.")
