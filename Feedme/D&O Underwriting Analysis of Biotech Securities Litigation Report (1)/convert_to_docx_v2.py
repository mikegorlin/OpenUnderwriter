#!/usr/bin/env python3
"""Convert Markdown files to professionally formatted DOCX documents."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import re
import os

def create_styles(doc):
    """Set up professional document styles."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    
    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)  # Dark navy
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    
    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(18)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x00, 0x4E, 0x8C)  # Medium blue
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    
    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x00, 0x6B, 0xA6)  # Lighter blue
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)
    
    # Heading 4
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Calibri'
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.color.rgb = RGBColor(0x00, 0x6B, 0xA6)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(4)

def set_cell_shading(cell, color):
    """Set background color of a table cell."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading_elm = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')

def add_formatted_run(paragraph, text):
    """Add a run with inline bold/italic formatting."""
    # Split on bold markers
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Handle italic within non-bold text
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*'):
                    run = paragraph.add_run(ip[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(ip)

def parse_table(lines):
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and not re.match(r'^\|[\s\-:]+\|', line):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)
    return rows

def add_table(doc, rows):
    """Add a professionally formatted table to the document."""
    if not rows:
        return
    
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Format header row
    for j, cell_text in enumerate(rows[0]):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text))
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '002B5C')
    
    # Format data rows
    for i, row_data in enumerate(rows[1:], 1):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Clean markdown formatting
                clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                run = p.add_run(clean_text)
                run.font.size = Pt(10)
                if i % 2 == 0:
                    set_cell_shading(cell, 'F0F4F8')

def md_to_docx(md_path, docx_path):
    """Convert a markdown file to a formatted DOCX document."""
    with open(md_path, 'r') as f:
        content = f.read()
    
    doc = Document()
    create_styles(doc)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            i += 1
            continue
        
        # Horizontal rule
        if stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run('_' * 80)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue
        
        # Headings
        if stripped.startswith('####'):
            text = stripped[4:].strip()
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            doc.add_heading(text, level=4)
            i += 1
            continue
        elif stripped.startswith('###'):
            text = stripped[3:].strip()
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            doc.add_heading(text, level=3)
            i += 1
            continue
        elif stripped.startswith('##'):
            text = stripped[2:].strip()
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            doc.add_heading(text, level=2)
            i += 1
            continue
        elif stripped.startswith('#'):
            text = stripped[1:].strip()
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            doc.add_heading(text, level=1)
            i += 1
            continue
        
        # Tables
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            if rows:
                add_table(doc, rows)
                doc.add_paragraph()  # spacing
            continue
        
        # Blockquotes
        if stripped.startswith('>'):
            text = stripped[1:].strip()
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue
        
        # Bullet points
        if stripped.startswith('* ') or stripped.startswith('- '):
            text = stripped[2:].strip()
            # Remove citation brackets for cleaner look
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_run(p, text)
            i += 1
            continue
        
        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            text = num_match.group(2).strip()
            p = doc.add_paragraph(style='List Number')
            add_formatted_run(p, text)
            i += 1
            continue
        
        # Indented bullets
        if stripped.startswith('  *') or stripped.startswith('  -'):
            text = stripped.lstrip(' *-').strip()
            p = doc.add_paragraph(style='List Bullet 2')
            add_formatted_run(p, text)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_run(p, stripped)
        i += 1
    
    doc.save(docx_path)
    print(f"Saved: {docx_path}")

# Convert both documents
os.makedirs('/home/ubuntu/output', exist_ok=True)

md_to_docx(
    '/home/ubuntu/output/Swiss_Re_BioTech_Critical_Analysis_Expanded.md',
    '/home/ubuntu/output/Swiss_Re_BioTech_Critical_Analysis_Expanded.docx'
)

md_to_docx(
    '/home/ubuntu/output/Biotech_DO_Underwriting_Guide_Expanded.md',
    '/home/ubuntu/output/Biotech_DO_Underwriting_Guide_Expanded.docx'
)

print("All conversions complete.")
